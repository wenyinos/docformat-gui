#!/usr/bin/env python3
"""
公文格式处理工具 - 纸质感极简风格 v2
优化：更大图标、更好排版、卡片式选择
"""

import os
import sys
import threading
import re
import uuid
import ctypes
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from pathlib import Path

# 添加scripts目录到路径
SCRIPT_DIR = Path(__file__).parent / "scripts"
sys.path.insert(0, str(SCRIPT_DIR))

from scripts.analyzer import analyze_punctuation, analyze_numbering, analyze_paragraph_format, analyze_font
from scripts.formatter import format_document, PRESETS


_DND_DISABLED_REASON = ""
_DPI_AWARENESS_SET = False
_UI_DENSITY = 1.0


def _enable_windows_high_dpi():
    """尽早启用 DPI 感知，兼容 Win7 到 Win11 的缩放行为。"""
    global _DPI_AWARENESS_SET
    if _DPI_AWARENESS_SET or os.name != 'nt':
        return
    try:
        ctypes.windll.user32.SetProcessDpiAwarenessContext(ctypes.c_void_p(-4))
        _DPI_AWARENESS_SET = True
        return
    except Exception:
        pass
    try:
        ctypes.windll.shcore.SetProcessDpiAwareness(2)
        _DPI_AWARENESS_SET = True
        return
    except Exception:
        pass
    try:
        ctypes.windll.user32.SetProcessDPIAware()
        _DPI_AWARENESS_SET = True
    except Exception:
        pass


def _configure_tk_high_dpi(root):
    """让 Tk 字体、控件和画布尺寸跟随当前显示器 DPI。"""
    global _UI_DENSITY
    try:
        dpi = float(root.winfo_fpixels('1i'))
        scaling = max(1.0, min(4.0, dpi / 72.0))
        root.tk.call('tk', 'scaling', scaling)
        _UI_DENSITY = max(1.0, min(2.5, dpi / 96.0))
        if 'Theme' in globals():
            Theme.apply_density(_UI_DENSITY)
        root._docformat_dpi = dpi
        root._docformat_scaling = scaling
    except Exception:
        root._docformat_dpi = 96.0
        root._docformat_scaling = 96.0 / 72.0


def ui_px(value):
    """把固定像素尺寸转换为适合当前显示器缩放的尺寸。"""
    return max(1, int(round(float(value) * _UI_DENSITY)))


_enable_windows_high_dpi()


def _is_windows_elevated():
    """返回当前进程是否以管理员权限运行。"""
    if os.name != 'nt':
        return False
    try:
        return bool(ctypes.windll.shell32.IsUserAnAdmin())
    except Exception:
        return False


def _should_enable_drag_drop():
    """判断当前运行环境是否启用文件拖拽。"""
    global _DND_DISABLED_REASON
    _DND_DISABLED_REASON = ""
    if os.environ.get('DOCFORMAT_DISABLE_DND') == '1':
        _DND_DISABLED_REASON = "已通过 DOCFORMAT_DISABLE_DND 禁用拖拽"
        return False
    if _is_windows_elevated():
        _DND_DISABLED_REASON = "当前为管理员模式，Windows 会拦截普通资源管理器拖拽"
        return False
    # v1.8.1: macOS 打包版先禁用 tkdnd，优先保证 .app/.dmg 可正常打开。
    # 源码运行仍可用拖拽，后续待打包稳定后再恢复。
    if getattr(sys, 'frozen', False) and sys.platform == 'darwin':
        _DND_DISABLED_REASON = "macOS 打包版暂时禁用拖拽"
        return False
    return True

# 拖拽支持（可选，没有安装时自动降级）
try:
    from tkinterdnd2 import TkinterDnD, DND_FILES, COPY
    _DND_AVAILABLE = _should_enable_drag_drop()
except Exception as e:
    _DND_DISABLED_REASON = f"拖拽运行库不可用：{e}"
    _DND_AVAILABLE = False

__version__ = '1.8.5'

def _open_file(path):
    """跨平台打开文件"""
    try:
        if os.name == 'nt':
            os.startfile(path)
        elif sys.platform == 'darwin':
            import subprocess
            subprocess.Popen(['open', path])
        else:
            import subprocess
            subprocess.Popen(['xdg-open', path])
    except Exception:
        pass


# ===== 设计系统 =====
class Theme:
    # 纸质色调
    BG = '#FBF9F6'              # 温暖米白纸张
    CARD = '#FFFFFF'            # 纯白卡片
    CARD_ALT = '#F7F4EF'        # 米黄卡片（推荐区）
    INPUT_BG = '#F2EFE9'        # 输入框背景（稍深米色）
    SURFACE = '#FFFDF9'         # 轻质面板
    
    # 陶土红
    PRIMARY = '#BC4B26'         # 朱砂/印泥色
    PRIMARY_HOVER = '#A3421F'   # 悬停加深
    PRIMARY_LIGHT = '#F9F0EC'   # 极淡红
    PRIMARY_SOFT = '#E7A28A'    # 柔和强调色
    ACCENT = '#2F6F73'          # 青绿色，用于开关选中状态
    ACCENT_LIGHT = '#EAF5F3'    # 极淡青绿色
    
    # 文字
    TEXT = '#2E2E2E'            # 深炭灰
    TEXT_SECONDARY = '#6B6B6B'  # 次要文字
    TEXT_MUTED = '#A0A0A0'      # 禁用/占位
    
    # 边框与分隔
    BORDER = '#E8E4DE'          # 温暖灰边框
    BORDER_LIGHT = '#F0EDE8'    # 更浅边框
    BORDER_SELECTED = '#BC4B26' # 选中边框
    SHADOW = '#E4DED5'          # 轻阴影
    
    # 日志区
    LOG_BG = '#1A1A1A'
    LOG_TEXT = '#C8C8C8'
    LOG_SUCCESS = '#7CB87C'
    LOG_WARNING = '#D4A656'
    LOG_ERROR = '#CF6B6B'
    
    # 字体 - 宋体优先
    if sys.platform == 'darwin':
        # 优先尝试 Windows 字体（用户可能已安装），再回退到 macOS 系统字体
        FONT_SERIF = ('宋体', 'SimSun', 'STSong', 'Songti SC', 'PingFang SC', 'Heiti SC', 'serif')
    else:
        FONT_SERIF = ('Noto Serif SC', 'Source Han Serif SC', 'SimSun', 'PMingLiU', 'serif')
    
    # 间距
    SPACE_XS = 4
    SPACE_SM = 8
    SPACE_MD = 16
    SPACE_LG = 24
    SPACE_XL = 40

    _BASE_SPACING = (4, 8, 16, 24, 40)

    @classmethod
    def apply_density(cls, density):
        values = [max(1, int(round(value * density))) for value in cls._BASE_SPACING]
        (
            cls.SPACE_XS,
            cls.SPACE_SM,
            cls.SPACE_MD,
            cls.SPACE_LG,
            cls.SPACE_XL,
        ) = values


def get_font(size=12, weight='normal'):
    """获取宋体字体"""
    return (Theme.FONT_SERIF[0], size, weight)


# ===== 配置管理 =====
import json

def _migrate_legacy_config(config_dir):
    """把旧版放在 exe 同目录的配置迁移到新的用户目录。

    仅当新目录还没有配置、且旧目录确实存在配置时才迁移，保证现有用户升级后不丢设置。
    迁移失败不应阻止程序启动。
    """
    try:
        import shutil
        new_file = Path(config_dir) / "custom_settings.json"
        if new_file.exists():
            return
        legacy_dir = Path(sys.executable).parent
        legacy_file = legacy_dir / "custom_settings.json"
        if not legacy_file.exists():
            return
        if legacy_file.resolve() == new_file.resolve():
            return
        shutil.copy2(legacy_file, new_file)
        # 一并迁移备份文件（如果有）
        legacy_bak = legacy_file.with_suffix('.json.v1bak')
        if legacy_bak.exists():
            shutil.copy2(legacy_bak, new_file.with_suffix('.json.v1bak'))
    except Exception:
        pass


def _get_config_dir():
    """获取配置文件存放目录（确保可写）"""
    if sys.platform == 'darwin':
        # macOS 打包后：.app 包内部只读，配置文件放到用户目录
        config_dir = Path.home() / 'Library' / 'Application Support' / 'DocFormatter'
    elif sys.platform == 'win32':
        # Windows 打包后：写到 %APPDATA%。
        # 不再写在 exe 同目录——一旦程序被装进 Program Files 等只读位置，
        # 写配置就会失败，逼用户以管理员身份运行；而管理员模式下系统会拦截
        # 普通权限程序（资源管理器）拖来的文件，导致拖拽功能失效。
        base = os.environ.get('APPDATA') or str(Path.home() / 'AppData' / 'Roaming')
        config_dir = Path(base) / 'DocFormatter'
    else:
        # Linux 打包后：遵循 XDG 规范，避免同样的只读目录 / 权限问题
        base = os.environ.get('XDG_CONFIG_HOME') or str(Path.home() / '.config')
        config_dir = Path(base) / 'DocFormatter'

    try:
        config_dir.mkdir(parents=True, exist_ok=True)
    except Exception:
        # 极端情况下用户目录不可写，退回到 exe 同目录，至少不致崩溃
        config_dir = Path(sys.executable).parent
    _migrate_legacy_config(config_dir)
    return config_dir

CONFIG_FILE = _get_config_dir() / "custom_settings.json"

# 字体列表 —— 在 Tk 根窗口创建后由 _init_system_fonts() 填充
COMMON_FONTS_CN = []
COMMON_FONTS_EN = []

# 公文常用字体（作为优先推荐，始终显示在列表顶部）
_PRIORITY_FONTS = [
    '仿宋_GB2312', '仿宋_GB32312', '仿宋', '宋体', '黑体',
    '楷体_GB2312', '楷体_GB32312', '楷体',
    '方正小标宋简体', '方正仿宋_GBK', '华文仿宋', '华文中宋',
    'Times New Roman', 'Arial', 'Calibri', 'Cambria',
]

PAGE_NUMBER_STYLE_OPTIONS = {
    'dash': '— 1 —（公文标准）',
    'plain': '1（纯数字）',
    'page_text': '第 1 页',
    'page_total': '1 / 总页数',
}
PAGE_NUMBER_POSITION_OPTIONS = {
    'outside': '奇数页右、偶数页左（外侧）',
    'center': '全部居中',
    'right': '全部居右',
    'left': '全部居左',
}

def _init_system_fonts():
    """
    从系统读取所有已安装字体，填充 COMMON_FONTS_CN 和 COMMON_FONTS_EN。
    必须在 tk.Tk() 根窗口创建之后调用，且只调用一次。
    """
    global COMMON_FONTS_CN, COMMON_FONTS_EN
    if COMMON_FONTS_CN:
        return
    try:
        import tkinter.font as tkfont
        all_fonts = sorted(
            f for f in set(tkfont.families())
            if f and not f.startswith('@')
        )
        priority = [f for f in _PRIORITY_FONTS if f in set(all_fonts)]
        rest = [f for f in all_fonts if f not in set(_PRIORITY_FONTS)]
        combined = priority + rest
        COMMON_FONTS_CN = combined
        COMMON_FONTS_EN = combined
    except Exception:
        COMMON_FONTS_CN = list(_PRIORITY_FONTS)
        COMMON_FONTS_EN = list(_PRIORITY_FONTS)

# 字号对照表
FONT_SIZES = [
    ('初号', 42), ('小初', 36), ('一号', 26), ('小一', 24),
    ('二号', 22), ('小二', 18), ('三号', 16), ('小三', 15),
    ('四号', 14), ('小四', 12), ('五号', 10.5), ('小五', 9),
]

DEFAULT_CUSTOM_SETTINGS = {
    'name': '自定义格式',
    'page': {'top': 3.7, 'bottom': 3.5, 'left': 2.8, 'right': 2.6},
    'title': {
        'font_cn': '方正小标宋简体', 'font_en': 'Times New Roman',
        'size': 22, 'bold': False, 'align': 'center', 'indent': 0,
        'line_spacing': 28, 'space_before': 0, 'space_after': 0
    },
    'recipient': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'left', 'indent': 0,
        'line_spacing': 28, 'space_before': 0, 'space_after': 0
    },
    'heading1': {
        'font_cn': '黑体', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'left', 'indent': 32,
        'line_spacing': 28, 'space_before': 0, 'space_after': 0
    },
    'heading2': {
        'font_cn': '楷体_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'left', 'indent': 32,
        'line_spacing': 28, 'space_before': 0, 'space_after': 0
    },
    'heading3': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': True, 'align': 'left', 'indent': 32,
        'line_spacing': 28, 'space_before': 0, 'space_after': 0
    },
    'heading4': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'left', 'indent': 32,
        'line_spacing': 28, 'space_before': 0, 'space_after': 0
    },
    'body': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'justify',
        'indent': 32, 'line_spacing': 28, 'space_before': 0, 'space_after': 0
    },
    'signature': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'right', 'indent': 0,
        'line_spacing': 28, 'space_before': 0, 'space_after': 0
    },
    'date': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'right', 'indent': 0,
        'line_spacing': 28, 'space_before': 0, 'space_after': 0
    },
    'attachment': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'justify', 'indent': 0,
        'line_spacing': 28, 'space_before': 0, 'space_after': 0
    },
    'closing': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 16, 'bold': False, 'align': 'left', 'indent': 32,
        'line_spacing': 28, 'space_before': 0, 'space_after': 0
    },
    'table': {
        'font_cn': '仿宋_GB2312', 'font_en': 'Times New Roman',
        'size': 12, 'bold': False, 'line_spacing': 22,
        'first_line_indent': 0, 'header_bold': True, 'smart_align': False
    },
    'space_handling': 'remove_all',
    'first_line_bold': False,
    'bold_serial': True,
    'split_heading_at_punct': False,
    'deep_clean': False,
    'page_number': True,
    'page_number_font': '宋体',
    'page_number_size': 14,
    'page_number_style': 'dash',
    'page_number_position': 'outside',
    'page_number_offset_mm': 7,
    'replace_existing_page_number': True,
}


# v1.8.0: 配置文件 schema 版本
CONFIG_SCHEMA_VERSION = 2

# 内置只读预设的 id（与 PRESETS dict key 对应）
BUILTIN_PRESET_IDS = ('official', 'academic', 'legal')


def _make_default_user_preset():
    """生成一份新的"用户自定义"预设（带新 uuid）。"""
    import copy
    p = copy.deepcopy(DEFAULT_CUSTOM_SETTINGS)
    p['id'] = str(uuid.uuid4())
    p['name'] = '我的自定义格式'
    p['is_builtin'] = False
    return p


def _make_empty_config():
    """生成一份全新的、空白的配置文件结构。"""
    return {
        'schema_version': CONFIG_SCHEMA_VERSION,
        'active_preset_id': None,   # 当前选中的自定义预设 id
        'presets': [_make_default_user_preset()],
    }


def _migrate_legacy_config(legacy_data):
    """把 v1.7.x 及更早的老格式配置迁移成 v1.8.0+ 的新 schema。"""
    import copy
    if not isinstance(legacy_data, dict):
        return _make_empty_config()
    
    # 老配置就是一份完整的 settings dict（无 schema_version 字段）
    migrated_preset = _merge_settings(DEFAULT_CUSTOM_SETTINGS, copy.deepcopy(legacy_data))
    migrated_preset['id'] = str(uuid.uuid4())
    migrated_preset['name'] = legacy_data.get('name', '我的自定义格式')
    migrated_preset['is_builtin'] = False
    
    # 备份老配置（防意外）
    backup_path = CONFIG_FILE.with_suffix('.json.v1bak')
    try:
        with open(backup_path, 'w', encoding='utf-8') as f:
            json.dump(legacy_data, f, ensure_ascii=False, indent=2)
        print(f"[信息] 老配置已备份到 {backup_path}")
    except Exception:
        pass
    
    return {
        'schema_version': CONFIG_SCHEMA_VERSION,
        'active_preset_id': migrated_preset['id'],
        'presets': [migrated_preset],
    }


def _merge_settings(defaults, custom):
    merged = {}
    for key, value in defaults.items():
        if key in custom:
            if isinstance(value, dict) and isinstance(custom.get(key), dict):
                merged[key] = _merge_settings(value, custom[key])
            else:
                merged[key] = custom[key]
        else:
            merged[key] = value
    return merged


def _ensure_page_number_defaults(preset):
    """补齐新版页码配置，并兼容旧版 footer_distance 字段。"""
    if not isinstance(preset, dict):
        return preset
    preset.setdefault('page_number', True)
    preset.setdefault('page_number_font', '宋体')
    preset.setdefault('page_number_size', 14)
    preset.setdefault('page_number_style', 'dash')
    preset.setdefault('page_number_position', 'outside')
    preset.setdefault('page_number_offset_mm', 7)
    preset.setdefault('replace_existing_page_number', True)
    return preset


def load_custom_settings():
    """加载用户自定义设置，并修复旧配置缺失的 active preset。"""
    if not CONFIG_FILE.exists():
        config = _make_empty_config()
    else:
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
        except Exception as e:
            print(f"[警告] 配置文件读取失败，使用默认: {e}")
            data = None

        if isinstance(data, dict) and data.get('schema_version') == CONFIG_SCHEMA_VERSION:
            config = dict(data)
        else:
            config = _migrate_legacy_config(data)

    presets = config.get('presets') or []
    if not presets:
        presets = [_make_default_user_preset()]

    merged_presets = []
    for preset in presets:
        merged = _merge_settings(DEFAULT_CUSTOM_SETTINGS, preset)
        merged['id'] = preset.get('id') or str(uuid.uuid4())
        merged['name'] = preset.get('name') or '我的自定义格式'
        merged['is_builtin'] = bool(preset.get('is_builtin', False))
        _ensure_page_number_defaults(merged)
        merged_presets.append(merged)

    config['schema_version'] = CONFIG_SCHEMA_VERSION
    config['presets'] = merged_presets
    active_id = config.get('active_preset_id')
    if not active_id or not any(p.get('id') == active_id for p in merged_presets):
        config['active_preset_id'] = merged_presets[0].get('id')
    return config


def save_custom_settings(config):
    """保存配置文件。
    
    Args:
        config: 完整的 config dict（包含 schema_version、presets 等）
    """
    try:
        # 兼容：如果调用方还在传旧的 settings dict（无 schema_version），
        # 自动包装成单 preset 的新 schema
        if 'schema_version' not in config:
            preset = {
                **config,
                'id': str(uuid.uuid4()),
                'name': config.get('name', '我的自定义格式'),
                'is_builtin': False,
            }
            _ensure_page_number_defaults(preset)
            config = {
                'schema_version': CONFIG_SCHEMA_VERSION,
                'active_preset_id': preset['id'],
                'presets': [preset]
            }
        else:
            config['schema_version'] = CONFIG_SCHEMA_VERSION
            presets = config.get('presets') or []
            for preset in presets:
                _ensure_page_number_defaults(preset)
            if presets and not config.get('active_preset_id'):
                config['active_preset_id'] = presets[0].get('id')
        
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        return CONFIG_FILE
    except Exception as e:
        print(f"[错误] 保存配置失败: {e}")


def get_active_user_preset(config):
    """根据 active_preset_id 从 config 中取出当前选中的用户预设。
    
    如果 id 不存在或没有 presets，返回第一个，再不行返回新建的默认预设。
    """
    presets = config.get('presets', [])
    active_id = config.get('active_preset_id')
    if active_id:
        for p in presets:
            if p.get('id') == active_id:
                return p
    if presets:
        return presets[0]
    return _make_default_user_preset()


def _fit_dialog_to_screen(dialog, parent, desired_w, desired_h, min_w, min_h):
    """让弹窗按屏幕可用空间自动取尺寸并居中显示。"""
    dialog.update_idletasks()
    screen_w = dialog.winfo_screenwidth()
    screen_h = dialog.winfo_screenheight()

    max_w = max(520, screen_w - 20)
    max_h = max(420, screen_h - 70)
    win_w = min(desired_w, max_w)
    win_h = min(desired_h, max_h)

    dialog.minsize(min(min_w, win_w), min(min_h, win_h))

    try:
        parent.update_idletasks()
        parent_w = parent.winfo_width()
        parent_h = parent.winfo_height()
        parent_x = parent.winfo_x()
        parent_y = parent.winfo_y()
        x = parent_x + (parent_w - win_w) // 2
        y = parent_y + (parent_h - win_h) // 2
    except Exception:
        x = (screen_w - win_w) // 2
        y = (screen_h - win_h) // 2

    x = max(0, min(x, screen_w - win_w))
    y = max(0, min(y, screen_h - win_h))
    dialog.geometry(f"{int(win_w)}x{int(win_h)}+{int(x)}+{int(y)}")



# ===== 快速设置中，正文字体联动的元素 =====
BODY_FONT_GROUP = ['body', 'heading3', 'heading4', 'closing', 'attachment', 'signature', 'date']


class CustomSettingsDialog(tk.Toplevel):
    """自定义格式设置弹窗 - 快速设置 + 高级设置（可折叠）"""
    
    def __init__(self, parent, on_save=None):
        super().__init__(parent)
        self.dialog = self
        
        self.on_save = on_save
        self._config = load_custom_settings()
        self.settings = get_active_user_preset(self._config)
        if not self._config.get('active_preset_id') and self.settings.get('id'):
            self._config['active_preset_id'] = self.settings['id']
        self._adv_vars = {}  # 高级模式的变量存储
        
        # 窗口设置
        self.title("自定义格式设置")
        self.configure(bg=Theme.BG)
        self.resizable(True, True)
        
        # 模态窗口
        self.transient(parent)
        self.grab_set()
        self.protocol("WM_DELETE_WINDOW", self._on_close)
        
        _fit_dialog_to_screen(
            self, parent,
            desired_w=1600, desired_h=1100,
            min_w=1040, min_h=720
        )
        
        self._create_widgets()
        self._refresh_preset_list()
        self.update_idletasks()   # 确保所有控件完成布局
        self._load_values()
        self.after_idle(self._load_values)  # 事件循环空闲后再刷新一次，兜底
    
    # ==================== 界面构建 ====================

    def _create_inline_toggle(self, parent, text, variable, font_size=11):
        """创建自绘小复选框，避免系统主题导致勾选框显示异常。"""
        return InlineToggle(
            parent,
            text=text,
            variable=variable,
            box_size=22,
            font=get_font(font_size),
            bg=Theme.BG,
        )
    
    def _create_widgets(self):
        """创建控件 - 快速设置 + 可折叠高级设置"""
        # ===== 顶部标题 + 按钮（固定） =====
        header = tk.Frame(self, bg=Theme.BG)
        header.pack(fill='x', padx=20, pady=(15, 5))
        
        tk.Label(
            header, text="⚙️ 自定义格式设置", font=get_font(16, 'bold'),
            bg=Theme.BG, fg=Theme.TEXT
        ).pack(side='left')
        
        # 保存按钮（顶部）
        save_top = tk.Frame(header, bg=Theme.PRIMARY, cursor='hand2')
        save_top.pack(side='right')
        save_top_label = tk.Label(
            save_top, text="  保存设置  ", font=get_font(12, 'bold'),
            bg=Theme.PRIMARY, fg='white', pady=6, cursor='hand2'
        )
        save_top_label.pack()
        for w in [save_top, save_top_label]:
            w.bind('<Button-1>', lambda e: self._save())
            w.bind('<Enter>', lambda e: (save_top.configure(bg=Theme.PRIMARY_HOVER), save_top_label.configure(bg=Theme.PRIMARY_HOVER)))
            w.bind('<Leave>', lambda e: (save_top.configure(bg=Theme.PRIMARY), save_top_label.configure(bg=Theme.PRIMARY)))
        
        cancel_top = tk.Label(
            header, text="取消", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY, cursor='hand2', padx=10
        )
        cancel_top.pack(side='right', padx=(0, 10))
        cancel_top.bind('<Button-1>', lambda e: self._on_close())

        self._build_preset_bar(self)
        
        # ===== 滚动区域 =====
        scroll_container = tk.Frame(self, bg=Theme.BG)
        scroll_container.pack(fill='both', expand=True, padx=5)
        
        self.canvas = tk.Canvas(scroll_container, bg=Theme.BG, highlightthickness=0)
        scrollbar = tk.Scrollbar(scroll_container, orient='vertical', command=self.canvas.yview)
        
        self.canvas.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side='right', fill='y')
        self.canvas.pack(side='left', fill='both', expand=True)
        
        self.content_frame = tk.Frame(self.canvas, bg=Theme.BG)
        self.canvas_window = self.canvas.create_window((0, 0), window=self.content_frame, anchor='nw')
        
        self.content_frame.bind('<Configure>', self._on_frame_configure)
        self.canvas.bind('<Configure>', self._on_canvas_configure)
        self.canvas.bind('<Enter>', lambda e: self._bind_mousewheel())
        self.canvas.bind('<Leave>', lambda e: self._unbind_mousewheel())
        
        main = self.content_frame
        pad_x = 15
        
        # ============================================================
        #  快速设置（始终显示）
        # ============================================================
        
        # --- 页面边距 ---
        self._create_section(main, "📄 页面边距 (cm)", pad_x)
        margin_frame = tk.Frame(main, bg=Theme.BG)
        margin_frame.pack(fill='x', pady=(0, 12), padx=pad_x)
        
        self.margin_vars = {}
        margins = [('top', '上'), ('bottom', '下'), ('left', '左'), ('right', '右')]
        for i, (key, label) in enumerate(margins):
            col = i % 4
            f = tk.Frame(margin_frame, bg=Theme.BG)
            f.grid(row=0, column=col, sticky='w', padx=(0, 15), pady=2)
            tk.Label(f, text=f"{label}:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, anchor='e').pack(side='left')
            var = tk.StringVar(value=str(self.settings.get('page', {}).get(key, 2.5)))
            self.margin_vars[key] = var
            tk.Entry(f, textvariable=var, font=get_font(11), width=6, relief='solid', bd=1).pack(side='left', padx=3)
        
        # --- 标题格式 ---
        self._create_section(main, "📝 标题", pad_x)
        title_frame = tk.Frame(main, bg=Theme.BG)
        title_frame.pack(fill='x', pady=(0, 12), padx=pad_x)
        
        row_t = tk.Frame(title_frame, bg=Theme.BG)
        row_t.pack(fill='x', pady=2)
        
        tk.Label(row_t, text="字体:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=6, anchor='e').pack(side='left')
        self.title_font_var = tk.StringVar()
        self._create_combobox(row_t, self.title_font_var, COMMON_FONTS_CN, width=16,
                              initial_value=self.settings.get('title', {}).get('font_cn', '方正小标宋简体')).pack(side='left', padx=3)
        
        tk.Label(row_t, text="字号:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=5, anchor='e').pack(side='left', padx=(10, 0))
        self.title_size_var = tk.StringVar()
        self._create_combobox(row_t, self.title_size_var, [f"{name}({pt}pt)" for name, pt in FONT_SIZES], width=11,
                              initial_value=self._size_display(self.settings.get('title', {}).get('size', 22))).pack(side='left', padx=3)
        
        tk.Label(row_t, text="行距:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=5, anchor='e').pack(side='left', padx=(10, 0))
        self.title_line_spacing_var = tk.StringVar(value=str(self.settings.get('title', {}).get('line_spacing', 28) or ''))
        tk.Entry(row_t, textvariable=self.title_line_spacing_var, font=get_font(11), width=5, relief='solid', bd=1).pack(side='left', padx=3)
        tk.Label(row_t, text="磅", font=get_font(10), bg=Theme.BG, fg=Theme.TEXT_MUTED).pack(side='left')
        
        self.title_bold_var = tk.BooleanVar(value=self.settings.get('title', {}).get('bold', False))
        self._create_inline_toggle(
            row_t, "加粗", self.title_bold_var
        ).pack(side='left', padx=(10, 0))
        
        # --- 一级标题 / 二级标题 ---
        self._create_section(main, "🔤 各级标题字体", pad_x)
        heading_frame = tk.Frame(main, bg=Theme.BG)
        heading_frame.pack(fill='x', pady=(0, 12), padx=pad_x)
        
        row_h1 = tk.Frame(heading_frame, bg=Theme.BG)
        row_h1.pack(fill='x', pady=2)
        tk.Label(row_h1, text="一级(一、):", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=10, anchor='e').pack(side='left')
        self.h1_font_var = tk.StringVar()
        self._create_combobox(row_h1, self.h1_font_var, COMMON_FONTS_CN, width=16,
                              initial_value=self.settings.get('heading1', {}).get('font_cn', '黑体')).pack(side='left', padx=3)
        tk.Label(row_h1, text="字号:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=5, anchor='e').pack(side='left', padx=(10, 0))
        self.h1_size_var = tk.StringVar()
        self._create_combobox(row_h1, self.h1_size_var, [f"{name}({pt}pt)" for name, pt in FONT_SIZES], width=11,
                              initial_value=self._size_display(self.settings.get('heading1', {}).get('size', 16))).pack(side='left', padx=3)
        
        self.h1_bold_var = tk.BooleanVar(value=self.settings.get('heading1', {}).get('bold', False))
        self._create_inline_toggle(
            row_h1, "加粗", self.h1_bold_var
        ).pack(side='left', padx=(10, 0))
        
        row_h2 = tk.Frame(heading_frame, bg=Theme.BG)
        row_h2.pack(fill='x', pady=2)
        tk.Label(row_h2, text="二级((一)):", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=10, anchor='e').pack(side='left')
        self.h2_font_var = tk.StringVar()
        self._create_combobox(row_h2, self.h2_font_var, COMMON_FONTS_CN, width=16,
                              initial_value=self.settings.get('heading2', {}).get('font_cn', '楷体_GB2312')).pack(side='left', padx=3)
        tk.Label(row_h2, text="字号:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=5, anchor='e').pack(side='left', padx=(10, 0))
        self.h2_size_var = tk.StringVar()
        self._create_combobox(row_h2, self.h2_size_var, [f"{name}({pt}pt)" for name, pt in FONT_SIZES], width=11,
                              initial_value=self._size_display(self.settings.get('heading2', {}).get('size', 16))).pack(side='left', padx=3)
        
        self.h2_bold_var = tk.BooleanVar(value=self.settings.get('heading2', {}).get('bold', False))
        self._create_inline_toggle(
            row_h2, "加粗", self.h2_bold_var
        ).pack(side='left', padx=(10, 0))
        
        # --- 正文格式 ---
        self._create_section(main, "📖 正文格式", pad_x)
        body_frame = tk.Frame(main, bg=Theme.BG)
        body_frame.pack(fill='x', pady=(0, 12), padx=pad_x)
        
        row_b1 = tk.Frame(body_frame, bg=Theme.BG)
        row_b1.pack(fill='x', pady=2)
        tk.Label(row_b1, text="字体:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=6, anchor='e').pack(side='left')
        self.body_font_var = tk.StringVar()
        self._create_combobox(row_b1, self.body_font_var, COMMON_FONTS_CN, width=16,
                              initial_value=self.settings.get('body', {}).get('font_cn', '仿宋_GB2312')).pack(side='left', padx=3)
        
        tk.Label(row_b1, text="字号:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=5, anchor='e').pack(side='left', padx=(10, 0))
        self.body_size_var = tk.StringVar()
        self._create_combobox(row_b1, self.body_size_var, [f"{name}({pt}pt)" for name, pt in FONT_SIZES], width=11,
                              initial_value=self._size_display(self.settings.get('body', {}).get('size', 16))).pack(side='left', padx=3)
        
        tk.Label(row_b1, text="行距:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=5, anchor='e').pack(side='left', padx=(10, 0))
        self.line_spacing_var = tk.StringVar(value=str(self.settings.get('body', {}).get('line_spacing', 28) or ''))
        tk.Entry(row_b1, textvariable=self.line_spacing_var, font=get_font(11), width=5, relief='solid', bd=1).pack(side='left', padx=3)
        tk.Label(row_b1, text="磅", font=get_font(10), bg=Theme.BG, fg=Theme.TEXT_MUTED).pack(side='left')
        
        self.body_bold_var = tk.BooleanVar(value=self.settings.get('body', {}).get('bold', False))
        self._create_inline_toggle(
            row_b1, "加粗", self.body_bold_var
        ).pack(side='left', padx=(10, 0))
        
        row_b_en = tk.Frame(body_frame, bg=Theme.BG)
        row_b_en.pack(fill='x', pady=2)
        tk.Label(
            row_b_en, text="英数字体:", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=6, anchor='e'
        ).pack(side='left')
        self.global_font_en_var = tk.StringVar()
        self._create_combobox(
            row_b_en, self.global_font_en_var, COMMON_FONTS_EN, width=16,
            initial_value=self.settings.get('body', {}).get('font_en', 'Times New Roman')
        ).pack(side='left', padx=3)
        tk.Label(
            row_b_en,
            text="  ⓘ 全文英文/数字统一使用此字体",
            font=get_font(9), bg=Theme.BG, fg=Theme.TEXT_MUTED
        ).pack(side='left', padx=(10, 0))
        
        row_b2 = tk.Frame(body_frame, bg=Theme.BG)
        row_b2.pack(fill='x', pady=2)
        tk.Label(row_b2, text="首行缩进:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=8, anchor='e').pack(side='left')
        self.indent_var = tk.StringVar()
        _body = self.settings.get('body', {})
        _indent = _body.get('indent', 32)
        _bsize = _body.get('size', 16) or 16
        _indent_chars = int(_indent / _bsize) if _bsize else 2
        self._create_combobox(row_b2, self.indent_var, ['0字符', '2字符', '4字符'], width=8,
                              initial_value=f'{_indent_chars}字符').pack(side='left', padx=3)

        row_spacing = tk.Frame(body_frame, bg=Theme.BG)
        row_spacing.pack(fill='x', pady=2)
        tk.Label(
            row_spacing, text="段前/段后:", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=8, anchor='e'
        ).pack(side='left')
        self.space_before_var = tk.StringVar(value='0')
        tk.Entry(
            row_spacing, textvariable=self.space_before_var,
            font=get_font(11), width=4, relief='solid', bd=1
        ).pack(side='left', padx=3)
        tk.Label(row_spacing, text="/ ", font=get_font(10),
                 bg=Theme.BG, fg=Theme.TEXT_MUTED).pack(side='left')
        self.space_after_var = tk.StringVar(value='0')
        tk.Entry(
            row_spacing, textvariable=self.space_after_var,
            font=get_font(11), width=4, relief='solid', bd=1
        ).pack(side='left', padx=3)
        tk.Label(row_spacing, text="磅  ⓘ 全文段前/段后统一设置",
                 font=get_font(9), bg=Theme.BG, fg=Theme.TEXT_MUTED).pack(side='left', padx=(4,0))
        
        tk.Label(row_b2, text="  ⓘ 正文字体/字号同时应用于: 三/四级标题、落款、附件、结束语",
                 font=get_font(9), bg=Theme.BG, fg=Theme.TEXT_MUTED).pack(side='left', padx=(10, 0))
        
        # --- 表格格式 ---
        self._create_section(main, "📊 表格格式", pad_x)
        table_frame = tk.Frame(main, bg=Theme.BG)
        table_frame.pack(fill='x', pady=(0, 12), padx=pad_x)
        
        row_tbl1 = tk.Frame(table_frame, bg=Theme.BG)
        row_tbl1.pack(fill='x', pady=2)
        
        tk.Label(row_tbl1, text="字体:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=6, anchor='e').pack(side='left')
        self.table_font_var = tk.StringVar()
        self._create_combobox(row_tbl1, self.table_font_var, COMMON_FONTS_CN, width=16,
                              initial_value=self.settings.get('table', {}).get('font_cn', '仿宋_GB2312')).pack(side='left', padx=3)
        
        tk.Label(row_tbl1, text="字号:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=5, anchor='e').pack(side='left', padx=(10, 0))
        self.table_size_var = tk.StringVar()
        self._create_combobox(row_tbl1, self.table_size_var, [f"{name}({pt}pt)" for name, pt in FONT_SIZES], width=11,
                              initial_value=self._size_display(self.settings.get('table', {}).get('size', 12))).pack(side='left', padx=3)
        
        tk.Label(row_tbl1, text="行距:", font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY, width=5, anchor='e').pack(side='left', padx=(10, 0))
        self.table_line_spacing_var = tk.StringVar(value=str(self.settings.get('table', {}).get('line_spacing', 22) or ''))
        tk.Entry(row_tbl1, textvariable=self.table_line_spacing_var, font=get_font(11), width=5, relief='solid', bd=1).pack(side='left', padx=3)
        tk.Label(row_tbl1, text="磅", font=get_font(10), bg=Theme.BG, fg=Theme.TEXT_MUTED).pack(side='left')
        
        row_tbl2 = tk.Frame(table_frame, bg=Theme.BG)
        row_tbl2.pack(fill='x', pady=2)
        self.table_header_bold_var = tk.BooleanVar(value=self.settings.get('table', {}).get('header_bold', True))
        self._create_inline_toggle(
            row_tbl2, "表头行加粗", self.table_header_bold_var
        ).pack(side='left', padx=(6, 0))
        
        self.table_smart_align_var = tk.BooleanVar(
            value=self.settings.get('table', {}).get('smart_align', False)
        )
        self._create_inline_toggle(
            row_tbl2, "智能调整单元格对齐", self.table_smart_align_var
        ).pack(side='left', padx=(16, 0))

        # ⓘ 说明按钮
        align_info_btn = tk.Label(
            row_tbl2, text=" ⓘ ",
            font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_MUTED,
            cursor='hand2'
        )
        align_info_btn.pack(side='left')
        align_info_btn.bind('<Enter>', lambda e: align_info_btn.configure(fg=Theme.PRIMARY))
        align_info_btn.bind('<Leave>', lambda e: align_info_btn.configure(fg=Theme.TEXT_MUTED))
        align_info_btn.bind('<Button-1>', lambda e: self._show_table_align_info(e))
        
        # --- 特殊选项 ---
        self._create_section(main, "✨ 特殊选项", pad_x)
        special_frame = tk.Frame(main, bg=Theme.BG)
        special_frame.pack(fill='x', pady=(0, 12), padx=pad_x)

        # 空格处理
        space_row = tk.Frame(special_frame, bg=Theme.BG)
        space_row.pack(anchor='w', pady=(0, 4))
        tk.Label(
            space_row, text="空格处理:", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY
        ).pack(side='left', padx=(6, 8))
        self.space_handling_var = tk.StringVar(
            value=self.settings.get('space_handling', 'remove_all')
        )
        for val, label in [
            ('remove_all',      '删除全部空格'),
                ('keep_en_boundary','规范英文/数字前后空格（补齐为一个空格）'),
            ('keep_all',        '不处理空格'),
        ]:
            tk.Radiobutton(
                space_row, text=label, value=val,
                variable=self.space_handling_var,
                font=get_font(11), bg=Theme.BG, fg=Theme.TEXT,
                activebackground=Theme.BG, selectcolor=Theme.CARD,
            ).pack(side='left', padx=(0, 8))

        self.first_bold_var = tk.BooleanVar(value=self.settings.get('first_line_bold', False))
        self._create_inline_toggle(
            special_frame, "正文段落首句加粗", self.first_bold_var, font_size=12
        ).pack(anchor='w', padx=6, pady=3)

        self.bold_serial_var = tk.BooleanVar(value=self.settings.get('bold_serial', True))
        self._create_inline_toggle(
            special_frame,
            "「一是/一要/第一条」等序列词自动加粗",
            self.bold_serial_var,
            font_size=12,
        ).pack(anchor='w', padx=6, pady=3)

        self.deep_clean_var = tk.BooleanVar(value=self.settings.get('deep_clean', False))
        self._create_inline_toggle(
            special_frame,
            "强力清洗模式（清除原文格式后再排版，适合复制粘贴的内容）",
            self.deep_clean_var,
            font_size=12,
        ).pack(anchor='w', padx=6, pady=3)

        self.page_number_var = tk.BooleanVar(value=self.settings.get('page_number', True))
        self._create_inline_toggle(
            special_frame, "添加页码", self.page_number_var, font_size=12
        ).pack(anchor='w', padx=6, pady=3)

        # 页码格式
        pn_row = tk.Frame(special_frame, bg=Theme.BG)
        pn_row.pack(fill='x', anchor='w', pady=(2, 4))
        tk.Label(
            pn_row, text="页码样式:", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY
        ).pack(side='left', padx=(6, 4))
        current_style = self.settings.get('page_number_style', 'dash')
        self.page_number_style_var = tk.StringVar(
            value=PAGE_NUMBER_STYLE_OPTIONS.get(current_style, PAGE_NUMBER_STYLE_OPTIONS['dash'])
        )
        self._create_combobox(
            pn_row, self.page_number_style_var,
            list(PAGE_NUMBER_STYLE_OPTIONS.values()), width=22,
            initial_value=self.page_number_style_var.get()
        ).pack(side='left')

        tk.Label(
            pn_row, text="位置:", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY
        ).pack(side='left', padx=(14, 4))
        current_position = self.settings.get('page_number_position', 'outside')
        self.page_number_position_var = tk.StringVar(
            value=PAGE_NUMBER_POSITION_OPTIONS.get(
                current_position, PAGE_NUMBER_POSITION_OPTIONS['outside']
            )
        )
        self._create_combobox(
            pn_row, self.page_number_position_var,
            list(PAGE_NUMBER_POSITION_OPTIONS.values()), width=28,
            initial_value=self.page_number_position_var.get()
        ).pack(side='left')

        pn_font_row = tk.Frame(special_frame, bg=Theme.BG)
        pn_font_row.pack(fill='x', anchor='w', pady=(2, 4))
        tk.Label(
            pn_font_row, text="页码字体:", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY
        ).pack(side='left', padx=(6, 4))
        self.page_number_font_var = tk.StringVar(value=self.settings.get('page_number_font', '宋体'))
        page_number_fonts = COMMON_FONTS_CN or ['宋体', '仿宋', '仿宋_GB2312', 'Times New Roman']
        self._create_combobox(
            pn_font_row, self.page_number_font_var, page_number_fonts, width=18,
            initial_value=self.page_number_font_var.get()
        ).pack(side='left')

        tk.Label(
            pn_font_row, text="字号:", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY
        ).pack(side='left', padx=(14, 4))
        self.page_number_size_var = tk.StringVar(
            value=self._size_display(self.settings.get('page_number_size', 14))
        )
        self._create_combobox(
            pn_font_row, self.page_number_size_var,
            [f"{name}({pt}pt)" for name, pt in FONT_SIZES],
            width=11, initial_value=self.page_number_size_var.get()
        ).pack(side='left')

        fd_row = tk.Frame(special_frame, bg=Theme.BG)
        fd_row.pack(anchor='w', pady=(2, 6))
        tk.Label(
            fd_row, text="距版心下边缘:", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY
        ).pack(side='left', padx=(6, 4))
        self.page_number_offset_var = tk.StringVar(
            value=str(self.settings.get('page_number_offset_mm', 7))
        )
        tk.Entry(
            fd_row, textvariable=self.page_number_offset_var,
            font=get_font(11), width=5, relief='solid', bd=1
        ).pack(side='left')
        tk.Label(fd_row, text=" mm", font=get_font(10),
                 bg=Theme.BG, fg=Theme.TEXT_MUTED).pack(side='left')
        tk.Label(
            fd_row,
            text="  公文通常为 7mm；程序会根据下边距自动换算，不是距纸底 7mm",
            font=get_font(9), bg=Theme.BG, fg=Theme.TEXT_MUTED
        ).pack(side='left', padx=(8, 0))

        self.replace_page_number_var = tk.BooleanVar(
            value=self.settings.get('replace_existing_page_number', True)
        )
        self._create_inline_toggle(
            special_frame,
            "重新应用页码设置（替换文档中已有的页码）",
            self.replace_page_number_var,
        ).pack(anchor='w', padx=6, pady=2)
        
        # ============================================================
        #  高级设置（可折叠）
        # ============================================================
        self._create_advanced_section(main, pad_x)
        
        # ===== 底部按钮 =====
        btn_frame = tk.Frame(self, bg=Theme.BG)
        btn_frame.pack(fill='x', padx=20, pady=(10, 10))
        
        tk.Frame(btn_frame, bg=Theme.BORDER, height=1).pack(fill='x', pady=(0, 12))
        
        btn_row = tk.Frame(btn_frame, bg=Theme.BG)
        btn_row.pack(fill='x')
        
        # 恢复默认
        reset_btn = tk.Label(
            btn_row, text="恢复默认公文格式", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY, cursor='hand2'
        )
        reset_btn.pack(side='left')
        reset_btn.bind('<Button-1>', lambda e: self._reset_defaults())
        reset_btn.bind('<Enter>', lambda e: reset_btn.configure(fg=Theme.PRIMARY))
        reset_btn.bind('<Leave>', lambda e: reset_btn.configure(fg=Theme.TEXT_SECONDARY))
        
        # 保存按钮
        save_btn = tk.Frame(btn_row, bg=Theme.PRIMARY, cursor='hand2')
        save_btn.pack(side='right')
        save_label = tk.Label(
            save_btn, text="  保存设置  ", font=get_font(12, 'bold'),
            bg=Theme.PRIMARY, fg='white', pady=8, cursor='hand2'
        )
        save_label.pack()
        for w in [save_btn, save_label]:
            w.bind('<Button-1>', lambda e: self._save())
            w.bind('<Enter>', lambda e: (save_btn.configure(bg=Theme.PRIMARY_HOVER), save_label.configure(bg=Theme.PRIMARY_HOVER)))
            w.bind('<Leave>', lambda e: (save_btn.configure(bg=Theme.PRIMARY), save_label.configure(bg=Theme.PRIMARY)))
        
        cancel_btn = tk.Label(
            btn_row, text="取消", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY, cursor='hand2', padx=15
        )
        cancel_btn.pack(side='right', padx=(0, 15))
        cancel_btn.bind('<Button-1>', lambda e: self._on_close())
        
        size_grip = ttk.Sizegrip(btn_frame)
        size_grip.pack(side='right', padx=(0, 2), pady=(2, 0))

    def _build_preset_bar(self, parent):
        """构建预设管理工具栏（v1.8.0 新增）。"""
        bar = tk.Frame(parent, bg=Theme.BG)
        bar.pack(fill='x', padx=Theme.SPACE_MD, pady=(Theme.SPACE_SM, 0))

        tk.Label(bar, text='当前预设：', bg=Theme.BG,
                 fg=Theme.TEXT, font=get_font(12)).pack(side='left')

        self.preset_var = tk.StringVar()
        self.preset_combo = ttk.Combobox(
            bar, textvariable=self.preset_var,
            state='readonly', width=24, font=get_font(12)
        )
        self.preset_combo.pack(side='left', padx=(Theme.SPACE_XS, Theme.SPACE_MD))
        self.preset_combo.bind('<<ComboboxSelected>>', self._on_preset_selected)

        btn_style = {'bg': Theme.BG, 'fg': Theme.TEXT,
                     'font': get_font(11), 'relief': 'flat',
                     'cursor': 'hand2', 'padx': Theme.SPACE_SM, 'pady': 2,
                     'borderwidth': 1, 'highlightthickness': 0}
        tk.Button(bar, text='+ 新建', command=self._on_new_preset,
                  **btn_style).pack(side='left', padx=2)
        tk.Button(bar, text='✎ 重命名', command=self._on_rename_preset,
                  **btn_style).pack(side='left', padx=2)
        tk.Button(bar, text='🗑 删除', command=self._on_delete_preset,
                  **btn_style).pack(side='left', padx=2)

        # 分隔
        tk.Frame(bar, width=1, bg=Theme.BORDER).pack(
            side='left', fill='y', padx=Theme.SPACE_MD, pady=4
        )

        tk.Button(bar, text='⬆ 导出', command=self._on_export_preset,
                  **btn_style).pack(side='left', padx=2)
        tk.Button(bar, text='⬇ 导入', command=self._on_import_preset,
                  **btn_style).pack(side='left', padx=2)

        return bar

    def _refresh_preset_list(self):
        """刷新下拉菜单显示。"""
        presets = [p for p in self._config.get('presets', []) if not p.get('is_builtin')]
        self._visible_presets = presets
        names = [p.get('name', '未命名') for p in presets]
        self.preset_combo['values'] = names

        # 选中当前 active 的
        active_id = self._config.get('active_preset_id')
        if active_id:
            for i, p in enumerate(presets):
                if p.get('id') == active_id:
                    self.preset_combo.current(i)
                    return
        if names:
            self.preset_combo.current(0)
            self._config['active_preset_id'] = presets[0].get('id')

    def _write_current_settings_to_config(self):
        """把当前 self.settings 写回 _config 中对应的 preset。"""
        active_id = self._config.get('active_preset_id')
        if active_id:
            for i, p in enumerate(self._config.get('presets', [])):
                if p.get('id') == active_id:
                    self.settings['id'] = active_id
                    self.settings['name'] = p.get('name', self.settings.get('name', '我的自定义格式'))
                    self.settings['is_builtin'] = False
                    self._config['presets'][i] = self.settings
                    return

    def _on_preset_selected(self, event=None):
        """切换预设：保存当前设置回旧 preset → 加载新 preset 的值。"""
        # 先把当前 UI 的值保存回 self.settings
        self._save_values()

        # 把 self.settings 写回当前 active preset
        self._write_current_settings_to_config()

        # 切到新选中的 preset
        new_idx = self.preset_combo.current()
        if new_idx < 0:
            return
        new_preset = self._visible_presets[new_idx]
        self._config['active_preset_id'] = new_preset.get('id')

        # 加载新 preset 的值到 UI
        import copy
        self.settings = copy.deepcopy(new_preset)
        self._load_values()

    def _on_new_preset(self):
        """新建一份预设（基于当前 UI 值复制一份）。"""
        from tkinter import simpledialog
        name = simpledialog.askstring(
            '新建预设', '请输入预设名称：', parent=self.dialog,
            initialvalue='新预设'
        )
        if not name or not name.strip():
            return
        name = name.strip()

        # 检查重名
        existing = [p.get('name') for p in self._config.get('presets', [])]
        if name in existing:
            from tkinter import messagebox
            messagebox.showwarning('重名', f'已存在名为「{name}」的预设', parent=self.dialog)
            return

        # 基于当前 UI 值新建
        self._save_values()
        self._write_current_settings_to_config()
        import copy
        new_preset = copy.deepcopy(self.settings)
        new_preset['id'] = str(uuid.uuid4())
        new_preset['name'] = name
        new_preset['is_builtin'] = False

        self._config.setdefault('presets', []).append(new_preset)
        self._config['active_preset_id'] = new_preset['id']
        self.settings = new_preset
        self._refresh_preset_list()

    def _on_rename_preset(self):
        """重命名当前选中的预设。"""
        from tkinter import simpledialog, messagebox
        active_id = self._config.get('active_preset_id')
        if not active_id:
            return
        cur = next((p for p in self._config['presets']
                    if p.get('id') == active_id), None)
        if not cur:
            return

        new_name = simpledialog.askstring(
            '重命名预设', '新名称：', parent=self.dialog,
            initialvalue=cur.get('name', '')
        )
        if not new_name or not new_name.strip():
            return
        new_name = new_name.strip()

        # 检查重名
        existing = [p.get('name') for p in self._config['presets']
                    if p.get('id') != active_id]
        if new_name in existing:
            messagebox.showwarning('重名', f'已存在名为「{new_name}」的预设', parent=self.dialog)
            return

        cur['name'] = new_name
        self.settings['name'] = new_name
        self._refresh_preset_list()

    def _on_delete_preset(self):
        """删除当前选中的预设。最后一份不能删（保护）。"""
        from tkinter import messagebox
        presets = self._config.get('presets', [])
        if len(presets) <= 1:
            messagebox.showinfo('提示', '至少需要保留一份预设，不能删除最后一份', parent=self.dialog)
            return

        active_id = self._config.get('active_preset_id')
        cur = next((p for p in presets if p.get('id') == active_id), None)
        if not cur:
            return

        if not messagebox.askyesno(
            '确认删除',
            f'确定要删除预设「{cur.get("name", "?")}」吗？此操作不可撤销。',
            parent=self.dialog
        ):
            return

        self._config['presets'] = [p for p in presets if p.get('id') != active_id]
        new_active = self._config['presets'][0]
        self._config['active_preset_id'] = new_active.get('id')

        import copy
        self.settings = copy.deepcopy(new_active)
        self._refresh_preset_list()
        self._load_values()

    def _on_export_preset(self):
        """导出当前选中的预设为 .json 文件。"""
        from tkinter import filedialog, messagebox
        import json

        self._save_values()
        self._write_current_settings_to_config()
        active_id = self._config.get('active_preset_id')
        cur = next((p for p in self._config.get('presets', [])
                    if p.get('id') == active_id), None)
        if not cur:
            return

        default_name = f"{cur.get('name', 'preset')}.docfmt.json"
        path = filedialog.asksaveasfilename(
            title='导出预设',
            defaultextension='.json',
            initialfile=default_name,
            filetypes=[('公文格式预设', '*.json'), ('所有文件', '*.*')],
            parent=self.dialog,
        )
        if not path:
            return

        # 导出时去掉 id（导入时会重新生成，避免冲突）
        export_data = {
            'schema_version': CONFIG_SCHEMA_VERSION,
            'export_type': 'docformat-preset',
            'preset': {k: v for k, v in cur.items() if k != 'id'},
        }
        try:
            with open(path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, ensure_ascii=False, indent=2)
            messagebox.showinfo('导出成功', f'预设已导出到\n{path}',
                                parent=self.dialog)
        except Exception as e:
            messagebox.showerror('导出失败', str(e), parent=self.dialog)

    def _on_import_preset(self):
        """导入 .json 预设文件。"""
        from tkinter import filedialog, messagebox
        import json
        import copy

        path = filedialog.askopenfilename(
            title='导入预设',
            filetypes=[('公文格式预设', '*.json'), ('所有文件', '*.*')],
            parent=self.dialog,
        )
        if not path:
            return

        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
        except Exception as e:
            messagebox.showerror('导入失败', f'文件无法解析：{e}',
                                 parent=self.dialog)
            return

        # 兼容三种格式：
        #   a) 标准导出格式：{schema_version, export_type, preset: {...}}
        #   b) 完整 config 格式：{schema_version, presets: [...]}
        #   c) 老 v1.7.x 单 preset 格式：直接是 settings dict
        preset_data = None
        if isinstance(data, dict):
            if 'preset' in data and isinstance(data['preset'], dict):
                preset_data = data['preset']
            elif 'presets' in data and isinstance(data['presets'], list):
                if data['presets']:
                    preset_data = data['presets'][0]
            elif 'page' in data or 'body' in data:
                preset_data = data

        if not preset_data:
            messagebox.showerror('导入失败', '文件格式不识别',
                                 parent=self.dialog)
            return

        new_preset = copy.deepcopy(preset_data)
        new_preset['id'] = str(uuid.uuid4())
        new_preset['is_builtin'] = False

        base_name = new_preset.get('name', '导入的预设')
        existing = {p.get('name') for p in self._config.get('presets', [])}
        final_name = base_name
        counter = 2
        while final_name in existing:
            final_name = f"{base_name} ({counter})"
            counter += 1
        new_preset['name'] = final_name

        self._save_values()
        self._write_current_settings_to_config()
        self._config.setdefault('presets', []).append(new_preset)
        self._config['active_preset_id'] = new_preset['id']
        self.settings = new_preset
        self._refresh_preset_list()
        self._load_values()

        messagebox.showinfo(
            '导入成功',
            f'预设已导入为「{final_name}」',
            parent=self.dialog
        )
    
    def _create_advanced_section(self, parent, pad_x):
        """创建可折叠的高级设置区域"""
        self._adv_expanded = False
        
        # 折叠按钮
        self._adv_toggle_frame = tk.Frame(parent, bg=Theme.BG)
        self._adv_toggle_frame.pack(fill='x', padx=pad_x, pady=(8, 0))
        
        # 分隔线
        tk.Frame(self._adv_toggle_frame, bg=Theme.BORDER, height=1).pack(fill='x', pady=(0, 8))
        
        self._adv_toggle_label = tk.Label(
            self._adv_toggle_frame,
            text="▸ 高级设置 — 按元素类型独立配置字体/行距",
            font=get_font(12, 'bold'), bg=Theme.BG, fg=Theme.TEXT_SECONDARY,
            cursor='hand2', anchor='w'
        )
        self._adv_toggle_label.pack(anchor='w')
        self._adv_toggle_label.bind('<Button-1>', lambda e: self._toggle_advanced())
        self._adv_toggle_label.bind('<Enter>', lambda e: self._adv_toggle_label.configure(fg=Theme.PRIMARY))
        self._adv_toggle_label.bind('<Leave>', lambda e: self._adv_toggle_label.configure(fg=Theme.TEXT_SECONDARY))
        
        # 高级内容区域（初始隐藏）
        self._adv_content = tk.Frame(parent, bg=Theme.BG)
        # 不 pack — 初始隐藏
        
        tk.Label(
            self._adv_content,
            text="ⓘ 此处可逐个元素类型覆盖上方快速设置的值。留空行距表示跟随正文行距。",
            font=get_font(9), bg=Theme.BG, fg=Theme.TEXT_MUTED
        ).pack(anchor='w', padx=pad_x, pady=(5, 8))
        
        # 元素类型列表
        elements = [
            ('recipient', '🏢 主送机关', '仿宋_GB2312', 16),
            ('heading1',  '1️⃣  一级标题 (一、)', '黑体', 16),
            ('heading2',  '2️⃣  二级标题 ((一))', '楷体_GB2312', 16),
            ('heading3',  '3️⃣  三级标题 (1.)', '仿宋_GB2312', 16),
            ('heading4',  '4️⃣  四级标题 ((1))', '仿宋_GB2312', 16),
            ('attachment', '📎 附件', '仿宋_GB2312', 16),
            ('closing',   '🧾 结束语', '仿宋_GB2312', 16),
            ('signature', '✒️  落款单位', '仿宋_GB2312', 16),
            ('date',      '📅 落款日期', '仿宋_GB2312', 16),
        ]
        
        for key, label, default_font, default_size in elements:
            self._create_adv_element_row(self._adv_content, pad_x, key, label, default_font, default_size)
    
    def _create_adv_element_row(self, parent, pad_x, key, label, default_font, default_size):
        """创建高级设置中的一个元素行：中文字体 + 英数字体 + 字号 + 行距 + 加粗"""
        row = tk.Frame(parent, bg=Theme.BG)
        row.pack(fill='x', padx=pad_x, pady=2)
        
        tk.Label(row, text=label, font=get_font(10), bg=Theme.BG, fg=Theme.TEXT, width=14, anchor='w').pack(side='left')
        
        # 中文字体
        font_var = tk.StringVar()
        self._create_combobox(row, font_var, COMMON_FONTS_CN, width=12,
                              initial_value=self.settings.get(key, {}).get('font_cn', default_font)).pack(side='left', padx=3)

        # 英数字体
        tk.Label(row, text="英数:", font=get_font(10), bg=Theme.BG, fg=Theme.TEXT_SECONDARY).pack(side='left', padx=(6, 0))
        font_en_var = tk.StringVar()
        self._create_combobox(row, font_en_var, COMMON_FONTS_EN, width=12,
                              initial_value=self.settings.get(key, {}).get('font_en', 'Times New Roman')).pack(side='left', padx=3)
        
        # 字号
        tk.Label(row, text="字号:", font=get_font(10), bg=Theme.BG, fg=Theme.TEXT_SECONDARY).pack(side='left', padx=(6, 0))
        size_var = tk.StringVar()
        self._create_combobox(row, size_var, [f"{name}({pt}pt)" for name, pt in FONT_SIZES], width=9,
                              initial_value=self._size_display(self.settings.get(key, {}).get('size', default_size))).pack(side='left', padx=3)
        
        # 行距
        tk.Label(row, text="行距:", font=get_font(10), bg=Theme.BG, fg=Theme.TEXT_SECONDARY).pack(side='left', padx=(6, 0))
        ls_val = self.settings.get(key, {}).get('line_spacing', '')
        ls_var = tk.StringVar(value=str(ls_val) if ls_val else '')
        tk.Entry(row, textvariable=ls_var, font=get_font(10), width=4, relief='solid', bd=1).pack(side='left', padx=3)
        
        # 加粗
        default_bold = DEFAULT_CUSTOM_SETTINGS.get(key, {}).get('bold', False)
        bold_var = tk.BooleanVar(value=self.settings.get(key, {}).get('bold', default_bold))
        self._create_inline_toggle(
            row, "粗", bold_var, font_size=10
        ).pack(side='left', padx=(6, 0))
        
        # 存储变量引用
        self._adv_vars[key] = {'font': font_var, 'font_en': font_en_var, 'size': size_var, 'line_spacing': ls_var, 'bold': bold_var}
        
        # 记录初始值，用于判断用户是否修改过高级设置
        if not hasattr(self, '_adv_initial_values'):
            self._adv_initial_values = {}
        self._adv_initial_values[key] = {
            'font': font_var.get(),
            'font_en': font_en_var.get(),
            'size': size_var.get(),
            'line_spacing': ls_var.get(),
            'bold': bold_var.get(),
        }
    
    def _show_table_align_info(self, event):
        """弹出表格智能对齐规则说明"""
        popup = tk.Toplevel(self)
        popup.overrideredirect(True)
        popup.configure(bg=Theme.CARD)

        border = tk.Frame(popup, bg=Theme.BORDER, padx=1, pady=1)
        border.pack()
        inner = tk.Frame(border, bg=Theme.CARD)
        inner.pack()

        msg = (
            "启用后，工具将按以下规则自动调整单元格对齐：\n\n"
            "  · 表头行（第一行）→ 居中\n"
            "  · 含\"合计\"/\"总计\"的单元格 → 居中\n"
            "  · 序号列 → 居中\n"
            "  · 纯数字/百分比/金额 → 靠右\n"
            "  · 4字以内的短文本 → 居中\n"
            "  · 其余较长文本 → 靠左\n\n"
            "默认关闭，保留文档原始对齐格式。"
        )
        tk.Label(
            inner, text=msg,
            font=get_font(10), bg=Theme.CARD, fg=Theme.TEXT,
            justify='left', padx=14, pady=10,
        ).pack()

        popup.update_idletasks()
        x = event.x_root + 10
        y = event.y_root + 10
        sw = popup.winfo_screenwidth()
        pw = popup.winfo_reqwidth()
        if x + pw > sw - 10:
            x = sw - pw - 10
        popup.geometry(f'+{x}+{y}')

        popup.bind('<Button-1>', lambda e: popup.destroy())
        popup.focus_set()
        popup.bind('<FocusOut>', lambda e: popup.destroy())
    
    def _toggle_advanced(self):
        """切换高级设置的折叠/展开"""
        if self._adv_expanded:
            self._adv_content.pack_forget()
            self._adv_toggle_label.configure(text="▸ 高级设置 — 按元素类型独立配置字体/行距")
            self._adv_expanded = False
        else:
            self._adv_content.pack(fill='x', after=self._adv_toggle_frame, pady=(0, 10))
            self._adv_toggle_label.configure(text="▾ 高级设置 — 按元素类型独立配置字体/行距")
            self._adv_expanded = True
        
        # 更新滚动区域
        self.content_frame.update_idletasks()
        self.canvas.configure(scrollregion=self.canvas.bbox('all'))
    
    # ==================== 滚动/辅助方法 ====================
    
    def _on_frame_configure(self, event):
        self.canvas.configure(scrollregion=self.canvas.bbox('all'))
    
    def _on_canvas_configure(self, event):
        self.canvas.itemconfig(self.canvas_window, width=event.width)
    
    def _bind_mousewheel(self):
        # 只绑定到自定义设置窗口自身，避免 unbind_all 把主窗口滚轮也解绑。
        self.bind('<MouseWheel>', self._on_mousewheel)
        self.bind('<Button-4>', self._on_mousewheel)
        self.bind('<Button-5>', self._on_mousewheel)
        self.bind('<Shift-MouseWheel>', self._on_shift_mousewheel)
        self.canvas.bind('<MouseWheel>', self._on_mousewheel)
        self.canvas.bind('<Button-4>', self._on_mousewheel)
        self.canvas.bind('<Button-5>', self._on_mousewheel)
        self.canvas.bind('<Shift-MouseWheel>', self._on_shift_mousewheel)
    
    def _unbind_mousewheel(self):
        self.unbind('<MouseWheel>')
        self.unbind('<Button-4>')
        self.unbind('<Button-5>')
        self.unbind('<Shift-MouseWheel>')
        self.canvas.unbind('<MouseWheel>')
        self.canvas.unbind('<Button-4>')
        self.canvas.unbind('<Button-5>')
        self.canvas.unbind('<Shift-MouseWheel>')
    
    def _on_mousewheel(self, event):
        if getattr(event, 'num', None) == 4:
            self.canvas.yview_scroll(-1, 'units')
        elif getattr(event, 'num', None) == 5:
            self.canvas.yview_scroll(1, 'units')
        else:
            delta = getattr(event, 'delta', 0)
            units = -int(delta / 120) if delta else 0
            if units == 0 and delta:
                units = -1 if delta > 0 else 1
            if units:
                self.canvas.yview_scroll(units, 'units')
        return "break"

    def _on_shift_mousewheel(self, event):
        delta = getattr(event, 'delta', 0)
        if delta:
            self.canvas.xview_scroll(int(-1 * (delta / 120)), 'units')
        return "break"
    
    def _create_section(self, parent, title, padx=0):
        tk.Label(
            parent, text=title, font=get_font(12, 'bold'),
            bg=Theme.BG, fg=Theme.TEXT
        ).pack(anchor='w', pady=(10, 4), padx=padx)
    
    def _create_combobox(self, parent, variable, values, width=15, initial_value=None):
        """创建可滚动下拉框（ttk.Combobox）"""
        if initial_value is not None:
            if initial_value in values:
                reordered = [initial_value] + [v for v in values if v != initial_value]
            else:
                reordered = [initial_value] + list(values)
        else:
            reordered = list(values)
        
        frame = tk.Frame(
            parent, bg=Theme.INPUT_BG,
            highlightbackground=Theme.BORDER, highlightthickness=1
        )
        
        combo = ttk.Combobox(
            frame,
            textvariable=variable,
            values=reordered,
            width=width,
            state='readonly',
            font=get_font(10),
        )
        combo.pack(fill='x', padx=1, pady=1)
        
        if initial_value is not None:
            variable.set(initial_value)
        
        combo.bind('<FocusIn>',  lambda e: frame.configure(highlightbackground=Theme.PRIMARY))
        combo.bind('<FocusOut>', lambda e: frame.configure(highlightbackground=Theme.BORDER))
        
        return frame
    
    def _size_display(self, pt_value):
        """pt值 → 显示字符串"""
        try:
            pt_value = float(pt_value)
        except (TypeError, ValueError):
            pt_value = 16.0
        for name, pt in FONT_SIZES:
            if abs(float(pt) - pt_value) < 0.01:
                return f"{name}({pt}pt)"
        return f"自定义({pt_value}pt)"
    
    def _get_size_from_var(self, var):
        """从字号下拉框获取pt值"""
        text = var.get()
        for name, pt in FONT_SIZES:
            if f"{name}({pt}pt)" == text:
                return pt
        import re as _re
        match = _re.search(r'\((\d+(?:\.\d+)?)\s*pt\)', text)
        if match:
            return float(match.group(1))
        return 16
    
    def _get_line_spacing(self, var, fallback=28):
        """从行距输入框获取值，空值返回 fallback"""
        val = var.get().strip()
        if not val:
            return fallback
        try:
            return int(float(val))
        except ValueError:
            return fallback
    
    # ==================== 加载/保存 ====================
    
    def _load_values(self):
        """加载设置到 UI"""
        s = self.settings
        try:
            # 页边距
            for key in ['top', 'bottom', 'left', 'right']:
                self.margin_vars[key].set(str(s.get('page', {}).get(key, 2.5)))
            
            # 标题
            self.title_font_var.set(s.get('title', {}).get('font_cn', '方正小标宋简体'))
            self._set_size_var(self.title_size_var, s.get('title', {}).get('size', 22))
            self.title_line_spacing_var.set(str(s.get('title', {}).get('line_spacing', 28) or ''))
            self.title_bold_var.set(s.get('title', {}).get('bold', False))
            
            # 一/二级标题
            self.h1_font_var.set(s.get('heading1', {}).get('font_cn', '黑体'))
            self._set_size_var(self.h1_size_var, s.get('heading1', {}).get('size', 16))
            self.h1_bold_var.set(s.get('heading1', {}).get('bold', False))
            self.h2_font_var.set(s.get('heading2', {}).get('font_cn', '楷体_GB2312'))
            self._set_size_var(self.h2_size_var, s.get('heading2', {}).get('size', 16))
            self.h2_bold_var.set(s.get('heading2', {}).get('bold', False))
            
            # 正文
            self.body_font_var.set(s.get('body', {}).get('font_cn', '仿宋_GB2312'))
            self._set_size_var(self.body_size_var, s.get('body', {}).get('size', 16))
            self.line_spacing_var.set(str(s.get('body', {}).get('line_spacing', 28) or ''))
            self.body_bold_var.set(s.get('body', {}).get('bold', False))
            self.global_font_en_var.set(
                s.get('body', {}).get('font_en', 'Times New Roman')
            )
            self.space_before_var.set(str(s.get('body', {}).get('space_before', 0)))
            self.space_after_var.set(str(s.get('body', {}).get('space_after', 0)))
            
            body_size = s.get('body', {}).get('size', 16) or 16
            indent = s.get('body', {}).get('indent', 32)
            indent_chars = int(indent / body_size) if body_size else 2
            self.indent_var.set(f'{indent_chars}字符')
            
            # 表格
            tbl = s.get('table', {})
            self.table_font_var.set(tbl.get('font_cn', '仿宋_GB2312'))
            self._set_size_var(self.table_size_var, tbl.get('size', 12))
            self.table_line_spacing_var.set(str(tbl.get('line_spacing', 22) or ''))
            self.table_header_bold_var.set(tbl.get('header_bold', True))
            self.table_smart_align_var.set(
                self.settings.get('table', {}).get('smart_align', False)
            )
            
            # 特殊选项
            self.first_bold_var.set(s.get('first_line_bold', False))
            self.bold_serial_var.set(s.get('bold_serial', True))
            self.deep_clean_var.set(s.get('deep_clean', False))
            self.space_handling_var.set(s.get('space_handling', 'remove_all'))
            self.page_number_var.set(s.get('page_number', True))
            self.page_number_style_var.set(
                PAGE_NUMBER_STYLE_OPTIONS.get(
                    s.get('page_number_style', 'dash'),
                    PAGE_NUMBER_STYLE_OPTIONS['dash']
                )
            )
            self.page_number_position_var.set(
                PAGE_NUMBER_POSITION_OPTIONS.get(
                    s.get('page_number_position', 'outside'),
                    PAGE_NUMBER_POSITION_OPTIONS['outside']
                )
            )
            self.page_number_font_var.set(s.get('page_number_font', '宋体'))
            self._set_size_var(self.page_number_size_var, s.get('page_number_size', 14))
            self.page_number_offset_var.set(str(s.get('page_number_offset_mm', 7)))
            self.replace_page_number_var.set(s.get('replace_existing_page_number', True))
            
            # 高级设置
            for key, vars_dict in self._adv_vars.items():
                elem = s.get(key, {})
                vars_dict['font'].set(elem.get('font_cn', '仿宋_GB2312'))
                vars_dict['font_en'].set(elem.get('font_en', 'Times New Roman'))
                self._set_size_var(vars_dict['size'], elem.get('size', 16))
                ls = elem.get('line_spacing', '')
                vars_dict['line_spacing'].set(str(ls) if ls else '')
                default_bold = DEFAULT_CUSTOM_SETTINGS.get(key, {}).get('bold', False)
                vars_dict['bold'].set(elem.get('bold', default_bold))
        except Exception as e:
            print(f"[警告] 加载设置到界面失败: {e}")
    
    def _set_size_var(self, var, pt_value):
        try:
            pt_value = float(pt_value)
        except (TypeError, ValueError):
            pt_value = 16.0
        for name, pt in FONT_SIZES:
            if abs(float(pt) - pt_value) < 0.01:
                var.set(f"{name}({pt}pt)")
                return
        var.set(f"自定义({pt_value}pt)")
    
    def _reset_defaults(self):
        import copy
        self.settings = copy.deepcopy(DEFAULT_CUSTOM_SETTINGS)
        self._load_values()

    def _save_values(self):
        """保存当前 UI 值到 self.settings，不写文件、不关闭弹窗。"""
        current_name = self.settings.get('name', '我的自定义格式')

        # 收集快速设置值
        page = {key: float(self.margin_vars[key].get()) for key in ['top', 'bottom', 'left', 'right']}

        title_size = self._get_size_from_var(self.title_size_var)
        h1_size = self._get_size_from_var(self.h1_size_var)
        h2_size = self._get_size_from_var(self.h2_size_var)
        body_size = self._get_size_from_var(self.body_size_var)
        body_ls = self._get_line_spacing(self.line_spacing_var, 28)
        title_ls = self._get_line_spacing(self.title_line_spacing_var, 28)
        try:
            space_before = int(float(self.space_before_var.get()))
            space_after = int(float(self.space_after_var.get()))
        except ValueError:
            space_before = space_after = 0

        # 首行缩进
        indent_text = self.indent_var.get()
        indent_chars = int(float(indent_text.replace('字符', '')))
        indent_pt = indent_chars * body_size

        body_font = self.body_font_var.get()
        global_font_en = self.global_font_en_var.get()
        body_bold = self.body_bold_var.get()
        page_number_style = next(
            (
                key for key, label in PAGE_NUMBER_STYLE_OPTIONS.items()
                if label == self.page_number_style_var.get()
            ),
            'dash',
        )
        page_number_position = next(
            (
                key for key, label in PAGE_NUMBER_POSITION_OPTIONS.items()
                if label == self.page_number_position_var.get()
            ),
            'outside',
        )
        try:
            page_number_offset_mm = float(self.page_number_offset_var.get())
        except ValueError:
            page_number_offset_mm = 7
        page_number_offset_mm = max(0, min(30, page_number_offset_mm))

        # 构建基础设置 — 正文字体联动到多个元素
        self.settings = {
            'name': current_name,
            'page': page,
            'title': {
                'font_cn': self.title_font_var.get(), 'font_en': global_font_en,
                'size': title_size, 'bold': self.title_bold_var.get(), 'align': 'center', 'indent': 0,
                'line_spacing': title_ls, 'space_before': space_before, 'space_after': space_after
            },
            'recipient': {
                'font_cn': body_font, 'font_en': global_font_en,
                'size': body_size, 'bold': body_bold, 'align': 'left', 'indent': 0,
                'line_spacing': body_ls, 'space_before': space_before, 'space_after': space_after
            },
            'heading1': {
                'font_cn': self.h1_font_var.get(), 'font_en': global_font_en,
                'size': h1_size, 'bold': self.h1_bold_var.get(), 'align': 'left', 'indent': indent_pt,
                'line_spacing': body_ls, 'space_before': space_before, 'space_after': space_after
            },
            'heading2': {
                'font_cn': self.h2_font_var.get(), 'font_en': global_font_en,
                'size': h2_size, 'bold': self.h2_bold_var.get(), 'align': 'left', 'indent': indent_pt,
                'line_spacing': body_ls, 'space_before': space_before, 'space_after': space_after
            },
            'heading3': {
                'font_cn': body_font, 'font_en': global_font_en,
                'size': body_size, 'bold': body_bold, 'align': 'left', 'indent': indent_pt,
                'line_spacing': body_ls, 'space_before': space_before, 'space_after': space_after
            },
            'heading4': {
                'font_cn': body_font, 'font_en': global_font_en,
                'size': body_size, 'bold': body_bold, 'align': 'left', 'indent': indent_pt,
                'line_spacing': body_ls, 'space_before': space_before, 'space_after': space_after
            },
            'body': {
                'font_cn': body_font, 'font_en': global_font_en,
                'size': body_size, 'bold': body_bold, 'align': 'justify', 'indent': indent_pt,
                'line_spacing': body_ls, 'space_before': space_before, 'space_after': space_after
            },
            'signature': {
                'font_cn': body_font, 'font_en': global_font_en,
                'size': body_size, 'bold': body_bold, 'align': 'right', 'indent': 0,
                'line_spacing': body_ls, 'space_before': space_before, 'space_after': space_after
            },
            'date': {
                'font_cn': body_font, 'font_en': global_font_en,
                'size': body_size, 'bold': body_bold, 'align': 'right', 'indent': 0,
                'line_spacing': body_ls, 'space_before': space_before, 'space_after': space_after
            },
            'attachment': {
                'font_cn': body_font, 'font_en': global_font_en,
                'size': body_size, 'bold': body_bold, 'align': 'justify', 'indent': 0,
                'line_spacing': body_ls, 'space_before': space_before, 'space_after': space_after
            },
            'closing': {
                'font_cn': body_font, 'font_en': global_font_en,
                'size': body_size, 'bold': body_bold, 'align': 'left', 'indent': indent_pt,
                'line_spacing': body_ls, 'space_before': space_before, 'space_after': space_after
            },
            'table': {
                'font_cn': self.table_font_var.get(), 'font_en': global_font_en,
                'size': self._get_size_from_var(self.table_size_var), 'bold': False,
                'line_spacing': self._get_line_spacing(self.table_line_spacing_var, 22),
                'first_line_indent': 0,
                'header_bold': self.table_header_bold_var.get(),
                'smart_align': self.table_smart_align_var.get()
            },
            'space_handling': self.space_handling_var.get(),
            'first_line_bold': self.first_bold_var.get(),
            'bold_serial': self.bold_serial_var.get(),
            'deep_clean': self.deep_clean_var.get(),
            'page_number': self.page_number_var.get(),
            'page_number_font': self.page_number_font_var.get(),
            'page_number_size': self._get_size_from_var(self.page_number_size_var),
            'page_number_style': page_number_style,
            'page_number_position': page_number_position,
            'page_number_offset_mm': page_number_offset_mm,
            'replace_existing_page_number': self.replace_page_number_var.get(),
        }

        # 应用高级设置覆盖（仅在用户真正修改过时）
        initial = getattr(self, '_adv_initial_values', {})
        for key, vars_dict in self._adv_vars.items():
            if key in self.settings and isinstance(self.settings[key], dict):
                key_initial = initial.get(key, {})

                adv_font = vars_dict['font'].get()
                adv_font_en = vars_dict['font_en'].get()
                adv_size = self._get_size_from_var(vars_dict['size'])
                adv_ls_str = vars_dict['line_spacing'].get().strip()

                # 只在值与初始值不同时才覆盖（说明用户主动修改了）
                if adv_font and adv_font != key_initial.get('font', ''):
                    self.settings[key]['font_cn'] = adv_font
                if adv_font_en and adv_font_en != key_initial.get('font_en', ''):
                    self.settings[key]['font_en'] = adv_font_en
                if adv_size and vars_dict['size'].get() != key_initial.get('size', ''):
                    self.settings[key]['size'] = adv_size
                if adv_ls_str and adv_ls_str != key_initial.get('line_spacing', '').strip():
                    try:
                        self.settings[key]['line_spacing'] = int(float(adv_ls_str))
                    except ValueError:
                        pass
                self.settings[key]['bold'] = vars_dict['bold'].get()
    
    def _save(self):
        """保存设置 - 快速设置为主，高级设置覆盖"""
        try:
            # 保存当前 UI 值回 self.settings
            self._save_values()

            # 写回 _config 中对应 preset
            active_id = self._config.get('active_preset_id')
            if active_id:
                for i, p in enumerate(self._config['presets']):
                    if p.get('id') == active_id:
                        self.settings['id'] = active_id
                        self.settings['is_builtin'] = False
                        self._config['presets'][i] = self.settings
                        break
            save_custom_settings(self._config)
            
            if self.on_save:
                self.on_save(self.settings)
            
            # 先释放模态锁，再关闭窗口（macOS 上 grab + messagebox + destroy 可能卡死）
            try:
                self.grab_release()
            except Exception:
                pass
            self.destroy()
            
        except ValueError as e:
            messagebox.showerror("输入错误", f"请检查输入的数值是否正确：\n{e}", parent=self)
        except Exception as e:
            messagebox.showerror("保存失败", f"保存设置时出错：\n{e}\n\n配置文件路径：{CONFIG_FILE}", parent=self)
    
    def _on_close(self):
        result = messagebox.askyesnocancel("保存设置", "是否保存当前设置？", parent=self)
        if result is None:
            return
        if result:
            self._save()
        else:
            try:
                self.grab_release()
            except Exception:
                pass
            self.destroy()


class PasteTextDialog(tk.Toplevel):
    """从粘贴文本生成 docx 的对话框（v1.8.0 新增）。

    用户粘贴文本 → 工具创建临时 docx → 走 format_document 流程
    → 输出格式化后的 docx。
    """

    def __init__(self, parent, on_generate=None):
        super().__init__(parent)
        self.on_generate = on_generate

        self.title("从文本生成 docx")
        self.configure(bg=Theme.BG)
        self.resizable(True, True)

        # 模态
        self.transient(parent)
        self.grab_set()
        self.protocol("WM_DELETE_WINDOW", self._on_close)

        _fit_dialog_to_screen(
            self, parent,
            desired_w=1100, desired_h=900,
            min_w=640, min_h=560
        )

        self._build_ui()

    def _build_ui(self):
        """构建 UI。"""
        # ===== 顶部标题 =====
        header = tk.Frame(self, bg=Theme.BG)
        header.pack(fill='x', padx=20, pady=(15, 8))

        tk.Label(
            header, text="📋 从文本生成 docx", font=get_font(16, 'bold'),
            bg=Theme.BG, fg=Theme.TEXT
        ).pack(side='left')

        cancel_top = tk.Label(
            header, text="取消", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY, cursor='hand2', padx=10
        )
        cancel_top.pack(side='right')
        cancel_top.bind('<Button-1>', lambda e: self._on_close())

        # ===== 主体区域（整体可滚动）=====
        body_container = tk.Frame(self, bg=Theme.BG)
        body_container.pack(fill='both', expand=True, padx=(20, 10))

        self.body_canvas = tk.Canvas(body_container, bg=Theme.BG, highlightthickness=0)
        body_scrollbar = tk.Scrollbar(
            body_container, orient='vertical', command=self.body_canvas.yview
        )
        self.body_canvas.configure(yscrollcommand=body_scrollbar.set)
        body_scrollbar.pack(side='right', fill='y')
        self.body_canvas.pack(side='left', fill='both', expand=True)

        main = tk.Frame(self.body_canvas, bg=Theme.BG)
        self.body_canvas_window = self.body_canvas.create_window(
            (0, 0), window=main, anchor='nw'
        )
        main.bind(
            '<Configure>',
            lambda e: self.body_canvas.configure(scrollregion=self.body_canvas.bbox('all'))
        )
        self.body_canvas.bind(
            '<Configure>',
            lambda e: self.body_canvas.itemconfig(self.body_canvas_window, width=e.width)
        )

        # --- 文档标题 ---
        title_label = tk.Label(
            main, text="文档标题：", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY
        )
        title_label.pack(anchor='w', pady=(0, 4))

        self.title_var = tk.StringVar(value="新建文档")
        title_entry = tk.Entry(
            main, textvariable=self.title_var,
            font=get_font(12), relief='solid', bd=1,
            bg=Theme.INPUT_BG, fg=Theme.TEXT,
        )
        title_entry.pack(fill='x', pady=(0, 4), ipady=4)

        tk.Label(
            main, text="ⓘ 此标题将作为文档第一段，居中、用标题字体",
            font=get_font(9), bg=Theme.BG, fg=Theme.TEXT_MUTED
        ).pack(anchor='w', pady=(0, 12))

        # --- 正文粘贴区 ---
        tk.Label(
            main, text="正文内容（可粘贴纯文本或 Markdown）：",
            font=get_font(11), bg=Theme.BG, fg=Theme.TEXT_SECONDARY
        ).pack(anchor='w', pady=(0, 4))

        # Text 控件 + 滚动条
        text_container = tk.Frame(
            main, bg=Theme.INPUT_BG,
            highlightbackground=Theme.BORDER, highlightthickness=1
        )
        text_container.pack(fill='both', expand=True, pady=(0, 4))

        self.text_widget = tk.Text(
            text_container,
            font=get_font(11), bg=Theme.INPUT_BG, fg=Theme.TEXT,
            relief='flat', wrap='word', padx=10, pady=8,
            insertbackground=Theme.TEXT,
        )
        scrollbar = tk.Scrollbar(text_container, orient='vertical',
                                 command=self.text_widget.yview)
        self.text_widget.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side='right', fill='y')
        self.text_widget.pack(side='left', fill='both', expand=True)

        # 占位提示文字
        placeholder = (
            "粘贴文本到这里...\n\n"
            "· 每个空行分隔的段落自动作为新段\n"
            "· 一、（一）1. 等编号自动识别为各级标题\n"
            "· 落款单位、日期自动识别"
        )
        self.text_widget.insert('1.0', placeholder)
        self.text_widget.configure(fg=Theme.TEXT_MUTED)
        self._placeholder_active = True

        self.text_widget.bind('<FocusIn>', self._clear_placeholder)
        self.text_widget.bind('<KeyRelease>', self._update_stats)
        self.text_widget.bind('<<Paste>>', self._on_paste)

        # 字数统计
        self.stats_label = tk.Label(
            main, text="ⓘ 已粘贴 0 字 · 0 段",
            font=get_font(9), bg=Theme.BG, fg=Theme.TEXT_MUTED
        )
        self.stats_label.pack(anchor='w', pady=(0, 12))

        # --- 保存位置 ---
        tk.Label(
            main, text="保存位置：", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY
        ).pack(anchor='w', pady=(0, 4))

        save_row = tk.Frame(main, bg=Theme.BG)
        save_row.pack(fill='x', pady=(0, 8))

        self.save_path_var = tk.StringVar()
        save_entry = tk.Entry(
            save_row, textvariable=self.save_path_var,
            font=get_font(11), relief='solid', bd=1,
            bg=Theme.INPUT_BG, fg=Theme.TEXT,
        )
        save_entry.pack(side='left', fill='x', expand=True, ipady=3)

        browse_btn = tk.Label(
            save_row, text=" 📁 浏览 ", font=get_font(11),
            bg=Theme.BG, fg=Theme.PRIMARY, cursor='hand2',
            padx=8, pady=4,
        )
        browse_btn.pack(side='left', padx=(8, 0))
        browse_btn.bind('<Button-1>', lambda e: self._browse_save_dir())
        browse_btn.bind('<Enter>', lambda e: browse_btn.configure(fg=Theme.PRIMARY_HOVER))
        browse_btn.bind('<Leave>', lambda e: browse_btn.configure(fg=Theme.PRIMARY))

        # 默认保存到桌面
        desktop = Path.home() / 'Desktop'
        if desktop.exists():
            self.save_path_var.set(str(desktop))

        # ===== 底部按钮 =====
        btn_frame = tk.Frame(main, bg=Theme.BG)
        btn_frame.pack(fill='x', pady=(8, 16))

        tk.Frame(btn_frame, bg=Theme.BORDER, height=1).pack(fill='x', pady=(0, 12))

        btn_row = tk.Frame(btn_frame, bg=Theme.BG)
        btn_row.pack(fill='x')

        # 生成按钮
        gen_btn = tk.Frame(btn_row, bg=Theme.PRIMARY, cursor='hand2')
        gen_btn.pack(side='right')
        gen_label = tk.Label(
            gen_btn, text="  生成 docx 并处理  ", font=get_font(12, 'bold'),
            bg=Theme.PRIMARY, fg='white', pady=8, cursor='hand2'
        )
        gen_label.pack()
        for w in [gen_btn, gen_label]:
            w.bind('<Button-1>', lambda e: self._generate())
            w.bind('<Enter>', lambda e: (gen_btn.configure(bg=Theme.PRIMARY_HOVER),
                                          gen_label.configure(bg=Theme.PRIMARY_HOVER)))
            w.bind('<Leave>', lambda e: (gen_btn.configure(bg=Theme.PRIMARY),
                                          gen_label.configure(bg=Theme.PRIMARY)))

        cancel_btn = tk.Label(
            btn_row, text="取消", font=get_font(11),
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY, cursor='hand2', padx=15
        )
        cancel_btn.pack(side='right', padx=(0, 15))
        cancel_btn.bind('<Button-1>', lambda e: self._on_close())

    def _clear_placeholder(self, event=None):
        """首次聚焦时清掉占位文字。"""
        if self._placeholder_active:
            self.text_widget.delete('1.0', 'end')
            self.text_widget.configure(fg=Theme.TEXT)
            self._placeholder_active = False

    def _on_paste(self, event=None):
        """粘贴事件：清占位 + 延迟更新统计。"""
        if self._placeholder_active:
            self.text_widget.delete('1.0', 'end')
            self.text_widget.configure(fg=Theme.TEXT)
            self._placeholder_active = False
        # 粘贴是异步的，延迟更新
        self.after(50, self._update_stats)

    def _update_stats(self, event=None):
        """更新字数/段数统计。"""
        if self._placeholder_active:
            return
        text = self.text_widget.get('1.0', 'end').strip()
        char_count = len(text)
        # 段数：以空行分隔
        if not text:
            para_count = 0
        else:
            blocks = [b for b in text.split('\n\n') if b.strip()]
            para_count = len(blocks) if blocks else len([
                line for line in text.split('\n') if line.strip()
            ])
        format_hint = ""
        if char_count > 20:
            is_md = _detect_markdown(text)
            format_hint = " · 📝 Markdown 格式" if is_md else " · 纯文本"
        self.stats_label.configure(
            text=f"ⓘ 已粘贴 {char_count} 字 · {para_count} 段{format_hint}"
        )

    def _browse_save_dir(self):
        """选择保存目录。"""
        d = filedialog.askdirectory(
            title='选择保存目录',
            initialdir=self.save_path_var.get() or str(Path.home()),
            parent=self,
        )
        if d:
            self.save_path_var.set(d)

    def _generate(self):
        """生成 docx 并触发主流程的格式化处理。"""
        # 验证输入
        title = self.title_var.get().strip()
        if not title:
            messagebox.showwarning('提示', '请填写文档标题', parent=self)
            return

        if self._placeholder_active:
            messagebox.showwarning('提示', '请粘贴文本内容', parent=self)
            return

        text = self.text_widget.get('1.0', 'end').strip()
        if not text:
            messagebox.showwarning('提示', '请粘贴文本内容', parent=self)
            return

        save_dir = self.save_path_var.get().strip()
        if not save_dir:
            messagebox.showwarning('提示', '请选择保存位置', parent=self)
            return

        save_dir_path = Path(save_dir)
        if not save_dir_path.is_dir():
            messagebox.showerror('错误', f'保存位置不是目录：{save_dir}', parent=self)
            return

        # 文件名清洗：去掉特殊字符
        safe_title = re.sub(r'[\\/:*?"<>|]', '_', title)
        output_path = save_dir_path / f"{safe_title}.docx"

        # v1.8.0: 检测格式
        is_md = _detect_markdown(text)

        # 调用回调，把数据交给主流程
        if self.on_generate:
            try:
                self.grab_release()
            except Exception:
                pass
            self.destroy()
            self.on_generate(title, text, str(output_path), is_md)

    def _on_close(self):
        try:
            self.grab_release()
        except Exception:
            pass
        self.destroy()


def _create_docx_from_text(title, body_text, output_path):
    """把"标题 + 正文文本"生成一个最简 docx 文件。

    每个空行分隔的块作为一个段落。后续的 format_document 会自动识别
    各级标题、落款等。

    v1.8.0 新增。
    """
    from docx import Document

    doc = Document()
    # v1.8.1: 给标题段加 CENTER 对齐 + 大字号，让 detect_para_type
    # 能稳定识别为 title（避免短标题被回退识别为 body）
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title.strip())
    title_run.font.size = Pt(22)

    # 正文：先按双换行拆段，每段内的单换行也拆成独立段。
    body_text = body_text.strip()
    if body_text:
        blocks = body_text.split('\n\n')
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            # 块内按单换行再拆，每行作为一段
            for line in block.split('\n'):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)

    doc.save(output_path)
    return output_path


# v1.8.0: Markdown 检测与解析
import re as _md_re


def _detect_markdown(text):
    """启发式判断文本是否为 Markdown 格式。"""
    if not text or not text.strip():
        return False

    score = 0
    lines = text.split('\n')

    # 标题
    h_pattern = _md_re.compile(r'^\s*#{1,6}\s*\S')
    h_count = sum(1 for line in lines if h_pattern.match(line))
    if h_count >= 1:
        score += 3

    # 加粗
    bold_count = len(_md_re.findall(r'\*\*[^*\n]+\*\*', text))
    if bold_count >= 2:
        score += 2

    # 无序列表
    ul_pattern = _md_re.compile(r'^\s*[-*+]\s+\S')
    ul_count = sum(1 for line in lines if ul_pattern.match(line))
    if ul_count >= 2:
        score += 2

    # 代码块
    if '```' in text:
        score += 2

    # 引用
    quote_count = sum(1 for line in lines if line.startswith('> '))
    if quote_count >= 1:
        score += 1

    return score >= 3


def _parse_markdown_inline(text):
    """解析单行内的 Markdown 行内格式，返回 [(text, is_bold), ...]。"""
    parts = []
    pattern = _md_re.compile(r'\*\*([^*\n]+)\*\*|__([^_\n]+)__')
    last_end = 0

    for m in pattern.finditer(text):
        # 加粗前的普通文本
        if m.start() > last_end:
            parts.append((text[last_end:m.start()], False))
        # 加粗内容
        bold_text = m.group(1) or m.group(2)
        if bold_text:
            parts.append((bold_text, True))
        last_end = m.end()

    # 末尾普通文本
    if last_end < len(text):
        parts.append((text[last_end:], False))

    return parts if parts else [(text, False)]


def _create_docx_from_markdown(title, md_text, output_path):
    """把 Markdown 文本解析为 docx。"""
    from docx import Document

    doc = Document()

    # 解析 Markdown 文本
    lines = md_text.split('\n')

    # 标题计数器，把 ##/###/#### 转成 一、/(一)/1. 等中文格式
    h2_counter = 0
    h3_counters = {}  # h2_index -> h3_count
    h4_counters = {}  # (h2_index, h3_index) -> h4_count
    current_h2 = 0
    current_h3 = 0

    # 中文数字
    cn_nums = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
               '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十']

    def cn(n):
        return cn_nums[n - 1] if 1 <= n <= len(cn_nums) else str(n)

    # 主标题
    title_added = False

    # 状态：是否在代码块内
    in_code_block = False
    code_buffer = []

    h_pattern = _md_re.compile(r'^\s*(#{1,6})\s*(.*)$')
    ul_pattern = _md_re.compile(r'^\s*[-*+]\s+(.*)$')
    ol_pattern = _md_re.compile(r'^\s*\d+\.\s+(.*)$')
    quote_pattern = _md_re.compile(r'^>\s+(.*)$')

    def add_para_with_inline(text, is_title=False, alignment=None):
        """添加段落，解析行内加粗。"""
        para = doc.add_paragraph()
        if alignment is not None:
            para.alignment = alignment
        parts = _parse_markdown_inline(text)
        for content, is_bold in parts:
            if not content:
                continue
            run = para.add_run(content)
            if is_bold:
                run.font.bold = True
        return para

    for line in lines:
        # 代码块边界
        if line.strip().startswith('```'):
            if in_code_block:
                # 结束代码块，把内容输出为段落
                for code_line in code_buffer:
                    if code_line.strip():
                        doc.add_paragraph(code_line)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        # 空行
        if not line.strip():
            continue

        # 标题
        h_match = h_pattern.match(line)
        if h_match:
            level = len(h_match.group(1))
            content = h_match.group(2).strip()

            if level == 1:
                # 主标题：覆盖用户输入的 title
                if not title_added:
                    # v1.8.1: 给 # 主标题加 CENTER + 大字号，确保
                    # detect_para_type 识别为 title
                    from docx.shared import Pt
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    title_para = add_para_with_inline(content)
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if title_para.runs:
                        for run in title_para.runs:
                            run.font.size = Pt(22)
                    title_added = True
                else:
                    # 文档已有标题，把额外的 # 转成一级标题
                    h2_counter += 1
                    add_para_with_inline(f"{cn(h2_counter)}、{content}")
                    current_h2 = h2_counter
                continue
            elif level == 2:
                h2_counter += 1
                current_h2 = h2_counter
                add_para_with_inline(f"{cn(h2_counter)}、{content}")
                continue
            elif level == 3:
                h3_counters[current_h2] = h3_counters.get(current_h2, 0) + 1
                current_h3 = h3_counters[current_h2]
                add_para_with_inline(f"（{cn(current_h3)}）{content}")
                continue
            elif level >= 4:
                key = (current_h2, current_h3)
                h4_counters[key] = h4_counters.get(key, 0) + 1
                add_para_with_inline(f"{h4_counters[key]}. {content}")
                continue

        # 列表项
        ul_match = ul_pattern.match(line)
        ol_match = ol_pattern.match(line)
        if ul_match:
            add_para_with_inline(ul_match.group(1).strip())
            continue
        if ol_match:
            add_para_with_inline(ol_match.group(1).strip())
            continue

        # 引用
        q_match = quote_pattern.match(line)
        if q_match:
            add_para_with_inline(q_match.group(1).strip())
            continue

        # 普通段落
        add_para_with_inline(line.strip())

    # 如果代码块没有闭合，也把缓冲内容输出
    if in_code_block:
        for code_line in code_buffer:
            if code_line.strip():
                doc.add_paragraph(code_line)

    # 如果用户给的 title 没被 # 标题覆盖，补一个标题段在开头
    # v1.8.1: 同时加 CENTER + 大字号让 detect_para_type 稳定识别
    if not title_added and title.strip():
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title.strip())
        title_run.font.size = Pt(22)
        body = doc.element.body
        body.remove(title_para._p)
        body.insert(0, title_para._p)

    doc.save(output_path)
    return output_path


# ===== 大尺寸线条图标 =====
class Icons:
    """用 Canvas 绘制的线条图标 - 48px 大尺寸"""
    
    @staticmethod
    def draw_magic(canvas, x, y, size=48, color='#2E2E2E'):
        """智能处理 - 魔法棒"""
        s = size
        lw = 2.5  # 线宽
        # 魔法棒主体
        canvas.create_line(x+s*0.15, y+s*0.85, x+s*0.65, y+s*0.35, fill=color, width=lw, capstyle='round')
        # 星星点缀
        stars = [(0.7, 0.2), (0.85, 0.35), (0.75, 0.5), (0.55, 0.15)]
        for px, py in stars:
            r = 3
            canvas.create_oval(x+s*px-r, y+s*py-r, x+s*px+r, y+s*py+r, fill=color, outline='')
        # 光芒线
        canvas.create_line(x+s*0.7, y+s*0.08, x+s*0.7, y+s*0.22, fill=color, width=1.5)
        canvas.create_line(x+s*0.9, y+s*0.28, x+s*0.78, y+s*0.35, fill=color, width=1.5)
    
    @staticmethod
    def draw_search(canvas, x, y, size=48, color='#2E2E2E'):
        """诊断 - 放大镜"""
        s = size
        lw = 2.5
        # 镜框
        canvas.create_oval(x+s*0.12, y+s*0.12, x+s*0.58, y+s*0.58, outline=color, width=lw)
        # 镜柄
        canvas.create_line(x+s*0.52, y+s*0.52, x+s*0.85, y+s*0.85, fill=color, width=lw, capstyle='round')
        # 高光
        canvas.create_arc(x+s*0.18, y+s*0.18, x+s*0.4, y+s*0.4, start=120, extent=60, style='arc', outline=color, width=1.5)
    
    @staticmethod
    def draw_edit(canvas, x, y, size=48, color='#2E2E2E'):
        """标点修复 - 铅笔"""
        s = size
        lw = 2.5
        # 笔身
        canvas.create_line(x+s*0.2, y+s*0.8, x+s*0.7, y+s*0.3, fill=color, width=lw, capstyle='round')
        # 笔尖
        canvas.create_polygon(
            x+s*0.15, y+s*0.85,
            x+s*0.2, y+s*0.8,
            x+s*0.25, y+s*0.85,
            fill=color, outline=''
        )
        # 笔头
        canvas.create_line(x+s*0.7, y+s*0.3, x+s*0.8, y+s*0.2, fill=color, width=lw, capstyle='round')
        canvas.create_line(x+s*0.75, y+s*0.35, x+s*0.85, y+s*0.25, fill=color, width=lw, capstyle='round')

    @staticmethod
    def draw_file(canvas, x, y, size=48, color='#2E2E2E'):
        """文件图标"""
        s = size
        lw = 2
        # 文件主体
        points = [
            x+s*0.2, y+s*0.1,   # 左上
            x+s*0.2, y+s*0.9,   # 左下
            x+s*0.8, y+s*0.9,   # 右下
            x+s*0.8, y+s*0.3,   # 右上（折角下）
            x+s*0.6, y+s*0.1,   # 折角
        ]
        canvas.create_polygon(points, fill='', outline=color, width=lw)
        # 折角线
        canvas.create_line(x+s*0.6, y+s*0.1, x+s*0.6, y+s*0.3, fill=color, width=lw)
        canvas.create_line(x+s*0.6, y+s*0.3, x+s*0.8, y+s*0.3, fill=color, width=lw)
    
    @staticmethod
    def draw_check(canvas, x, y, size=32, color='#7CB87C'):
        """勾选"""
        s = size
        canvas.create_line(x+s*0.15, y+s*0.5, x+s*0.4, y+s*0.75, fill=color, width=3, capstyle='round')
        canvas.create_line(x+s*0.4, y+s*0.75, x+s*0.85, y+s*0.25, fill=color, width=3, capstyle='round')


class FileInputField(tk.Frame):
    """文件输入框 - 带明显容器"""
    
    def __init__(self, parent, label_text, placeholder, variable, command, drop_command=None, **kwargs):
        super().__init__(parent, bg=Theme.BG, **kwargs)
        
        self.variable = variable
        self.command = command
        self.drop_command = drop_command
        self.placeholder = placeholder
        
        # 标签
        tk.Label(
            self,
            text=label_text,
            font=get_font(11),
            bg=Theme.BG,
            fg=Theme.TEXT_SECONDARY,
            width=4,
            anchor='w'
        ).pack(side='left')
        
        # 输入框容器
        self.container = tk.Frame(
            self,
            bg=Theme.INPUT_BG,
            highlightbackground=Theme.BORDER,
            highlightcolor=Theme.PRIMARY,
            highlightthickness=1
        )
        self.container.pack(side='left', fill='x', expand=True, padx=(Theme.SPACE_SM, 0))
        
        inner = tk.Frame(self.container, bg=Theme.INPUT_BG)
        inner.pack(fill='both', expand=True, padx=Theme.SPACE_MD, pady=Theme.SPACE_SM + 2)
        
        # 文件名显示
        self.filename_label = tk.Label(
            inner,
            text="未选择",
            font=get_font(11),
            bg=Theme.INPUT_BG,
            fg=Theme.TEXT_MUTED,
            anchor='w'
        )
        self.filename_label.pack(side='left', fill='x', expand=True)
        
        # 分隔线
        tk.Frame(inner, bg=Theme.BORDER, width=1).pack(side='left', fill='y', padx=Theme.SPACE_MD)
        
        # 操作按钮
        self.action_btn = tk.Label(
            inner,
            text=placeholder,
            font=get_font(10),
            bg=Theme.INPUT_BG,
            fg=Theme.PRIMARY,
            cursor='hand2'
        )
        self.action_btn.pack(side='right')
        self._drop_widgets = [self.container, inner, self.filename_label, self.action_btn]
        
        # 绑定点击
        for widget in [self.container, inner, self.filename_label, self.action_btn]:
            widget.bind('<Button-1>', self._on_click)
            widget.configure(cursor='hand2')
        
        # 悬停效果
        self.container.bind('<Enter>', lambda e: self.container.configure(highlightbackground='#D0CCC6'))
        self.container.bind('<Leave>', lambda e: self.container.configure(highlightbackground=Theme.BORDER))
        
        # 监听变量
        self.variable.trace_add('write', self._update_display)
        
        # 拖拽支持
        if _DND_AVAILABLE:
            self._enable_drag_drop()
    
    def _on_click(self, event=None):
        if self.command:
            self.command()
    
    def _enable_drag_drop(self):
        """启用拖拽文件支持"""
        global _DND_AVAILABLE, _DND_DISABLED_REASON
        root = self.winfo_toplevel()
        if not hasattr(root, 'drop_target_register') or not hasattr(root, 'dnd_bind'):
            _DND_AVAILABLE = False
            _DND_DISABLED_REASON = "当前 Tk 根窗口不支持 tkinterdnd2"
            return
        try:
            if not getattr(root, '_dnd_registered', False):
                root.drop_target_register(DND_FILES)
                root.dnd_bind('<<Drop>>', self._dispatch_drop)
                root._dnd_registered = True
                root._dnd_fields = []
            if not hasattr(root, '_dnd_fields'):
                root._dnd_fields = []
            if self not in root._dnd_fields:
                root._dnd_fields.append(self)

            for widget in self._drop_widgets:
                widget.drop_target_register(DND_FILES)
                widget.dnd_bind('<<DropEnter>>', self._on_drag_enter)
                widget.dnd_bind('<<DropLeave>>', self._on_drag_leave)
                widget.dnd_bind('<<Drop>>', self._on_drop)
        except Exception as exc:
            _DND_AVAILABLE = False
            _DND_DISABLED_REASON = f"拖拽初始化失败：{exc}"
    
    def _on_drag_enter(self, event):
        """拖拽进入时高亮边框"""
        self.container.configure(highlightbackground=Theme.PRIMARY)
        return COPY
    
    def _on_drag_leave(self, event):
        """拖拽离开时恢复边框"""
        self.container.configure(highlightbackground=Theme.BORDER)
        return COPY
    
    def _dispatch_drop(self, event):
        """
        根据鼠标坐标判断文件落在哪个 FileInputField，然后分发给对应字段处理。
        绑定在根窗口上，避免 canvas create_window 截获事件。
        """
        root = self.winfo_toplevel()
        fields = getattr(root, '_dnd_fields', [])
        for field in fields:
            try:
                x = field.container.winfo_rootx()
                y = field.container.winfo_rooty()
                w = field.container.winfo_width()
                h = field.container.winfo_height()
                if x <= event.x_root <= x + w and y <= event.y_root <= y + h:
                    # 高亮然后处理
                    field.container.configure(highlightbackground=Theme.PRIMARY)
                    field.after(300, lambda f=field: f.container.configure(
                        highlightbackground=Theme.BORDER))
                    field._on_drop(event)
                    return
            except Exception:
                continue
    
    def _on_drop(self, event):
        """处理拖拽放下的文件"""
        self.container.configure(highlightbackground=Theme.BORDER)
        # tkinterdnd2 返回的路径可能带花括号（多文件或路径含空格时）
        raw = event.data.strip()
        # 解析路径列表
        if raw.startswith('{'):
            # 多文件或含空格路径：{path1} {path2}
            import re as _re
            paths = _re.findall(r'\{([^}]+)\}', raw)
            if not paths:
                paths = [raw.strip('{}')]
        else:
            paths = raw.split()
        
        if not paths:
            return
        
        # 只取第一个文件路径（输入框单文件模式）
        path = paths[0].strip()
        if path:
            if self.drop_command:
                self.drop_command(path)
            else:
                self.variable.set(path)
        return COPY
    
    def _update_display(self, *args):
        path = self.variable.get()
        if path:
            # 显示文件名，路径过长则截断
            filename = Path(path).name
            if len(filename) > 40:
                filename = filename[:37] + "..."
            self.filename_label.configure(text=filename, fg=Theme.TEXT)
        else:
            self.filename_label.configure(text="未选择", fg=Theme.TEXT_MUTED)


class SelectableCard(tk.Frame):
    """可选择的卡片 - 大图标版"""
    
    def __init__(self, parent, title, description, value, variable,
                 icon_draw_func=None, is_featured=False, command=None, **kwargs):
        
        bg_color = Theme.CARD_ALT if is_featured else Theme.CARD
        super().__init__(parent, bg=bg_color, **kwargs)
        
        self.value = value
        self.variable = variable
        self.command = command
        self.is_featured = is_featured
        self.bg_color = bg_color
        self.selected = False
        
        # 边框
        self.configure(
            highlightbackground=Theme.BORDER,
            highlightcolor=Theme.BORDER_SELECTED,
            highlightthickness=1
        )
        
        # 内容 - 水平布局：左图标 + 右文字
        content = tk.Frame(self, bg=bg_color)
        content.pack(fill='both', expand=True, padx=Theme.SPACE_LG, pady=Theme.SPACE_LG)
        
        # 左侧：图标
        if icon_draw_func:
            icon_size = 56 if is_featured else 48
            self.icon_canvas = tk.Canvas(
                content,
                width=icon_size + 8,
                height=icon_size + 8,
                bg=bg_color,
                highlightthickness=0
            )
            self.icon_canvas.pack(side='left', padx=(0, Theme.SPACE_MD))
            icon_draw_func(self.icon_canvas, 4, 4, icon_size, Theme.TEXT)
            self._bind_click(self.icon_canvas)
        
        # 右侧：文字区域
        text_frame = tk.Frame(content, bg=bg_color)
        text_frame.pack(side='left', fill='both', expand=True)
        
        # 标题行（标题 + 推荐标签）
        title_row = tk.Frame(text_frame, bg=bg_color)
        title_row.pack(fill='x', anchor='w')
        
        title_size = 16 if is_featured else 14
        self.title_label = tk.Label(
            title_row,
            text=title,
            font=get_font(title_size, 'bold'),
            bg=bg_color,
            fg=Theme.TEXT,
            anchor='w'
        )
        self.title_label.pack(side='left')
        
        # 推荐标签
        if is_featured:
            tag = tk.Label(
                title_row,
                text=" 推荐 ",
                font=get_font(10, 'bold'),
                bg=Theme.PRIMARY,
                fg='white',
                padx=10,
                pady=3
            )
            tag.pack(side='left', padx=(Theme.SPACE_SM, 0))
            self._bind_click(tag)
        
        self._bind_click(title_row)
        
        # 描述
        desc_size = 12 if is_featured else 11
        self.desc_label = tk.Label(
            text_frame,
            text=description,
            font=get_font(desc_size),
            bg=bg_color,
            fg=Theme.TEXT_SECONDARY,
            anchor='w',
            justify='left'
        )
        self.desc_label.pack(fill='x', anchor='w', pady=(Theme.SPACE_SM, 0))
        
        # 绑定事件
        self._bind_click(self)
        self._bind_click(content)
        self._bind_click(text_frame)
        self._bind_click(self.title_label)
        self._bind_click(self.desc_label)
        
        # 监听变量
        self.variable.trace_add('write', self._on_variable_change)
        self._update_style()
    
    def _bind_click(self, widget):
        widget.bind('<Button-1>', self._on_click)
        widget.bind('<Enter>', self._on_enter)
        widget.bind('<Leave>', self._on_leave)
        widget.configure(cursor='hand2')
    
    def _on_click(self, event=None):
        self.variable.set(self.value)
        if self.command:
            self.command()
    
    def _on_enter(self, event=None):
        if not self.selected:
            self.configure(highlightbackground='#D0CCC6')
    
    def _on_leave(self, event=None):
        self._update_style()
    
    def _on_variable_change(self, *args):
        self._update_style()
    
    def _update_style(self):
        self.selected = (self.variable.get() == self.value)
        if self.selected:
            self.configure(highlightbackground=Theme.BORDER_SELECTED, highlightthickness=2)
        else:
            self.configure(highlightbackground=Theme.BORDER, highlightthickness=1)


class PresetCard(tk.Frame):
    """格式预设卡片"""
    
    def __init__(self, parent, text, value, variable, command=None, **kwargs):
        super().__init__(parent, bg=Theme.CARD, **kwargs)
        
        self.value = value
        self.variable = variable
        self.selected = False
        self.command = command  # 自定义点击命令
        
        self.configure(
            highlightbackground=Theme.BORDER,
            highlightcolor=Theme.BORDER_SELECTED,
            highlightthickness=1
        )
        
        self.label = tk.Label(
            self,
            text=text,
            font=get_font(12),
            bg=Theme.CARD,
            fg=Theme.TEXT,
            padx=Theme.SPACE_LG,
            pady=Theme.SPACE_MD
        )
        self.label.pack()
        
        # 绑定
        for widget in [self, self.label]:
            widget.bind('<Button-1>', self._on_click)
            widget.bind('<Enter>', self._on_enter)
            widget.bind('<Leave>', self._on_leave)
            widget.configure(cursor='hand2')
        
        self.variable.trace_add('write', self._update_style)
        self._update_style()
    
    def _on_click(self, event=None):
        self.variable.set(self.value)
        # 如果有自定义命令，执行它
        if self.command:
            self.command()
    
    def _on_enter(self, event=None):
        if not self.selected:
            self.configure(highlightbackground='#D0CCC6')
    
    def _on_leave(self, event=None):
        self._update_style()
    
    def _update_style(self, *args):
        self.selected = (self.variable.get() == self.value)
        if self.selected:
            self.configure(bg=Theme.PRIMARY_LIGHT, highlightbackground=Theme.PRIMARY, highlightthickness=2)
            self.label.configure(bg=Theme.PRIMARY_LIGHT, fg=Theme.TEXT, font=get_font(12, 'bold'))
        else:
            self.configure(bg=Theme.CARD, highlightbackground=Theme.BORDER, highlightthickness=1)
            self.label.configure(bg=Theme.CARD, fg=Theme.TEXT, font=get_font(12))
    
    def set_enabled(self, enabled):
        if enabled:
            self.label.configure(fg=Theme.TEXT, cursor='hand2')
            self.configure(cursor='hand2')
        else:
            self.label.configure(fg=Theme.TEXT_MUTED, cursor='arrow')
            self.configure(cursor='arrow', highlightbackground=Theme.BORDER_LIGHT)


class CollapsibleLog(tk.Frame):
    """可折叠的日志区域"""
    
    def __init__(self, parent, **kwargs):
        super().__init__(parent, bg=Theme.BG, **kwargs)
        
        self.expanded = False
        
        # 折叠条
        self.toggle_bar = tk.Frame(self, bg='#E8E4DE', height=36)
        self.toggle_bar.pack(fill='x')
        self.toggle_bar.pack_propagate(False)
        
        self.toggle_btn = tk.Label(
            self.toggle_bar,
            text="＋  展开运行日志",
            font=get_font(11),
            bg='#E8E4DE',
            fg=Theme.TEXT_SECONDARY,
            cursor='hand2'
        )
        self.toggle_btn.pack(side='left', padx=Theme.SPACE_MD, pady=Theme.SPACE_SM)
        self.toggle_btn.bind('<Button-1>', self._toggle)
        self.toggle_bar.bind('<Button-1>', self._toggle)
        self.toggle_bar.configure(cursor='hand2')
        
        # 日志面板
        self.log_panel = tk.Frame(self, bg=Theme.LOG_BG)
        
        # 日志文本
        self.log_text = tk.Text(
            self.log_panel,
            font=('Consolas', 11),
            bg=Theme.LOG_BG,
            fg=Theme.LOG_TEXT,
            relief='flat',
            padx=Theme.SPACE_LG,
            pady=Theme.SPACE_MD,
            wrap='word',
            height=10,
            highlightthickness=0,
            insertbackground=Theme.LOG_TEXT
        )
        self.log_text.pack(side='left', fill='both', expand=True)
        
        # 配置颜色标签
        self.log_text.tag_configure('info', foreground=Theme.LOG_TEXT)
        self.log_text.tag_configure('success', foreground=Theme.LOG_SUCCESS)
        self.log_text.tag_configure('warning', foreground=Theme.LOG_WARNING)
        self.log_text.tag_configure('error', foreground=Theme.LOG_ERROR)
    
    def _toggle(self, event=None):
        self.expanded = not self.expanded
        if self.expanded:
            self.log_panel.pack(fill='both', expand=True)
            self.toggle_btn.configure(text="－  收起运行日志")
        else:
            self.log_panel.pack_forget()
            self.toggle_btn.configure(text="＋  展开运行日志")
    
    def log(self, message, tag='info'):
        """线程安全的日志输出"""
        import threading
        if threading.current_thread() is threading.main_thread():
            self.log_text.insert(tk.END, message + "\n", tag)
            self.log_text.see(tk.END)
        else:
            # 非主线程，调度到主线程执行
            self.after(0, lambda: self.log(message, tag))
    
    def clear(self):
        self.log_text.delete(1.0, tk.END)


class ProcessingPanel(tk.Frame):
    """运行中的动态状态面板。"""

    TIPS = [
        "正在整理文档结构",
        "正在处理段落格式",
        "正在应用公文规范",
        "正在保存输出文件",
    ]

    def __init__(self, parent, **kwargs):
        super().__init__(parent, bg=Theme.BG, **kwargs)
        self.running = False
        self.frame_index = 0
        self.percent = 0
        self.stage_text = "准备中..."

        self.shadow = tk.Frame(self, bg=Theme.SHADOW)
        self.card = tk.Frame(
            self,
            bg=Theme.SURFACE,
            highlightbackground=Theme.BORDER,
            highlightthickness=1
        )
        self.shadow.place(relx=0, rely=0, relwidth=1, relheight=1, x=3, y=3)
        self.card.place(relx=0, rely=0, relwidth=1, relheight=1)
        self.configure(height=ui_px(138))
        self.pack_propagate(False)

        inner = tk.Frame(self.card, bg=Theme.SURFACE)
        inner.pack(fill='both', expand=True, padx=Theme.SPACE_LG, pady=ui_px(18))

        self.canvas = tk.Canvas(
            inner, width=ui_px(76), height=ui_px(76),
            bg=Theme.SURFACE, highlightthickness=0
        )
        self.canvas.pack(side='left', padx=(0, Theme.SPACE_MD))

        text_area = tk.Frame(inner, bg=Theme.SURFACE)
        text_area.pack(side='left', fill='both', expand=True)
        text_area.bind('<Configure>', self._on_text_area_configure)

        self.title_label = tk.Label(
            text_area,
            text="正在处理文档",
            font=get_font(14, 'bold'),
            bg=Theme.SURFACE,
            fg=Theme.TEXT,
            anchor='w'
        )
        self.title_label.pack(fill='x', anchor='w')

        self.stage_label = tk.Label(
            text_area,
            text=self.stage_text,
            font=get_font(11),
            bg=Theme.SURFACE,
            fg=Theme.TEXT_SECONDARY,
            anchor='w',
            justify='left',
            wraplength=520
        )
        self.stage_label.pack(fill='x', anchor='w', pady=(7, 0))

        self.tip_label = tk.Label(
            text_area,
            text="",
            font=get_font(10),
            bg=Theme.SURFACE,
            fg=Theme.TEXT_MUTED,
            anchor='w',
            justify='left',
            wraplength=520
        )
        self.tip_label.pack(fill='x', anchor='w', pady=(9, 0))

    def _on_text_area_configure(self, event):
        wrap = max(240, event.width - 8)
        self.stage_label.configure(wraplength=wrap)
        self.tip_label.configure(wraplength=wrap)

    def start(self, title=None):
        self.running = True
        self.frame_index = 0
        self.percent = 0
        self.stage_text = "准备中..."
        self.title_label.configure(text=title or "正在处理文档")
        self.stage_label.configure(text=self.stage_text)
        self.tip_label.configure(text="")
        self.pack(fill='x', pady=(0, Theme.SPACE_MD))
        self._animate()

    def update_status(self, percent, stage_text):
        self.percent = max(0, min(int(percent), 100))
        self.stage_text = stage_text or "处理中..."
        self.stage_label.configure(text=self.stage_text)

    def stop(self):
        self.running = False
        self.pack_forget()

    def _animate(self):
        if not self.running:
            return

        self.frame_index += 1
        self.canvas.delete('all')
        self._draw_animation()
        if _UI_DENSITY > 1.01:
            self.canvas.scale('all', 0, 0, _UI_DENSITY, _UI_DENSITY)

        dot_count = (self.frame_index // 6) % 4
        tip = self.TIPS[(self.frame_index // 28) % len(self.TIPS)]
        self.tip_label.configure(text=f"{tip}{'.' * dot_count}")
        self.after(90, self._animate)

    def _draw_animation(self):
        self.canvas.create_rectangle(
            19, 10, 57, 64,
            fill=Theme.PRIMARY_LIGHT,
            outline=Theme.PRIMARY,
            width=2
        )
        self.canvas.create_line(45, 10, 45, 24, fill=Theme.PRIMARY, width=2)
        self.canvas.create_line(45, 24, 57, 24, fill=Theme.PRIMARY, width=2)

        fill_h = max(4, int(38 * self.percent / 100))
        self.canvas.create_rectangle(
            25, 57 - fill_h, 51, 57,
            fill=Theme.PRIMARY_SOFT,
            outline=''
        )
        sweep_x = 23 + (self.frame_index * 2 % 30)
        self.canvas.create_line(
            sweep_x, 19, sweep_x, 56,
            fill=Theme.PRIMARY,
            width=2,
            capstyle='round'
        )


class InlineToggle(tk.Frame):
    """与主题一致的轻量开关行，用于替代系统默认复选框。"""

    def __init__(self, parent, text, variable, command=None,
                 box_size=28, font=None, bg=None, **kwargs):
        self._bg = bg or Theme.BG
        super().__init__(parent, bg=self._bg, **kwargs)
        self.variable = variable
        self.command = command
        self.enabled = True

        self._box_size = ui_px(box_size)
        self.box = tk.Canvas(
            self, width=self._box_size, height=self._box_size,
            bg=self._bg, highlightthickness=0
        )
        self.box.pack(side='left', padx=(0, Theme.SPACE_SM))

        self.label = tk.Label(
            self,
            text=text,
            font=font or get_font(11, 'bold'),
            bg=self._bg,
            fg=Theme.TEXT_SECONDARY,
            anchor='w'
        )
        self.label.pack(side='left', fill='x', expand=True)

        for widget in (self, self.box, self.label):
            widget.bind('<Button-1>', self._toggle)
            widget.bind('<Enter>', self._on_enter)
            widget.bind('<Leave>', self._on_leave)
            widget.configure(cursor='hand2')

        self.variable.trace_add('write', lambda *_: self._draw())
        self._draw()

    def configure(self, cnf=None, **kwargs):
        state = kwargs.pop('state', None)
        fg = kwargs.pop('fg', None)
        result = super().configure(cnf or {}, **kwargs)
        if state is not None:
            self.set_enabled(state != 'disabled')
        if fg is not None:
            self.label.configure(fg=fg)
        return result

    config = configure

    def set_enabled(self, enabled):
        self.enabled = bool(enabled)
        cursor = 'hand2' if self.enabled else 'arrow'
        for widget in (self, self.box, self.label):
            widget.configure(cursor=cursor)
        self._draw()

    def _toggle(self, event=None):
        if not self.enabled:
            return
        self.variable.set(not self.variable.get())
        if self.command:
            self.command()

    def _on_enter(self, event=None):
        if self.enabled:
            self.label.configure(fg=Theme.ACCENT if self.variable.get() else Theme.TEXT)

    def _on_leave(self, event=None):
        self._draw()

    def _draw(self):
        self.box.delete('all')
        checked = bool(self.variable.get())
        s = self._box_size
        if not self.enabled:
            fill = Theme.BORDER_LIGHT
            outline = Theme.BORDER
            fg = Theme.TEXT_MUTED
        elif checked:
            fill = Theme.ACCENT
            outline = Theme.ACCENT
            fg = Theme.TEXT
        else:
            fill = Theme.SURFACE
            outline = Theme.BORDER
            fg = Theme.TEXT_SECONDARY

        self.box.create_rectangle(
            s * 0.11, s * 0.11, s * 0.89, s * 0.89,
            fill=fill, outline=outline, width=max(2, ui_px(1))
        )
        if checked:
            self.box.create_line(
                s * 0.29, s * 0.50, s * 0.43, s * 0.64, s * 0.71, s * 0.32,
                fill='white', width=max(2.5, ui_px(1.3)),
                capstyle='round', joinstyle='round'
            )
        self.label.configure(fg=fg)


class ProcessingTaskList(tk.Frame):
    """运行时任务清单，按进度把已完成项置灰并划线。"""

    ICON_SIZE = 34

    TASKS = {
        'smart': [
            (6, "读取输入文档"),
            (22, "修复标点与空格"),
            (86, "应用格式规范"),
            (100, "保存输出文件"),
        ],
        'analyze': [
            (20, "读取输入文档"),
            (80, "诊断格式问题"),
            (100, "生成诊断结果"),
        ],
        'punctuation': [
            (20, "读取输入文档"),
            (82, "修复标点与空格"),
            (100, "保存输出文件"),
        ],
        'format': [
            (20, "读取输入文档"),
            (86, "应用格式规范"),
            (100, "保存输出文件"),
        ],
        'normal': [
            (20, "读取输入文档"),
            (80, "处理文档内容"),
            (100, "保存输出文件"),
        ],
    }

    def __init__(self, parent, **kwargs):
        super().__init__(parent, bg=Theme.BG, **kwargs)
        self.rows = []
        self.thresholds = []

    def _get_tasks(self, mode, has_input_conversion=False, has_output_conversion=False):
        tasks = list(self.TASKS.get(mode, self.TASKS['normal']))

        if has_input_conversion:
            tasks.insert(0, (4, "转换为 DOCX 工作稿"))

        if has_output_conversion:
            final_task = tasks[-1] if tasks else (100, "保存输出文件")
            tasks = tasks[:-1] + [(96, "转换输出格式"), final_task]

        return tasks

    def start(self, mode, has_input_conversion=False, has_output_conversion=False):
        tasks = self._get_tasks(
            mode,
            has_input_conversion=has_input_conversion,
            has_output_conversion=has_output_conversion
        )

        for widget in self.winfo_children():
            widget.destroy()

        self.rows = []
        self.thresholds = [threshold for threshold, _text in tasks]

        tk.Label(
            self,
            text="任务清单",
            font=get_font(11, 'bold'),
            bg=Theme.BG,
            fg=Theme.TEXT_SECONDARY,
            anchor='w'
        ).pack(fill='x', pady=(0, Theme.SPACE_SM))

        for _threshold, text in tasks:
            row = tk.Frame(self, bg=Theme.BG)
            row.pack(fill='x', pady=4)

            icon = tk.Canvas(
                row,
                width=ui_px(self.ICON_SIZE),
                height=ui_px(self.ICON_SIZE),
                bg=Theme.BG,
                highlightthickness=0
            )
            icon.pack(side='left', padx=(0, Theme.SPACE_MD))

            label = tk.Label(
                row,
                text=text,
                font=get_font(10),
                bg=Theme.BG,
                fg=Theme.TEXT_SECONDARY,
                anchor='w'
            )
            label.pack(side='left', fill='x', expand=True)
            self.rows.append((icon, label))

        self.pack(fill='x', pady=(0, Theme.SPACE_MD))
        self.update_status(0, "准备中...")

    def update_status(self, percent, stage_text=""):
        if not self.rows:
            return

        active_index = 0
        for idx, threshold in enumerate(self.thresholds):
            if percent >= threshold:
                active_index = idx + 1
        active_index = min(active_index, len(self.rows) - 1)
        if percent >= 100:
            active_index = len(self.rows)

        for idx, (icon, label) in enumerate(self.rows):
            icon.delete('all')
            size = ui_px(self.ICON_SIZE)
            line_width = max(3, ui_px(2))

            if idx < active_index:
                icon.create_oval(
                    size * .10, size * .10, size * .90, size * .90,
                    fill=Theme.LOG_SUCCESS,
                    outline=''
                )
                icon.create_line(
                    size * .27, size * .50,
                    size * .44, size * .67,
                    size * .74, size * .31,
                    fill='white',
                    width=max(4, ui_px(2.2)),
                    capstyle='round',
                    joinstyle='round'
                )
                label.configure(
                    fg=Theme.TEXT_MUTED,
                    font=(Theme.FONT_SERIF[0], 10, 'normal', 'overstrike')
                )
            elif idx == active_index:
                icon.create_oval(
                    size * .10, size * .10, size * .90, size * .90,
                    fill=Theme.PRIMARY_LIGHT,
                    outline=Theme.PRIMARY,
                    width=line_width
                )
                icon.create_oval(
                    size * .36, size * .36, size * .64, size * .64,
                    fill=Theme.PRIMARY,
                    outline=''
                )
                label.configure(fg=Theme.TEXT, font=get_font(10, 'bold'))
            else:
                icon.create_oval(
                    size * .14, size * .14, size * .86, size * .86,
                    fill=Theme.CARD,
                    outline=Theme.BORDER,
                    width=line_width
                )
                label.configure(fg=Theme.TEXT_SECONDARY, font=get_font(10))

    def stop(self):
        self.pack_forget()


class ResultPanel(tk.Frame):
    """结果反馈面板"""
    
    def __init__(self, parent, **kwargs):
        super().__init__(parent, bg=Theme.BG, **kwargs)
        
        # 占位状态
        self.placeholder = tk.Label(
            self,
            text="处理结果将在此处显示",
            font=get_font(12),
            bg=Theme.BG,
            fg=Theme.TEXT_MUTED,
            pady=Theme.SPACE_XL
        )
        self.placeholder.pack()
        
        # 结果卡片
        self.result_card = tk.Frame(self, bg=Theme.CARD, highlightbackground=Theme.BORDER, highlightthickness=1)
        self.result_content = tk.Frame(self.result_card, bg=Theme.CARD)
        self.result_content.pack(fill='both', expand=True, padx=Theme.SPACE_LG, pady=Theme.SPACE_LG)
        self._wrapped_labels = []
        self.bind('<Configure>', self._on_configure)
        self.result_content.bind('<Configure>', self._on_configure)

    def _wraplength(self):
        width = self.result_content.winfo_width()
        if width < 120:
            width = self.winfo_width()
        if width < 120:
            width = self.winfo_toplevel().winfo_width() - 120
        return max(360, width - 12)

    def _register_wrapped(self, label):
        self._wrapped_labels.append(label)
        label.configure(wraplength=self._wraplength())
        return label

    def _clear_result_content(self):
        for widget in self.result_content.winfo_children():
            widget.destroy()
        self._wrapped_labels = []

    def _on_configure(self, event=None):
        wrap = self._wraplength()
        for label in list(getattr(self, '_wrapped_labels', [])):
            try:
                label.configure(wraplength=wrap)
            except tk.TclError:
                self._wrapped_labels.remove(label)
    
    def show_success(self, message, filepath=None):
        self.placeholder.pack_forget()
        
        self._clear_result_content()
        
        # 成功图标 + 消息
        header = tk.Frame(self.result_content, bg=Theme.CARD)
        header.pack(fill='x', anchor='w')
        
        icon_canvas = tk.Canvas(header, width=36, height=36, bg=Theme.CARD, highlightthickness=0)
        icon_canvas.pack(side='left')
        Icons.draw_check(icon_canvas, 2, 2, 32, Theme.LOG_SUCCESS)
        
        tk.Label(
            header,
            text=message,
            font=get_font(15, 'bold'),
            bg=Theme.CARD,
            fg=Theme.TEXT,
            anchor='w'
        ).pack(side='left', padx=(Theme.SPACE_SM, 0))
        
        if filepath:
            self._register_wrapped(tk.Label(
                self.result_content,
                text=f"输出文件：{filepath}",
                font=get_font(11),
                bg=Theme.CARD,
                fg=Theme.TEXT_SECONDARY,
                anchor='w',
                justify='left',
            )).pack(fill='x', anchor='w', pady=(Theme.SPACE_SM, 0))
        
        self.result_card.pack(fill='x', pady=(Theme.SPACE_MD, 0))
    
    def show_diagnosis(self, results):
        self.placeholder.pack_forget()
        
        self._clear_result_content()
        
        tk.Label(
            self.result_content,
            text="诊断报告",
            font=get_font(15, 'bold'),
            bg=Theme.CARD,
            fg=Theme.TEXT,
            anchor='w'
        ).pack(fill='x', anchor='w', pady=(0, Theme.SPACE_MD))
        
        total = 0
        categories = [
            ('标点问题', results.get('punctuation', [])),
            ('序号问题', results.get('numbering', [])),
            ('段落问题', results.get('paragraph', [])),
            ('字体问题', results.get('font', [])),
        ]
        
        for name, issues in categories:
            count = len(issues)
            total += count
            
            row = tk.Frame(self.result_content, bg=Theme.CARD)
            row.pack(fill='x', pady=3)
            
            tk.Label(
                row,
                text=name,
                font=get_font(12),
                bg=Theme.CARD,
                fg=Theme.TEXT,
                width=10,
                anchor='w'
            ).pack(side='left')
            
            count_color = Theme.LOG_WARNING if count > 0 else Theme.LOG_SUCCESS
            tk.Label(
                row,
                text=f"{count} 处" if count > 0 else "无问题",
                font=get_font(12),
                bg=Theme.CARD,
                fg=count_color,
                anchor='w'
            ).pack(side='left')
        
        tk.Frame(self.result_content, bg=Theme.BORDER, height=1).pack(fill='x', pady=Theme.SPACE_MD)
        
        summary_color = Theme.LOG_SUCCESS if total == 0 else Theme.LOG_WARNING
        summary_text = "文档格式规范，未发现问题" if total == 0 else f"共发现 {total} 处格式问题"
        
        tk.Label(
            self.result_content,
            text=summary_text,
            font=get_font(13, 'bold'),
            bg=Theme.CARD,
            fg=summary_color,
            anchor='w'
        ).pack(fill='x', anchor='w')
        
        self.result_card.pack(fill='x', pady=(Theme.SPACE_MD, 0))

    def reset(self):
        self.result_card.pack_forget()
        self._clear_result_content()
        self.placeholder.pack()


class DocFormatApp:
    def __init__(self, root):
        self.root = root
        self.root.title("公文格式处理工具")
        _init_system_fonts()
        self._set_initial_window_geometry()
        self.root.configure(bg=Theme.BG)
        self._scrollbar_drag_offset = None
        
        # 变量
        self.input_file = tk.StringVar()
        self.output_file = tk.StringVar()
        self.operation = tk.StringVar(value="smart")
        self.preset = tk.StringVar(value="official")
        self.input_files = []   # 多文件模式下存储路径列表
        
        self.preset_cards = []
        
        self.create_widgets()
        self.root._on_file_selected = self._on_file_selected

    def _set_initial_window_geometry(self):
        """让主窗口默认以更宽的工作区打开，同时兼容小屏和高 DPI。"""
        try:
            screen_w = self.root.winfo_screenwidth()
            screen_h = self.root.winfo_screenheight()
        except Exception:
            screen_w, screen_h = 1366, 768

        width = min(screen_w - 80, max(1120, int(screen_w * 0.72)))
        height = min(screen_h - 80, max(780, int(screen_h * 0.82)))
        width = max(980, width)
        height = max(700, height)
        width = min(width, max(760, screen_w - 40))
        height = min(height, max(640, screen_h - 40))
        x = max(0, (screen_w - width) // 2)
        y = max(0, (screen_h - height) // 2)

        self.root.geometry(f"{width}x{height}+{x}+{y}")
        self.root.minsize(min(960, width), min(700, height))
    
    def create_widgets(self):
        """构建界面"""
        # 主容器 - 带滚动
        container = tk.Frame(self.root, bg=Theme.BG)
        container.pack(fill='both', expand=True)
        
        # Canvas + 自定义滚动条
        self.canvas = tk.Canvas(container, bg=Theme.BG, highlightthickness=0)
        self.scrollbar_canvas = tk.Canvas(
            container, width=ui_px(20), bg=Theme.BG,
            highlightthickness=0, cursor='sb_v_double_arrow'
        )
        
        self.canvas.pack(side='left', fill='both', expand=True)
        self.scrollbar_canvas.pack(side='right', fill='y')
        
        # 内容Frame
        self.main_frame = tk.Frame(self.canvas, bg=Theme.BG)
        self.canvas_window = self.canvas.create_window((0, 0), window=self.main_frame, anchor='nw')
        
        # 绑定滚动
        self.canvas.configure(yscrollcommand=self._on_canvas_yview)
        self.main_frame.bind('<Configure>', self._on_frame_configure)
        self.canvas.bind('<Configure>', self._on_canvas_configure)
        self.root.bind_all('<MouseWheel>', self._on_mousewheel, add='+')
        self.root.bind_all('<Button-4>', self._on_mousewheel, add='+')
        self.root.bind_all('<Button-5>', self._on_mousewheel, add='+')
        self.scrollbar_canvas.bind('<Button-1>', self._on_scrollbar_click)
        self.scrollbar_canvas.bind('<B1-Motion>', self._on_scrollbar_drag)
        self.scrollbar_canvas.bind('<ButtonRelease-1>', self._on_scrollbar_release)
        
        # 内容区域
        content = tk.Frame(self.main_frame, bg=Theme.BG)
        content.pack(fill='both', expand=True, padx=Theme.SPACE_XL, pady=Theme.SPACE_LG)
        
        # ===== 1. 头部 =====
        tk.Label(
            content,
            text="公文格式处理工具",
            font=get_font(24, 'bold'),
            bg=Theme.BG,
            fg=Theme.TEXT
        ).pack(anchor='w', pady=(0, Theme.SPACE_XL))
        
        # ===== 2. 文件选择区 =====
        file_section = tk.Frame(content, bg=Theme.BG)
        file_section.pack(fill='x', pady=(0, Theme.SPACE_LG))
        if _DND_AVAILABLE:
            placeholder_text = "点击选择文件，或将文件拖到这里"
        elif os.name == 'nt' and _DND_DISABLED_REASON:
            if "管理员" in _DND_DISABLED_REASON:
                placeholder_text = "点击选择文件（管理员模式下拖拽不可用）"
            else:
                placeholder_text = "点击选择文件（拖拽不可用）"
        else:
            placeholder_text = "点击选择文件"
        
        self.input_field = FileInputField(
            file_section,
            label_text="输入",
            placeholder=placeholder_text,
            variable=self.input_file,
            command=self.browse_input,
            drop_command=self._on_file_selected
        )
        self.input_field.pack(fill='x', pady=(0, Theme.SPACE_SM))

        # v1.8.0: 粘贴文本生成 docx 入口
        paste_row = tk.Frame(file_section, bg=Theme.BG)
        paste_row.pack(fill='x', pady=(0, Theme.SPACE_SM))

        tk.Frame(paste_row, bg=Theme.BG, width=40).pack(side='left')  # 占位对齐

        paste_btn = tk.Label(
            paste_row,
            text="📋 没有文件？粘贴文本生成 docx",
            font=get_font(11),
            bg=Theme.BG, fg=Theme.PRIMARY,
            cursor='hand2', anchor='w', padx=4,
        )
        paste_btn.pack(side='left')
        paste_btn.bind('<Button-1>', lambda e: self._open_paste_dialog())
        paste_btn.bind('<Enter>', lambda e: paste_btn.configure(fg=Theme.PRIMARY_HOVER))
        paste_btn.bind('<Leave>', lambda e: paste_btn.configure(fg=Theme.PRIMARY))

        folder_row = tk.Frame(file_section, bg=Theme.BG)
        folder_row.pack(fill='x', pady=(0, Theme.SPACE_SM))
        tk.Button(
            folder_row, text='📁 选择文件夹',
            command=self._on_select_folder,
            bg=Theme.BG, fg=Theme.TEXT_SECONDARY,
            font=get_font(11), relief='flat', cursor='hand2',
            padx=Theme.SPACE_SM, pady=4,
            borderwidth=1, highlightbackground=Theme.BORDER,
        ).pack(side='left', padx=4)
        
        self.output_field = FileInputField(
            file_section,
            label_text="输出",
            placeholder="文档修改后的储存位置",
            variable=self.output_file,
            command=self.browse_output
        )
        self.output_field.pack(fill='x')
        
        # 分隔
        tk.Frame(content, bg=Theme.BORDER, height=1).pack(fill='x', pady=Theme.SPACE_LG)
        
        # ===== 3. 功能选择区 =====
        mode_section = tk.Frame(content, bg=Theme.BG)
        mode_section.pack(fill='x', pady=(0, Theme.SPACE_LG))
        
        # 大卡片 - 智能一键处理
        smart_card = SelectableCard(
            mode_section,
            title="智能一键处理",
            description="自动修复标点符号，并应用标准格式规范，一步到位完成文档处理",
            value="smart",
            variable=self.operation,
            icon_draw_func=Icons.draw_magic,
            is_featured=True,
            command=self._on_mode_change
        )
        smart_card.pack(fill='x', pady=(0, Theme.SPACE_MD))

        # 两个小卡片
        small_cards = tk.Frame(mode_section, bg=Theme.BG)
        small_cards.pack(fill='x')
        small_cards.columnconfigure(0, weight=1)
        small_cards.columnconfigure(1, weight=1)
        
        diag_card = SelectableCard(
            small_cards,
            title="格式诊断",
            description="仅分析文档问题，不修改文件",
            value="analyze",
            variable=self.operation,
            icon_draw_func=Icons.draw_search,
            command=self._on_mode_change
        )
        diag_card.grid(row=0, column=0, sticky='nsew', padx=(0, Theme.SPACE_SM))
        
        punct_card = SelectableCard(
            small_cards,
            title="标点修复",
            description="仅修复中英文标点混用",
            value="punctuation",
            variable=self.operation,
            icon_draw_func=Icons.draw_edit,
            command=self._on_mode_change
        )
        punct_card.grid(row=0, column=1, sticky='nsew')

        # ===== 4. 格式预设 =====
        preset_section = tk.Frame(content, bg=Theme.BG)
        preset_section.pack(fill='x', pady=(0, Theme.SPACE_LG))
        
        # 标题行
        tk.Label(
            preset_section,
            text="格式预设",
            font=get_font(12),
            bg=Theme.BG,
            fg=Theme.TEXT_SECONDARY
        ).pack(anchor='w', pady=(0, Theme.SPACE_SM))
        
        preset_row = tk.Frame(preset_section, bg=Theme.BG)
        preset_row.pack(fill='x')
        
        presets = [
            ('official', 'GB/T 公文标准'),
            ('academic', '学术论文'),
            ('legal', '法律文书'),
        ]
        
        for i, (value, text) in enumerate(presets):
            card = PresetCard(preset_row, text, value, self.preset)
            card.pack(side='left', padx=(0 if i == 0 else Theme.SPACE_SM, 0))
            self.preset_cards.append(card)
        
        # 自定义卡片 - 点击直接打开设置窗口
        self.custom_card = PresetCard(
            preset_row, '⚙️ 自定义', 'custom', self.preset,
            command=self._open_custom_settings  # 点击时打开设置窗口
        )
        self.custom_card.pack(side='left', padx=(Theme.SPACE_SM, 0))
        self.preset_cards.append(self.custom_card)
        
        # ===== 修订标记开关 =====
        revision_row = tk.Frame(content, bg=Theme.BG)
        revision_row.pack(fill='x', pady=(0, Theme.SPACE_SM))

        self.revision_mode_var = tk.BooleanVar(value=False)
        self.revision_cb = InlineToggle(
            revision_row,
            text="输出修订标记（在 Word 中可逐条接受 / 拒绝格式更改）",
            variable=self.revision_mode_var,
        )
        self.revision_cb.pack(side='left')
        # ⓘ 说明按钮
        info_btn = tk.Label(
            revision_row,
            text=" ⓘ ",
            font=get_font(11),
            bg=Theme.BG,
            fg=Theme.TEXT_MUTED,
            cursor='hand2',
        )
        info_btn.pack(side='left')
        info_btn.bind('<Enter>', lambda e: info_btn.configure(fg=Theme.PRIMARY))
        info_btn.bind('<Leave>', lambda e: info_btn.configure(fg=Theme.TEXT_MUTED))
        info_btn.bind('<Button-1>', lambda e: self._show_revision_info(e))
        # ===== 修订标记开关结束 =====

        # ===== 完成后自动打开文件开关 =====
        auto_open_row = tk.Frame(content, bg=Theme.BG)
        auto_open_row.pack(fill='x', pady=(0, Theme.SPACE_SM))
        self.auto_open_var = tk.BooleanVar(value=False)
        self.auto_open_cb = InlineToggle(
            auto_open_row,
            text="处理完成后自动打开输出文件",
            variable=self.auto_open_var,
        )
        self.auto_open_cb.pack(side='left')

        
        # ===== 5. 执行按钮 =====
        self.run_btn = tk.Frame(content, bg=Theme.PRIMARY, cursor='hand2')
        self.run_btn.pack(fill='x', pady=Theme.SPACE_LG)
        
        self.run_label = tk.Label(
            self.run_btn,
            text="开始处理",
            font=get_font(15, 'bold'),
            bg=Theme.PRIMARY,
            fg='white',
            pady=Theme.SPACE_MD + 2
        )
        self.run_label.pack()
        
        for widget in [self.run_btn, self.run_label]:
            widget.bind('<Button-1>', lambda e: self.run_operation())
            widget.bind('<Enter>', lambda e: self._btn_hover(True))
            widget.bind('<Leave>', lambda e: self._btn_hover(False))
        self.run_label.configure(cursor='hand2')
        
        # ===== 5.5 运行状态 =====
        self.status_section = tk.Frame(content, bg=Theme.BG)
        self.status_section.pack(fill='x', pady=(0, Theme.SPACE_SM))

        self.task_list = ProcessingTaskList(self.status_section)
        self.processing_panel = ProcessingPanel(self.status_section)

        self.progress_frame = tk.Frame(self.status_section, bg=Theme.BG)
        # 默认不 pack，处理时才显示
        
        self.progress_stage = tk.Label(
            self.progress_frame,
            text="",
            font=get_font(10),
            bg=Theme.BG,
            fg=Theme.TEXT_SECONDARY,
            anchor='w'
        )
        self.progress_stage.pack(fill='x', pady=(0, 4))
        
        progress_bar_bg = tk.Frame(self.progress_frame, bg=Theme.BORDER, height=8)
        progress_bar_bg.pack(fill='x')
        progress_bar_bg.pack_propagate(False)
        
        self.progress_bar_fill = tk.Frame(progress_bar_bg, bg=Theme.PRIMARY, height=8, width=0)
        self.progress_bar_fill.place(x=0, y=0, relheight=1.0, relwidth=0.0)
        
        self.progress_pct = tk.Label(
            self.progress_frame,
            text="",
            font=get_font(9),
            bg=Theme.BG,
            fg=Theme.TEXT_MUTED,
            anchor='e'
        )
        self.progress_pct.pack(fill='x', pady=(2, 0))
        
        # ===== 6. 结果反馈区 =====
        self.result_panel = ResultPanel(content)
        self.result_panel.pack(fill='x', pady=(0, Theme.SPACE_LG))
        
        # ===== 7. 日志区 =====
        self.log_panel = CollapsibleLog(content)
        self.log_panel.pack(fill='x', pady=(Theme.SPACE_MD, 0))
        
        # ===== 8. 底部版权信息 =====
        footer = tk.Frame(content, bg=Theme.BG)
        footer.pack(fill='x', pady=(Theme.SPACE_LG, Theme.SPACE_SM))
        
        tk.Label(
            footer,
            text="© 2026 今日之明日 · MIT License",
            font=get_font(9),
            bg=Theme.BG,
            fg=Theme.TEXT_MUTED
        ).pack(side='left')
        
        about_label = tk.Label(
            footer,
            text="关于",
            font=get_font(9),
            bg=Theme.BG,
            fg=Theme.TEXT_MUTED,
            cursor='hand2'
        )
        about_label.pack(side='right')
        about_label.bind('<Button-1>', lambda e: self._show_about())
        about_label.bind('<Enter>', lambda e: about_label.configure(fg=Theme.PRIMARY))
        about_label.bind('<Leave>', lambda e: about_label.configure(fg=Theme.TEXT_MUTED))
        
        # 初始化
        self._on_mode_change()
        self.log_panel.log("工具已就绪，请选择文件", 'info')
        if _DND_DISABLED_REASON:
            hint = "如需拖拽，请用普通权限启动程序。" if "管理员" in _DND_DISABLED_REASON else "仍可点击选择文件。"
            self.log_panel.log(f"拖拽提示：{_DND_DISABLED_REASON}。{hint}", 'warning')
    
    def _show_about(self):
        """显示关于对话框"""
        about_text = (
            f"公文格式处理工具  v{__version__}\n\n"
            "一键将 Word 文档排版为标准公文格式\n\n"
            "开发者：KaguraNanaga\n"
            "许可证：MIT License\n"
            "项目地址：github.com/KaguraNanaga/docformat-gui\n\n"
            "─────────────────────\n\n"
            "🔒 数据安全声明\n\n"
            "所有文档处理均在本地完成。\n"
            "本工具不联网、不上传、不收集任何数据。\n\n"
            "─────────────────────\n\n"
            "⚠ 免责声明\n\n"
            "处理结果仅供参考，建议人工复核。\n"
            "本软件按「原样」提供，不附带任何担保。\n"
            "详见项目目录下 DISCLAIMER.md。"
        )
        messagebox.showinfo("关于", about_text)
    
    def _show_progress(self):
        """显示进度条"""
        self.progress_frame.pack(fill='x', pady=(0, Theme.SPACE_SM))
        self._update_progress(0, 100, '准备中...')

    def _show_processing(self, title=None):
        """显示运行中的动态状态面板。"""
        if hasattr(self, 'processing_panel'):
            self.processing_panel.start(title=title)

    def _show_task_list(self, mode, has_input_conversion=False, has_output_conversion=False):
        """显示运行任务清单。"""
        if hasattr(self, 'task_list'):
            self.task_list.start(
                mode,
                has_input_conversion=has_input_conversion,
                has_output_conversion=has_output_conversion
            )
    
    def _update_progress(self, current, total, stage_text):
        """更新进度条（线程安全）"""
        import threading
        if threading.current_thread() is not threading.main_thread():
            self.root.after(0, lambda: self._update_progress(current, total, stage_text))
            return
        
        pct = int(current * 100 / total) if total > 0 else 0
        pct = min(pct, 100)
        self.progress_bar_fill.place(x=0, y=0, relheight=1.0, relwidth=pct / 100)
        self.progress_stage.configure(text=stage_text)
        self.progress_pct.configure(text=f"{pct}%")
        if hasattr(self, 'task_list'):
            self.task_list.update_status(pct, stage_text)
        if hasattr(self, 'processing_panel') and self.processing_panel.running:
            self.processing_panel.update_status(pct, stage_text)
    
    def _hide_progress(self):
        """隐藏进度条"""
        self.progress_frame.pack_forget()
        if hasattr(self, 'processing_panel'):
            self.processing_panel.stop()
        if hasattr(self, 'task_list'):
            self.task_list.stop()
    
    def _on_frame_configure(self, event):
        self.canvas.configure(scrollregion=self.canvas.bbox('all'))
        self._draw_scrollbar()
    
    def _on_canvas_configure(self, event):
        self.canvas.itemconfig(self.canvas_window, width=event.width)
        self._draw_scrollbar()

    def _on_canvas_yview(self, first, last):
        self._draw_scrollbar()

    def _can_scroll(self):
        try:
            top, bottom = self.canvas.yview()
            return bottom - top < 0.99
        except Exception:
            return False
    
    def _on_mousewheel(self, event):
        try:
            if str(event.widget.winfo_toplevel()) != str(self.root):
                return
        except Exception:
            return
        if not self._can_scroll():
            return "break"

        if getattr(event, 'num', None) == 4:
            units = -3
        elif getattr(event, 'num', None) == 5:
            units = 3
        else:
            delta = getattr(event, 'delta', 0)
            if delta == 0:
                units = 0
            else:
                units = -int(delta / 120)
                if units == 0:
                    units = -1 if delta > 0 else 1
                units *= 3
        if units:
            self.canvas.yview_scroll(units, 'units')
        self._draw_scrollbar()
        return "break"
    
    def _draw_scrollbar(self):
        """绘制自定义滚动条"""
        self.scrollbar_canvas.delete('all')
        
        try:
            top, bottom = self.canvas.yview()
        except:
            return
        
        if bottom - top >= 0.99:
            self._scrollbar_thumb = None
            return
        
        w = 20
        h = self.scrollbar_canvas.winfo_height()
        
        if h < 10:
            return
        
        track_top = 8
        track_bottom = max(track_top + 24, h - 8)
        track_h = track_bottom - track_top
        bar_h = max(44, (bottom - top) * track_h)
        bar_h = min(bar_h, track_h)
        bar_y = track_top + top * (track_h - bar_h)
        self._scrollbar_thumb = (bar_y, bar_h, track_top, track_h)
        
        # 轨道
        self.scrollbar_canvas.create_rectangle(
            5, track_top, w - 5, track_bottom,
            fill='#E8E4DE', outline='', tags='track'
        )
        
        # 滑块（更深的颜色）
        self.scrollbar_canvas.create_rectangle(
            4, bar_y, w - 4, bar_y + bar_h,
            fill='#8E867E', outline='', tags='thumb'
        )
    
    def _on_scrollbar_click(self, event):
        """滚动条点击"""
        if not self._can_scroll():
            return "break"
        try:
            thumb = getattr(self, '_scrollbar_thumb', None)
            if not thumb:
                return "break"
            bar_y, bar_h, _track_top, _track_h = thumb
            if bar_y <= event.y <= bar_y + bar_h:
                self._scrollbar_drag_offset = event.y - bar_y
            else:
                self._scrollbar_drag_offset = bar_h / 2
                direction = -1 if event.y < bar_y else 1
                self.canvas.yview_scroll(direction, 'pages')
            self._draw_scrollbar()
        except Exception:
            return "break"
        return "break"
    
    def _on_scrollbar_drag(self, event):
        """滚动条拖动"""
        if not self._can_scroll():
            return "break"
        try:
            thumb = getattr(self, '_scrollbar_thumb', None)
            if not thumb:
                return "break"
            bar_y, bar_h, track_top, track_h = thumb
            offset = self._scrollbar_drag_offset
            if offset is None:
                offset = bar_h / 2
            available = max(1, track_h - bar_h)
            fraction = (event.y - offset - track_top) / available
            fraction = max(0.0, min(1.0, fraction))
            self.canvas.yview_moveto(fraction)
            self._draw_scrollbar()
        except Exception:
            return "break"
        return "break"

    def _on_scrollbar_release(self, event):
        self._scrollbar_drag_offset = None
        return "break"
    
    def _btn_hover(self, is_hover):
        color = Theme.PRIMARY_HOVER if is_hover else Theme.PRIMARY
        self.run_btn.configure(bg=color)
        self.run_label.configure(bg=color)
    
    def _on_mode_change(self):
        mode = self.operation.get()
        enabled = mode in ('smart',)
        for card in self.preset_cards:
            card.set_enabled(enabled)
        # 修订标记仅对"智能一键处理"有效
        if hasattr(self, 'revision_cb'):
            if mode == 'smart':
                self.revision_cb.configure(state='normal', fg=Theme.TEXT_SECONDARY)
            else:
                self.revision_mode_var.set(False)
                self.revision_cb.configure(state='disabled', fg=Theme.TEXT_MUTED)

    def _show_revision_info(self, event):
        """弹出修订标记说明小窗"""
        popup = tk.Toplevel(self.root)
        popup.overrideredirect(True)   # 无标题栏
        popup.configure(bg=Theme.CARD)

        # 描边效果：外层 Frame
        border = tk.Frame(popup, bg=Theme.BORDER, padx=1, pady=1)
        border.pack()

        inner = tk.Frame(border, bg=Theme.CARD)
        inner.pack()

        msg = (
            "修订标记仅记录以下格式变更：\n"
            "  · 字体、字号\n"
            "  · 首行缩进、对齐方式\n"
            "  · 行距、段前段后间距\n\n"
            "以下改动不会显示为修订标记：\n"
            "  · 标点符号替换\n"
            "  · 空格删除"
        )
        tk.Label(
            inner, text=msg,
            font=get_font(10),
            bg=Theme.CARD, fg=Theme.TEXT,
            justify='left',
            padx=14, pady=10,
        ).pack()

        # 定位到点击位置附近
        popup.update_idletasks()
        x = event.x_root + 10
        y = event.y_root + 10
        # 防止超出屏幕右边
        sw = popup.winfo_screenwidth()
        pw = popup.winfo_reqwidth()
        if x + pw > sw - 10:
            x = sw - pw - 10
        popup.geometry(f'+{x}+{y}')

        # 点击任意处关闭
        popup.bind('<Button-1>', lambda e: popup.destroy())
        popup.focus_set()
        popup.bind('<FocusOut>', lambda e: popup.destroy())
    
    def _open_custom_settings(self):
        """打开自定义设置窗口"""
        def on_save(settings):
            self.preset.set('custom')
            self.log_panel.log("自定义格式设置已保存", 'success')
        
        CustomSettingsDialog(self.root, on_save=on_save)

    def _open_paste_dialog(self):
        """打开粘贴文本对话框。"""
        PasteTextDialog(self.root, on_generate=self._on_text_generated)

    def _on_text_generated(self, title, body_text, output_path, is_markdown=False):
        """粘贴对话框生成 docx 后的回调：触发主流程格式化。"""
        import tempfile

        self.log_panel.log(f"\n{'─' * 35}", 'info')
        self.log_panel.log(f"从文本生成 docx：{title}", 'info')

        # 1) 先生成临时 docx 作为 format_document 的输入
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                temp_input = tmp.name
            if is_markdown:
                _create_docx_from_markdown(title, body_text, temp_input)
                self.log_panel.log("  临时文件已生成（Markdown 格式）", 'info')
            else:
                _create_docx_from_text(title, body_text, temp_input)
                self.log_panel.log("  临时文件已生成（纯文本格式）", 'info')
        except Exception as e:
            messagebox.showerror('生成失败', f'创建临时 docx 失败：{e}')
            self.log_panel.log(f"  生成临时文件失败: {e}", 'error')
            return

        # 2) 设置输入输出路径，触发现有的 run_operation 流程
        self.input_files.clear()
        self.input_files.append(temp_input)
        self.input_file.set(temp_input)
        self.output_file.set(output_path)

        # 显示在输入框（用更友好的名称）
        self.input_field.filename_label.configure(
            text=f"📋 文本生成: {title[:30]}",
            fg=Theme.TEXT,
        )
        self.output_field.filename_label.configure(
            text=Path(output_path).name,
            fg=Theme.TEXT,
        )

        # 强制使用 smart 模式（粘贴的文本必走标点修复+格式化全流程）
        self.operation.set('smart')
        self._on_mode_change()

        # 自动开始处理
        self.log_panel.log("  开始格式化处理...", 'info')
        self._pending_temp_input = temp_input
        self.run_operation()
    
    def _on_file_selected(self, filename):
        """文件选定后的统一处理（点击选择和拖拽共用）"""
        if not filename:
            return
        self._add_files_to_list([filename])

    def _add_files_to_list(self, filenames):
        """把一组文件加入批量处理列表，并刷新输入/输出显示。"""
        if not filenames:
            return

        filenames = [str(p) for p in filenames]

        # 用 clear+extend 原地修改，避免外部引用失效
        self.input_files.clear()
        self.input_files.extend(filenames)

        first = Path(filenames[0])

        if len(filenames) == 1:
            self.input_file.set(str(filenames[0]))
            output_name = f"{first.stem}_processed{first.suffix}"
            self.output_file.set(str(first.parent / output_name))
            self.log_panel.log(f"已选择: {first.name}", 'info')
        else:
            # 多文件：手动更新显示，输出设为第一个文件的目录
            self.input_file.set(str(filenames[0]))
            self.input_field.filename_label.configure(
                text=f"已选择 {len(filenames)} 个文件  ({first.name} ...)",
                fg=Theme.TEXT
            )
            self.output_file.set(str(first.parent))
            self.output_field.filename_label.configure(
                text=f"输出目录: {first.parent.name}",
                fg=Theme.TEXT
            )
            self.log_panel.log(f"已选择 {len(filenames)} 个文件，输出目录: {first.parent}", 'info')

        self.result_panel.reset()

    def _on_select_folder(self):
        """选择一个文件夹，递归扫描 .doc/.docx/.wps 文件作为批量输入。"""
        from tkinter import filedialog, messagebox
        from pathlib import Path

        folder = filedialog.askdirectory(
            title='选择包含 Word 文档的文件夹（会递归搜索子文件夹）',
            parent=self.root,
        )
        if not folder:
            return

        folder_path = Path(folder)
        extensions = {'.docx', '.doc', '.wps'}
        found = []
        for p in folder_path.rglob('*'):
            if p.is_file() and p.suffix.lower() in extensions and not p.name.startswith('~$'):
                found.append(str(p))

        found.sort(key=lambda x: x.lower())

        if not found:
            messagebox.showinfo(
                '没找到文件',
                f'文件夹「{folder_path.name}」下没有 Word 文档',
                parent=self.root,
            )
            return

        # 简单确认（防止误选大文件夹）
        if len(found) > 30:
            if not messagebox.askyesno(
                '确认',
                f'共找到 {len(found)} 个 Word 文档。是否全部加入处理列表？',
                parent=self.root,
            ):
                return

        self._add_files_to_list(found)
    
    def browse_input(self):
        is_windows = (os.name == 'nt')
        if is_windows:
            filetypes = [
                ("所有支持格式", "*.docx *.doc *.wps"),
                ("Word 文档 (.docx)", "*.docx"),
                ("Word 97-2003 (.doc)", "*.doc"),
                ("WPS 文档 (.wps)", "*.wps"),
                ("所有文件", "*.*"),
            ]
        else:
            filetypes = [
                ("Word 文档 (.docx)", "*.docx"),
                ("所有文件", "*.*"),
            ]

        filenames = filedialog.askopenfilenames(
            title="选择Word文档（可多选）",
            filetypes=filetypes
        )
        if not filenames:
            return

        self._add_files_to_list(filenames)
        if len(filenames) == 1:
            self.log_panel.log(f"输出格式已自动设置为: {Path(filenames[0]).suffix or '.docx'}", 'info')
    
    def browse_output(self):
        # 多文件模式：选择输出目录
        if len(self.input_files) > 1:
            directory = filedialog.askdirectory(
                title="选择输出目录",
                initialdir=str(Path(self.input_files[0]).parent)
            )
            if directory:
                self.output_file.set(directory)
                self.output_field.filename_label.configure(
                    text=f"输出目录: {Path(directory).name}",
                    fg=Theme.TEXT
                )
                self.log_panel.log(f"输出目录: {directory}", 'info')
            return

        # 单文件模式：原有行为
        is_windows = (os.name == 'nt')
        if is_windows:
            filetypes = [
                ("所有支持格式", "*.docx *.doc *.wps"),
                ("Word 文档 (.docx)", "*.docx"),
                ("Word 97-2003 (.doc)", "*.doc"),
                ("WPS 文档 (.wps)", "*.wps"),
            ]
        else:
            filetypes = [("Word 文档 (.docx)", "*.docx")]

        default_name = ""
        if self.input_files:
            p = Path(self.input_files[0])
            default_name = f"{p.stem}_processed{p.suffix}"

        filename = filedialog.asksaveasfilename(
            title="保存为",
            defaultextension=".docx",
            filetypes=filetypes,
            initialfile=default_name
        )
        if filename:
            self.output_file.set(filename)
    
    def run_operation(self):
        # 确定输入列表
        if self.input_files:
            input_paths = self.input_files[:]
        else:
            p = self.input_file.get().strip()
            if not p:
                messagebox.showerror("提示", "请先选择输入文件")
                return
            input_paths = [p]

        output_base = self.output_file.get().strip()
        mode = self.operation.get()

        # 验证文件存在
        for p in input_paths:
            if not os.path.exists(p):
                messagebox.showerror("错误", f"文件不存在:\n{p}")
                return

        # Linux 格式限制
        if os.name != 'nt':
            for p in input_paths:
                if Path(p).suffix.lower() in ('.doc', '.wps'):
                    messagebox.showerror(
                        "不支持的格式",
                        "当前系统仅支持 .docx 文件。\n.doc/.wps 请先另存为 .docx。"
                    )
                    return

        if mode != 'analyze' and not output_base:
            messagebox.showerror("提示", "请指定输出文件或目录")
            return

        # 诊断模式仅支持单文件
        if mode == 'analyze' and len(input_paths) > 1:
            messagebox.showwarning("提示", "诊断模式每次仅处理一个文件，将只分析第一个文件。")
            input_paths = input_paths[:1]

        self.run_btn.configure(bg=Theme.TEXT_MUTED)
        self.run_label.configure(bg=Theme.TEXT_MUTED, text="处理中...")
        self._show_progress()

        rev_mode = self.revision_mode_var.get() if hasattr(self, 'revision_mode_var') else False
        if mode == 'analyze':
            processing_title = "正在诊断文档"
        elif mode == 'punctuation':
            processing_title = "正在修复标点"
        elif mode == 'format':
            processing_title = "正在应用格式"
        else:
            processing_title = "正在处理文档"
        self._show_processing(title=processing_title)

        legacy_exts = ('.doc', '.wps')
        has_input_conversion = any(Path(p).suffix.lower() in legacy_exts for p in input_paths)
        if mode == 'analyze':
            has_output_conversion = False
        elif len(input_paths) > 1 or os.path.isdir(output_base):
            has_output_conversion = has_input_conversion
        else:
            has_output_conversion = Path(output_base).suffix.lower() in legacy_exts

        self._show_task_list(
            mode,
            has_input_conversion=has_input_conversion,
            has_output_conversion=has_output_conversion
        )
        thread = threading.Thread(
            target=self._do_operation,
            args=(input_paths, output_base, mode, rev_mode)
        )
        thread.start()
    
    def _do_operation(self, input_paths, output_base, mode, revision_mode=False):
        """批量调度器：循环处理每个文件，汇总结果。"""
        total = len(input_paths)
        success_paths = []
        failed_files = []

        try:
            self.log_panel.log(f"\n{'─' * 35}", 'info')
            self.log_panel.log(
                f"开始处理 {total} 个文件..." if total > 1 else f"开始处理: {Path(input_paths[0]).name}",
                'info'
            )

            for idx, input_path in enumerate(input_paths):
                offset = int(idx / total * 100)
                per_range = int(1 / total * 100) or 1

                def make_progress_fn(off, rng):
                    return lambda pct, text: self._update_progress(
                        off + pct * rng // 100, 100, text
                    )
                progress_fn = make_progress_fn(offset, per_range)

                if total > 1 or os.path.isdir(output_base):
                    out_dir = Path(output_base)
                    in_p = Path(input_path)
                    out_path = str(out_dir / f"{in_p.stem}_processed{in_p.suffix}")
                else:
                    out_path = output_base

                self.log_panel.log(
                    f"\n[{idx + 1}/{total}] {Path(input_path).name}", 'info'
                )
                try:
                    actual_out, summary = self._process_single_file(
                        input_path, out_path, mode, progress_fn,
                        revision_mode=revision_mode,
                    )
                    success_paths.append(actual_out)
                    self.log_panel.log(
                        f"  ✓ 已保存: {Path(actual_out).name}", 'success'
                    )
                except Exception as e:
                    self.log_panel.log(f"  ✗ 失败: {e}", 'error')
                    failed_files.append((Path(input_path).name, str(e)))

            self._update_progress(100, 100, '完成')

            if mode == 'analyze':
                if failed_files:
                    name, err = failed_files[0]
                    self.root.after(0, lambda n=name, e=err: messagebox.showerror(
                        "处理失败", f"{n}\n\n{e}"
                    ))
            elif failed_files:
                summary = "\n".join(f"  • {n}: {e}" for n, e in failed_files)
                self.log_panel.log(
                    f"\n完成: 成功 {len(success_paths)} 个 / 失败 {len(failed_files)} 个", 'warning'
                )
                self.root.after(0, lambda: messagebox.showwarning(
                    "部分文件处理失败",
                    f"成功 {len(success_paths)} 个，失败 {len(failed_files)} 个:\n{summary}"
                ))
            elif total == 1:
                fp = success_paths[0]
                self.root.after(0, lambda: self.result_panel.show_success(
                    "处理完成", Path(fp).name
                ))
                self.root.after(0, lambda: messagebox.showinfo(
                    "完成", f"文件已保存至:\n{fp}"
                ))
                if self.auto_open_var.get():
                    self.root.after(100, lambda p=fp: _open_file(p))
                self.log_panel.log("全部完成", 'success')
            else:
                out_dir = Path(output_base) if os.path.isdir(output_base) \
                          else Path(success_paths[0]).parent
                self.root.after(0, lambda: self.result_panel.show_success(
                    f"批量处理完成", f"成功处理 {total} 个文件"
                ))
                self.root.after(0, lambda: messagebox.showinfo(
                    "完成", f"成功处理 {total} 个文件\n输出目录:\n{out_dir}"
                ))
                self.log_panel.log(f"全部完成，共处理 {total} 个文件", 'success')

        except Exception as e:
            error_msg = str(e)
            self.log_panel.log(f"错误: {error_msg}", 'error')
            import traceback
            self.log_panel.log(traceback.format_exc(), 'error')
            self.root.after(0, lambda msg=error_msg: messagebox.showerror("错误", msg))

        finally:
            self.root.after(0, self._reset_btn)
            # v1.8.0: 清理粘贴文本生成的临时输入文件
            temp_input = getattr(self, '_pending_temp_input', None)
            if temp_input and os.path.exists(temp_input):
                try:
                    os.unlink(temp_input)
                except Exception:
                    pass
                self._pending_temp_input = None

    def _process_single_file(self, input_path, output_path, mode, progress_fn, revision_mode=False):
        """
        处理单个文件。
        progress_fn(pct: int, text: str) 由调用方传入，负责映射到全局进度条。
        返回 (实际输出路径, summary)（可能因回退而改变后缀）。
        不调用 _reset_btn，不显示完成 messagebox（由调用方统一处理）。
        """
        temp_docx = None
        temp_output_docx = None

        # 确定空格处理模式
        preset_name = self.preset.get() if hasattr(self, 'preset') else 'official'
        if preset_name == 'custom':
            try:
                _config = load_custom_settings()
                _cs = get_active_user_preset(_config)
                space_mode = _cs.get('space_handling', 'remove_all')
            except Exception:
                space_mode = 'remove_all'
        else:
            space_mode = 'remove_all'
        try:
            from docx import Document

            ext = Path(input_path).suffix.lower()
            if ext in ('.doc', '.wps'):
                progress_fn(0, f'转换 {ext} 为 .docx...')
                self.log_panel.log(f"检测到 {ext} 格式，正在转换...", 'info')
                from scripts.converter import convert_to_docx
                try:
                    temp_docx = convert_to_docx(input_path)
                except RuntimeError:
                    self.root.after(0, lambda: messagebox.showerror(
                        "转换失败", "未检测到 WPS 或 Microsoft Office，请先安装后再试。"
                    ))
                    raise
                input_path = temp_docx
                self.log_panel.log("转换成功", 'success')

            output_ext = Path(output_path).suffix.lower()
            needs_convert_back = output_ext in ('.doc', '.wps')
            if needs_convert_back:
                import tempfile
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                    temp_output_docx = tmp.name
                output_path_docx = temp_output_docx
            else:
                output_path_docx = output_path

            if mode == 'analyze':
                progress_fn(10, '正在诊断...')
                doc = Document(input_path)
                results = {
                    'punctuation': analyze_punctuation(doc),
                    'numbering': analyze_numbering(doc),
                    'paragraph': analyze_paragraph_format(doc),
                    'font': analyze_font(doc)
                }
                progress_fn(100, '诊断完成')
                self.root.after(0, lambda: self.result_panel.show_diagnosis(results))
                self.log_panel.log("诊断完成", 'success')

            elif mode == 'punctuation':
                progress_fn(10, '修复标点...')
                self._run_punctuation(input_path, output_path_docx, space_mode=space_mode)
                progress_fn(100, '完成')

            elif mode == 'smart':
                import tempfile
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                    temp_path = tmp.name
                self.log_panel.log("步骤 1/2: 修复标点...", 'info')
                progress_fn(0, '步骤 1/2: 修复标点...')
                self._run_punctuation(input_path, temp_path, quiet=True, space_mode=space_mode)
                self.log_panel.log("步骤 2/2: 应用格式...", 'info')
                progress_fn(5, '步骤 2/2: 应用格式...')

                def scaled_progress(pct, total, text):
                    progress_fn(5 + int(pct * 90 / 100), text)
                para_stats = self._run_format(temp_path, output_path_docx,
                                 progress_callback=scaled_progress,
                                 revision_mode=revision_mode)
                os.unlink(temp_path)

            if mode != 'analyze' and needs_convert_back:
                from scripts.converter import convert_from_docx
                try:
                    progress_fn(90, f'转换回 {output_ext} 格式...')
                    self.log_panel.log(f"正在转换回 {output_ext} 格式...", 'info')
                    actual_output = convert_from_docx(
                        output_path_docx, output_path,
                        format=output_ext.lstrip('.')
                    )
                    if actual_output and actual_output != output_path:
                        output_path = actual_output
                        self.log_panel.log(
                            f"保存 {output_ext} 需要 WPS Office，已自动保存为 .doc 格式", 'info'
                        )
                except RuntimeError as e:
                    if "未检测到" in str(e):
                        self.root.after(0, lambda: messagebox.showerror(
                            "转换失败", "未检测到 WPS 或 Microsoft Office，请先安装后再试。"
                        ))
                        raise
                    self._fallback_to_docx(output_path, output_path_docx)
                    output_path = str(Path(output_path).with_suffix('.docx'))
                except Exception as e:
                    self.log_panel.log(f"转换回 {output_ext} 失败: {e}", 'info')
                    self._fallback_to_docx(output_path, output_path_docx)
                    output_path = str(Path(output_path).with_suffix('.docx'))
                finally:
                    if os.path.exists(output_path_docx) and output_path_docx != output_path:
                        try:
                            os.unlink(output_path_docx)
                        except Exception:
                            pass

            return output_path, None

        finally:
            if temp_docx and os.path.exists(temp_docx):
                try:
                    os.unlink(temp_docx)
                except Exception:
                    pass
            if temp_output_docx and os.path.exists(temp_output_docx):
                try:
                    os.unlink(temp_output_docx)
                except Exception:
                    pass
    
    def _fallback_to_docx(self, original_output_path, docx_source_path):
        """转换回原格式失败时，将已处理好的 .docx 直接保存"""
        import shutil
        fallback_path = str(Path(original_output_path).with_suffix('.docx'))
        try:
            shutil.copy2(docx_source_path, fallback_path)
            self.log_panel.log(
                f"已回退保存为 .docx 格式: {Path(fallback_path).name}", 'info'
            )
        except Exception as e:
            self.log_panel.log(f"回退保存也失败: {e}", 'error')
    
    def _reset_btn(self):
        self.run_btn.configure(bg=Theme.PRIMARY)
        self.run_label.configure(bg=Theme.PRIMARY, text="开始处理")
        self._hide_progress()
    
    def _run_punctuation(self, input_path, output_path, quiet=False, space_mode='remove_all'):
        from docx import Document
        from scripts.punctuation import process_paragraph
        
        doc = Document(input_path)
        changes = 0
        
        for para in doc.paragraphs:
            if process_paragraph(para, space_mode=space_mode):
                changes += 1
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if process_paragraph(para, space_mode=space_mode):
                            changes += 1
        
        doc.save(output_path)
        if not quiet:
            self.log_panel.log(f"修复了 {changes} 处标点", 'success')
    
    def _run_format(self, input_path, output_path, progress_callback=None, revision_mode=False):
        preset_name = self.preset.get()
        
        # 设置 logging handler，让 formatter 的日志输出到日志面板
        import logging
        
        class LogPanelHandler(logging.Handler):
            def __init__(self, log_panel):
                super().__init__()
                self.log_panel = log_panel
            def emit(self, record):
                msg = self.format(record)
                self.log_panel.log(msg, 'info')
        
        formatter_logger = logging.getLogger('docformat.formatter')
        handler = LogPanelHandler(self.log_panel)
        handler.setFormatter(logging.Formatter('%(message)s'))
        formatter_logger.addHandler(handler)
        formatter_logger.setLevel(logging.INFO)
        
        try:
            cb = progress_callback if progress_callback is not None else self._update_progress
            bold_serial = True
            custom_settings = None
            if preset_name == 'custom':
                try:
                    _config = load_custom_settings()
                    _cs = get_active_user_preset(_config)
                    custom_settings = _cs
                    bold_serial = _cs.get('bold_serial', True)
                except Exception:
                    pass
            format_document(input_path, output_path, preset_name,
                           progress_callback=cb, revision_mode=revision_mode,
                           bold_serial=bold_serial, custom_settings=custom_settings)
        finally:
            formatter_logger.removeHandler(handler)
        
        if preset_name == 'custom':
            try:
                _config = load_custom_settings()
                custom = get_active_user_preset(_config)
                preset_label = custom.get('name', '自定义格式') if custom else '自定义格式'
            except Exception:
                preset_label = '自定义格式'
        else:
            preset_label = PRESETS.get(preset_name, {}).get('name', preset_name)
        self.log_panel.log(f"应用格式: {preset_label}", 'success')


def main():
    global _DND_AVAILABLE
    _enable_windows_high_dpi()
    
    if _DND_AVAILABLE:
        try:
            root = TkinterDnD.Tk()
        except Exception as e:
            # 某些打包环境下 python 模块可导入，但 tkdnd 运行库实际缺失。
            # 这里自动降级到普通 Tk，保证程序至少能启动使用。
            print(f"[警告] 拖拽运行库加载失败，已降级为普通模式: {e}")
            _DND_AVAILABLE = False
            root = tk.Tk()
    else:
        if getattr(sys, 'frozen', False) and sys.platform == 'darwin':
            print("[信息] macOS 打包版当前默认关闭拖拽功能，以优先保证应用可正常启动。")
        root = tk.Tk()
    _configure_tk_high_dpi(root)
    app = DocFormatApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
