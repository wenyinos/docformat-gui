import zipfile

from docx import Document

from scripts.ai_checker import annotate_document


def test_annotate_document_anchors_comment_to_quote(tmp_path):
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("第一句没有问题。")
    paragraph.add_run("第二句存在明显问题。")
    paragraph.add_run("第三句也没有问题。")

    stats = annotate_document(doc, [{
        "para": 1,
        "category": "病句",
        "quote": "第二句存在明显问题。",
        "message": "测试句子级批注",
        "suggestion": "",
    }])

    assert stats["annotated"] == 1
    assert stats["comment"] == 1
    assert stats["skipped"] == 0
    assert paragraph.text == "第一句没有问题。第二句存在明显问题。第三句也没有问题。"

    output_path = tmp_path / "quote_anchor.docx"
    doc.save(output_path)

    with zipfile.ZipFile(output_path) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")
        comments_xml = zf.read("word/comments.xml").decode("utf-8")

    assert "【AI病句】测试句子级批注" in comments_xml
    assert "<w:commentRangeStart" in document_xml
    assert "<w:t>第二句存在明显问题。</w:t>" in document_xml
