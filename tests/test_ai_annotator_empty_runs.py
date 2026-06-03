import zipfile

from docx import Document
from docx.oxml import OxmlElement

from scripts.ai_checker import annotate_document


def _add_hyperlink_text(paragraph, text):
    """构造 paragraph.text 非空、paragraph.runs 为空的特殊段落。

    python-docx 1.2.0 会把 w:hyperlink 内的文字计入 paragraph.text，
    但 paragraph.runs 不包含 hyperlink 内部的 w:r，因此可以稳定触发兜底分支。
    """
    hyperlink = OxmlElement("w:hyperlink")
    run = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def test_annotate_document_rebuilds_text_run_when_paragraph_runs_empty(tmp_path):
    text = "这是一个段落有文字但 paragraph.runs 为空的特殊段落"
    doc = Document()
    paragraph = doc.add_paragraph()
    _add_hyperlink_text(paragraph, text)

    assert paragraph.text.strip() == text
    assert paragraph.runs == []

    stats = annotate_document(doc, [{
        "para": 1,
        "category": "逻辑",
        "message": "测试空 runs 兜底",
        "suggestion": "请检查该段逻辑。",
    }])

    assert stats == {
        "annotated": 1,
        "skipped": 0,
        "mode": "inline",
        "comment": 0,
        "inline": 1,
    }
    assert paragraph.text.startswith(text)
    assert "【AI逻辑】测试空 runs 兜底" in paragraph.text

    output_path = tmp_path / "empty_runs_comment.docx"
    doc.save(output_path)

    reopened = Document(output_path)
    comments = list(reopened.comments)
    assert len(comments) == 0

    with zipfile.ZipFile(output_path) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")

    assert "word/comments.xml" not in zipfile.ZipFile(output_path).namelist()
    assert "<w:commentRangeStart" not in document_xml
    assert "【AI逻辑】测试空 runs 兜底" in document_xml
