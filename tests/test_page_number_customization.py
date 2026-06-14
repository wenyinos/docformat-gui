from copy import deepcopy
from pathlib import Path
import tempfile
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from scripts.formatter import PRESETS, add_page_number, format_document


def _footer_text(footer):
    return "".join(paragraph.text for paragraph in footer.paragraphs)


def test_official_margins_form_156_by_225_mm_text_area():
    page = PRESETS["official"]["page"]
    assert round(21.0 - page["left"] - page["right"], 1) == 15.6
    assert round(29.7 - page["top"] - page["bottom"], 1) == 22.5


def test_official_page_number_uses_text_area_offset_and_correct_outside_spaces():
    doc = Document()
    section = doc.sections[0]
    section.bottom_margin = Cm(3.5)

    add_page_number(
        doc,
        style="dash",
        position="outside",
        offset_from_text_mm=7,
    )

    assert abs(section.footer_distance.cm - 2.8) < 0.02
    odd = section.footer.paragraphs[0]
    even = section.even_page_footer.paragraphs[0]
    assert odd.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert even.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert _footer_text(section.footer).endswith(" —　")
    assert _footer_text(section.even_page_footer).startswith("　— ")


def test_page_number_styles_support_plain_text_and_total_pages():
    doc = Document()
    add_page_number(doc, style="page_text", position="center")
    assert "第 " in _footer_text(doc.sections[0].footer)
    assert " 页" in _footer_text(doc.sections[0].footer)

    add_page_number(
        doc,
        style="page_total",
        position="center",
        replace_existing=True,
    )
    xml = doc.sections[0].footer._element.xml
    assert "PAGE" in xml
    assert "NUMPAGES" in xml
    assert " / " in _footer_text(doc.sections[0].footer)


def test_custom_format_replaces_existing_page_number():
    with tempfile.TemporaryDirectory() as folder:
        source = Path(folder) / "source.docx"
        output = Path(folder) / "output.docx"
        doc = Document()
        doc.add_paragraph("测试正文。")
        add_page_number(doc, style="dash", position="outside")
        doc.save(source)

        custom = deepcopy(PRESETS["official"])
        custom.update({
            "name": "页码测试",
            "page_number_style": "page_text",
            "page_number_position": "center",
            "page_number_size": 12,
            "page_number_offset_mm": 10,
            "replace_existing_page_number": True,
        })
        format_document(
            str(source),
            str(output),
            preset_name="custom",
            custom_settings=custom,
        )

        result = Document(output)
        section = result.sections[0]
        assert section.footer.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert "第 " in _footer_text(section.footer)
        assert " 页" in _footer_text(section.footer)
        assert abs(section.footer_distance.cm - 2.5) < 0.02

        with zipfile.ZipFile(output) as archive:
            footer_xml = "".join(
                archive.read(name).decode("utf-8")
                for name in archive.namelist()
                if name.startswith("word/footer") and name.endswith(".xml")
            )
        assert footer_xml.count("第 ") == 1
        assert "— " not in footer_xml


def test_non_page_footer_content_is_not_overwritten():
    doc = Document()
    footer = doc.sections[0].footer
    footer.paragraphs[0].text = "内部资料"

    add_page_number(doc, style="plain", replace_existing=True)

    assert footer.paragraphs[0].text == "内部资料"
    assert "PAGE" not in footer._element.xml


def test_static_page_like_footer_can_be_replaced():
    doc = Document()
    footer = doc.sections[0].footer
    footer.paragraphs[0].text = "— 1 —"

    add_page_number(doc, style="page_text", position="center", replace_existing=True)

    assert "第 " in _footer_text(footer)
    assert "PAGE" in footer._element.xml
