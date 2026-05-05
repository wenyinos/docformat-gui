"""验证段落 style 重置和 Autospacing 清理（v1.7.2 修复）。"""
import sys
import zipfile
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))


def test_format_document_resets_style_to_normal():
    """处理后的所有段落 style 应该被重置为 Normal。"""
    from docx import Document
    from scripts.formatter import format_document

    doc = Document()
    doc.add_paragraph('一、第一项工作', style='Heading 3')
    doc.add_paragraph('正文段落。')
    doc.add_paragraph('（一）小标题', style='Heading 4')

    with tempfile.TemporaryDirectory() as tmp:
        ip = Path(tmp) / 'in.docx'
        op = Path(tmp) / 'out.docx'
        doc.save(ip)
        format_document(str(ip), str(op), preset_name='official')

        out = Document(op)
        for i, para in enumerate(out.paragraphs):
            if para.text.strip():
                assert para.style.name == 'Normal', \
                    f'段 {i} style 是 {para.style.name}，应为 Normal'


def test_format_document_strips_autospacing_from_styles():
    """处理后的 styles.xml 不应再有 Autospacing 属性。"""
    from docx import Document
    from scripts.formatter import format_document

    doc = Document()
    doc.add_paragraph('一、第一项工作', style='Heading 3')
    doc.add_paragraph('正文。')

    with tempfile.TemporaryDirectory() as tmp:
        ip = Path(tmp) / 'in.docx'
        op = Path(tmp) / 'out.docx'
        doc.save(ip)
        format_document(str(ip), str(op), preset_name='official')

        with zipfile.ZipFile(op, 'r') as z:
            styles_xml = z.read('word/styles.xml').decode('utf-8')
        assert 'Autospacing' not in styles_xml, \
            'styles.xml 仍含 Autospacing'
