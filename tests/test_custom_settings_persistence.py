from copy import deepcopy
from pathlib import Path

from docx import Document

import docformat_gui
from scripts import formatter
from scripts.formatter import PRESETS, format_document


def test_gui_saves_and_loads_active_custom_preset(tmp_path, monkeypatch):
    config_file = tmp_path / "custom_settings.json"
    monkeypatch.setattr(docformat_gui, "CONFIG_FILE", config_file)

    preset_a = deepcopy(docformat_gui.DEFAULT_CUSTOM_SETTINGS)
    preset_a.update({"id": "preset-a", "name": "A"})
    preset_a["page"]["left"] = 1.1

    preset_b = deepcopy(docformat_gui.DEFAULT_CUSTOM_SETTINGS)
    preset_b.update({"id": "preset-b", "name": "B"})
    preset_b["page"]["left"] = 2.2

    docformat_gui.save_custom_settings({
        "schema_version": docformat_gui.CONFIG_SCHEMA_VERSION,
        "active_preset_id": "preset-b",
        "presets": [preset_a, preset_b],
    })

    loaded = docformat_gui.load_custom_settings()
    active = docformat_gui.get_active_user_preset(loaded)

    assert loaded["active_preset_id"] == "preset-b"
    assert active["name"] == "B"
    assert active["page"]["left"] == 2.2


def test_formatter_reads_active_custom_preset_from_user_config(tmp_path, monkeypatch):
    appdata = tmp_path / "appdata"
    config_dir = appdata / "DocFormatter"
    config_dir.mkdir(parents=True)
    monkeypatch.setenv("APPDATA", str(appdata))
    monkeypatch.setattr(formatter.sys, "platform", "win32")

    preset = deepcopy(PRESETS["official"])
    preset.update({"id": "custom-active", "name": "用户保存的自定义格式"})
    preset["page"] = {"top": 1.1, "bottom": 1.2, "left": 1.3, "right": 1.4}

    monkeypatch.setattr(docformat_gui, "CONFIG_FILE", config_dir / "custom_settings.json")
    docformat_gui.save_custom_settings({
        "schema_version": docformat_gui.CONFIG_SCHEMA_VERSION,
        "active_preset_id": "custom-active",
        "presets": [preset],
    })

    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("这是正文内容。")
    doc.save(source)

    format_document(str(source), str(output), preset_name="custom")

    result = Document(output)
    section = result.sections[0]
    assert abs(section.top_margin.cm - 1.1) < 0.02
    assert abs(section.bottom_margin.cm - 1.2) < 0.02
    assert abs(section.left_margin.cm - 1.3) < 0.02
    assert abs(section.right_margin.cm - 1.4) < 0.02
