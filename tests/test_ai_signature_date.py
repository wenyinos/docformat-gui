from datetime import date

from docx import Document

from scripts import ai_checker


def test_signature_date_before_system_date_is_reported(monkeypatch):
    monkeypatch.setattr(ai_checker, "_get_beijing_today", lambda: date(2026, 6, 2))

    doc = Document()
    doc.add_paragraph("关于测试落款日期的通知")
    doc.add_paragraph("正文内容。")
    doc.add_paragraph("示例单位")
    doc.add_paragraph("5月30日")

    issues = ai_checker._check_signature_date(doc)

    assert len(issues) == 1
    assert issues[0]["para"] == 4
    assert issues[0]["quote"] == "5月30日"
    assert "早于当前系统日期" in issues[0]["message"]


def test_signature_date_equal_system_date_is_not_reported(monkeypatch):
    monkeypatch.setattr(ai_checker, "_get_beijing_today", lambda: date(2026, 6, 2))

    doc = Document()
    doc.add_paragraph("示例单位")
    doc.add_paragraph("2026年6月2日")

    assert ai_checker._check_signature_date(doc) == []
