#!/usr/bin/env python3
"""AI 逻辑/病句检查模块。"""

import json
import logging
import re
import sys
from copy import deepcopy
from datetime import datetime, timezone, timedelta
from pathlib import Path
from pprint import pprint
from urllib import error, request

from docx import Document
from docx.shared import RGBColor
from docx.text.run import Run

try:
    from scripts import ai_config
except ImportError:
    import ai_config


logger = logging.getLogger("docformat.ai_checker")

CATEGORIES = {"病句", "逻辑", "其他"}
COMMENT_AUTHOR = "AI校对助手"


def _get_beijing_today():
    """从本机系统时间获取北京时间日期对象，避免额外联网取时。"""
    beijing_tz = timezone(timedelta(hours=8))
    return datetime.now(beijing_tz).date()


def _get_beijing_date_text():
    """从本机系统时间获取北京时间日期文本，避免额外联网取时。"""
    return _get_beijing_today().strftime("%Y年%m月%d日")


_CN_DIGITS = {
    "〇": 0, "○": 0, "零": 0,
    "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
    "六": 6, "七": 7, "八": 8, "九": 9,
}


def _parse_cn_year(text):
    """解析二〇二六 / 二零二六这类年份。"""
    digits = []
    for char in text:
        if char in _CN_DIGITS:
            digits.append(str(_CN_DIGITS[char]))
        elif char.isdigit():
            digits.append(char)
        else:
            return None
    if len(digits) == 4:
        return int("".join(digits))
    return None


def _parse_cn_number(text):
    """解析一到三十一之间的中文数字，用于月份和日期。"""
    text = text.strip()
    if text.isdigit():
        return int(text)
    if text in _CN_DIGITS:
        return _CN_DIGITS[text]
    if "十" in text:
        left, right = text.split("十", 1)
        tens = 1 if not left else _CN_DIGITS.get(left)
        ones = 0 if not right else _CN_DIGITS.get(right)
        if tens is None or ones is None:
            return None
        return tens * 10 + ones
    return None


def _parse_date_from_text(text, default_year=None):
    """从文本中解析一个日期，支持常见阿拉伯数字和中文日期。"""
    from datetime import date

    patterns = [
        r"(?P<year>\d{4})年(?P<month>\d{1,2})月(?P<day>\d{1,2})[日号]?",
        r"(?P<year>\d{4})[./-](?P<month>\d{1,2})[./-](?P<day>\d{1,2})",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            try:
                return date(
                    int(match.group("year")),
                    int(match.group("month")),
                    int(match.group("day")),
                )
            except ValueError:
                return None

    match = re.search(r"(?P<year>[〇○零一二三四五六七八九\d]{4})年(?P<month>[一二三四五六七八九十\d]{1,3})月(?P<day>[一二三四五六七八九十\d]{1,3})[日号]?", text)
    if match:
        year = _parse_cn_year(match.group("year"))
        month = _parse_cn_number(match.group("month"))
        day = _parse_cn_number(match.group("day"))
        if year and month and day:
            try:
                return date(year, month, day)
            except ValueError:
                return None

    if default_year:
        match = re.search(r"(?P<month>\d{1,2})月(?P<day>\d{1,2})[日号]?", text)
        if match:
            try:
                return date(default_year, int(match.group("month")), int(match.group("day")))
            except ValueError:
                return None

    return None


def _is_likely_signature_date(text):
    """判断一行文本是否像落款日期。

    只对文末短行做本地兜底检查，避免把正文中历史日期误判成落款。
    """
    compact = re.sub(r"\s+", "", text)
    if not compact or len(compact) > 32:
        return False
    return bool(re.search(
        r"(\d{4}年\d{1,2}月\d{1,2}[日号]?|\d{4}[./-]\d{1,2}[./-]\d{1,2}|[〇○零一二三四五六七八九\d]{4}年[一二三四五六七八九十\d]{1,3}月[一二三四五六七八九十\d]{1,3}[日号]?|\d{1,2}月\d{1,2}[日号]?)",
        compact,
    ))


def _check_signature_date(doc):
    """本地检查落款日期是否与当前系统日期一致。

    这项不依赖模型，确保类似“系统日期 6 月 2 日、落款 5 月 30 日”的情况稳定提示。
    """
    today = _get_beijing_today()
    non_empty = [(idx + 1, para.text.strip()) for idx, para in enumerate(doc.paragraphs) if para.text.strip()]
    for para_index, text in reversed(non_empty[-10:]):
        if not _is_likely_signature_date(text):
            continue
        parsed_date = _parse_date_from_text(text, default_year=today.year)
        if not parsed_date or parsed_date == today:
            return []

        if parsed_date < today:
            relation = "早于"
            suggestion = f"建议核对是否应更新为 {today.strftime('%Y年%m月%d日')}。"
        else:
            relation = "晚于"
            suggestion = "建议核对落款日期是否误写为未来日期。"

        return [{
            "para": para_index,
            "category": "逻辑",
            "quote": text,
            "message": f"落款日期{relation}当前系统日期（当前北京时间为 {_get_beijing_date_text()}），请核对。",
            "suggestion": suggestion,
        }]
    return []


def _report_progress(progress_fn, percent, message):
    """统一上报进度，便于以后保持 GUI 和命令行流程一致。"""
    if progress_fn:
        progress_fn(percent, message)


def _fake_check_document(doc, progress_fn=None):
    """假实现：不联网，只返回占位问题。

    此处保留为测试入口。第二步接入真实 API 后，仍可通过 ai_config.USE_FAKE=True
    强制走这里，方便没有 API Key 或网络不可用时验证后续流程。
    """
    issues = _check_signature_date(doc)
    paragraphs = doc.paragraphs
    total = len(paragraphs)

    if total == 0:
        _report_progress(progress_fn, 100, "AI 检查完成：文档没有段落")
        return {"issues": issues, "total_chunks": 0, "failed_chunks": 0, "used_fake": True}

    _report_progress(progress_fn, 0, "开始 AI 逻辑/病句检查（假实现）")

    for index, para in enumerate(paragraphs, start=1):
        text = para.text.strip()

        if text and len(text) > 30:
            # 此处为假实现，真实 API 调用见 _call_api。
            # 这里不修改正文，只返回统一格式的问题数据，方便后续流程插入批注或展示结果。
            issues.append({
                "para": index,
                "category": "逻辑",
                "quote": text,
                "message": "这是一条用于测试的占位问题，正文不会被修改",
                "suggestion": "",
            })

        percent = int(index * 100 / total)
        _report_progress(progress_fn, percent, f"正在检查第 {index}/{total} 段")

    _report_progress(progress_fn, 100, f"AI 检查完成，共发现 {len(issues)} 条占位问题")
    return {"issues": issues, "total_chunks": 1 if total else 0, "failed_chunks": 0, "used_fake": True}


def _build_chunks(paragraphs):
    """按最大字符数把非空段落切成块。

    每个块是一个 list，元素为 (para_index, para_text)。
    para_index 使用文档真实段落索引，从 1 开始；空段落不进入块。
    如果单个段落超过 MAX_CHARS_PER_CHUNK，该段落会单独成块，不切断段落。
    """
    chunks = []
    current_chunk = []
    current_chars = 0
    max_chars = max(1, ai_config.MAX_CHARS_PER_CHUNK)

    for index, para in enumerate(paragraphs, start=1):
        text = para.text.strip()
        if not text:
            continue

        text_chars = len(text)
        if current_chunk and current_chars + text_chars > max_chars:
            chunks.append(current_chunk)
            current_chunk = []
            current_chars = 0

        current_chunk.append((index, text))
        current_chars += text_chars

        if text_chars > max_chars:
            chunks.append(current_chunk)
            current_chunk = []
            current_chars = 0

    if current_chunk:
        chunks.append(current_chunk)

    return chunks


def _call_api(chunk, glossary=""):
    """调用 OpenAI 兼容的 chat/completions 接口，返回模型消息正文。

    当前请求格式为 OpenAI 兼容格式：
    - POST ai_config.API_BASE_URL
    - Header: Authorization: Bearer <API_KEY>
    - JSON Body: {"model": MODEL_NAME, "messages": [...]}

    将来切换供应商时，通常只需要改 ai_config.py 中的 API_BASE_URL / MODEL_NAME，
    如果供应商的请求体或响应结构不兼容，再集中调整本函数的 payload 和响应解析部分。
    """
    current_date = _get_beijing_date_text()
    system_prompt = (
        "你是一名公文校对助手，只做句子层面的校对。\n"
        f"当前北京时间日期：{current_date}。判断“今年、去年、近期、已发生、尚未发生”等时间关系时，以该日期为准；"
        "不要因为模型知识库截止时间较早，就否定文本中明确说明的近期或当前事项。\n"
        "默认语境：除非文本明确表明是外部单位来文、代拟外部材料或跨公司/跨部门联合文本，"
        "否则按本公司从自身角度发文理解。\n"
        "检查范围仅限：明显的病句、语法错误、确定性较高的错别字；相邻语句之间明显的事实或逻辑矛盾"
        "（例如前后数字、日期、称谓不一致，或指代不清）。\n"
        "日期类问题也属于本次检查范围：落款日期不得早于当前北京时间日期；"
        "同一事项中的日期顺序不得倒置；不得出现明显不合理的提前开展、提前完成、提前报送等时间关系。"
        "如果文中明确说明是在回顾历史、补发旧文或引用既有材料，不要仅因日期早于当前日期就判错。\n"
        "明确不要做以下事情：不要评价文章的写作目的、立场或态度；不要评价是否得体；"
        "不要评价整体谋篇布局、详略安排或说服力；不要做风格润色；不要挑无关紧要的小毛病。\n"
        "如果某处问题需要联系全文意图或上级精神才能判断，不属于本次检查范围，请跳过。\n"
        "错别字检查只提示确定性较高的错误，尤其是专有名词、部门名称、人员姓名、职务称谓中的明显错写；"
        "不要把常见同义表达、简称、语序差异或正常公文用语当作错别字。\n"
        "对每个问题给出它属于哪一段（必须使用用户提供的段落编号）、类别、问题描述、可选修改建议。"
        "类别只能是：病句、逻辑、其他。quote 字段请填写原文中出问题的完整短句或最小连续片段，"
        "必须逐字摘自原文，不要改写；如果无法确定，可留空。\n"
        "请严格返回 JSON 数组，每个元素形如 "
        '{"para": 段落编号, "category": "病句", "quote": "原文短句", "message": "...", "suggestion": "..."}。'
        "不要返回 JSON 以外的任何内容，包括解释、markdown 代码块标记。如果没有发现问题，返回空数组 []。"
    )
    if glossary:
        system_prompt += (
            "\n以下是本单位已知正确的专有名词、简称、惯用表述和单位情况，"
            "请勿将它们判定为错误或建议修改。"
            "默认情况下，当前文档视为本公司发文，应优先参考术语表校对本公司的部门名称、人员姓名和职务称谓。"
            "如果术语表写明某单位有全称和简称，通常首次出现可用全称，后文可用简称；"
            "不要主动建议把全文简称全部改成全称，除非原文确实造成指代不清：\n"
            "如果正文明确是在给其他公司、事业单位或外部机构发函，且相关部门或人名显然属于对方单位，"
            "请谨慎处理，不要强行套用本单位术语表；如果正文涉及跨公司、跨部门联合事项，"
            "只对能够判断为本公司内部的名称使用术语表校对。\n"
            f"{glossary}"
        )
    user_content = "\n\n".join(f"【第{para_index}段】{text}" for para_index, text in chunk)
    payload = {
        "model": ai_config.MODEL_NAME,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content},
        ],
    }
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    req = request.Request(
        ai_config.API_BASE_URL,
        data=body,
        headers={
            "Authorization": f"Bearer {ai_config.API_KEY}",
            "Content-Type": "application/json",
        },
        method="POST",
    )

    with request.urlopen(req, timeout=ai_config.REQUEST_TIMEOUT) as resp:
        resp_text = resp.read().decode("utf-8", errors="replace")

    data = json.loads(resp_text)
    choices = data.get("choices") or []
    if not choices:
        logger.warning("AI API 响应中没有 choices")
        return ""

    message = choices[0].get("message") or {}
    return message.get("content", "")


def _parse_response(raw_text, valid_para_indexes=None):
    """解析模型返回的 JSON 数组，并过滤不合法结果。

    解析策略：
    1. 先直接 json.loads；
    2. 失败后截取第一个 [ 到最后一个 ] 之间的内容再解析；
    3. 仍失败则记录原因并返回 None，不抛异常中断整体流程。

    返回 None 表示本分段解析失败；返回 [] 表示模型给出了合法空数组。
    """
    if not raw_text:
        logger.warning("AI API 返回内容为空")
        return None

    try:
        parsed = json.loads(raw_text)
    except json.JSONDecodeError as exc:
        match = re.search(r"\[.*\]", raw_text, re.S)
        if not match:
            logger.warning("AI API 返回内容不是 JSON 数组: %s", exc)
            return None
        try:
            parsed = json.loads(match.group(0))
        except json.JSONDecodeError as inner_exc:
            logger.warning("AI API 返回内容 JSON 解析失败: %s", inner_exc)
            return None

    if not isinstance(parsed, list):
        logger.warning("AI API 返回 JSON 不是数组")
        return None

    valid_para_indexes = set(valid_para_indexes or [])
    issues = []
    for item in parsed:
        if not isinstance(item, dict):
            continue

        para = item.get("para")
        if isinstance(para, str) and para.isdigit():
            para = int(para)
        if not isinstance(para, int):
            continue
        if valid_para_indexes and para not in valid_para_indexes:
            continue

        category = item.get("category", "其他")
        if category not in CATEGORIES:
            category = "其他"

        message = str(item.get("message", "")).strip()
        if not message:
            continue

        suggestion = str(item.get("suggestion", "")).strip()
        quote = str(item.get("quote", "")).strip()
        issues.append({
            "para": para,
            "category": category,
            "quote": quote,
            "message": message,
            "suggestion": suggestion,
        })

    return issues


def check_document(doc, progress_fn=None):
    """检查文档中的逻辑/病句问题。

    Args:
        doc: python-docx 的 Document 对象。
        progress_fn: 可选进度回调，形如 progress_fn(百分比整数, '提示文字')。

    Returns:
        dict，字段如下：
        - issues: 问题列表。每个问题均为 dict，字段如下：
        - para: 段落索引，从 1 开始，与现有 analyzer 模块保持一致。
        - category: 问题类别，取值为 '病句'、'逻辑' 或 '其他'。
        - quote: 原文中出问题的短句或最小连续片段，可为空。
        - message: 给用户看的中文问题描述。
        - suggestion: 可选修改建议；没有建议时为空字符串。
        - total_chunks: 实际分段数。
        - failed_chunks: 调用或解析失败的分段数。
        - used_fake: 是否使用假实现。
    """
    if ai_config.USE_FAKE or not ai_config.API_KEY:
        return _fake_check_document(doc, progress_fn=progress_fn)

    glossary = ai_config.load_glossary()
    chunks = _build_chunks(doc.paragraphs)
    if not chunks:
        _report_progress(progress_fn, 100, "AI 检查完成：文档没有非空段落")
        return {"issues": [], "total_chunks": 0, "failed_chunks": 0, "used_fake": False}

    issues = _check_signature_date(doc)
    total_chunks = len(chunks)
    failed_chunks = 0
    _report_progress(progress_fn, 0, f"开始 AI 逻辑/病句检查，共 {total_chunks} 个分段")

    for chunk_index, chunk in enumerate(chunks, start=1):
        try:
            raw_text = _call_api(chunk, glossary=glossary)
            valid_para_indexes = {para_index for para_index, _ in chunk}
            parsed_issues = _parse_response(raw_text, valid_para_indexes=valid_para_indexes)
            if parsed_issues is None:
                failed_chunks += 1
            else:
                issues.extend(parsed_issues)
        except (TimeoutError, OSError, error.URLError, error.HTTPError, json.JSONDecodeError) as exc:
            logger.warning("AI 检查第 %s/%s 个分段失败: %s", chunk_index, total_chunks, exc)
            failed_chunks += 1
        except Exception as exc:
            logger.warning("AI 检查第 %s/%s 个分段发生未知错误: %s", chunk_index, total_chunks, exc)
            failed_chunks += 1

        percent = int(chunk_index * 100 / total_chunks)
        _report_progress(progress_fn, percent, f"已完成第 {chunk_index}/{total_chunks} 个分段")

    _report_progress(progress_fn, 100, f"AI 检查完成，共发现 {len(issues)} 条问题")
    return {
        "issues": issues,
        "total_chunks": total_chunks,
        "failed_chunks": failed_chunks,
        "used_fake": False,
    }


def _comment_api_available(doc):
    """判断当前 python-docx 是否支持原生 Word 批注 API。

    python-docx 1.2.0 提供 Document.add_comment(runs, text, author, initials)。
    如果用户环境降到旧版，缺少这个 API，则不要手写底层 XML 伪造批注，
    以免生成 Word 无法打开的损坏 docx，改用内联红色提示降级。
    """
    return hasattr(doc, "add_comment")


def _format_issue_comment(issue):
    """把问题 dict 转成写入 Word 批注或内联提示的文本。"""
    category = issue.get("category", "其他")
    if category not in CATEGORIES:
        category = "其他"

    message = str(issue.get("message", "")).strip()
    suggestion = str(issue.get("suggestion", "")).strip()
    comment_text = f"【AI{category}】{message}"
    if suggestion:
        comment_text += f"\n建议：{suggestion}"
    return comment_text


def _add_inline_hint(paragraph, text):
    """降级方案：在段落末尾追加红色内联提示。

    这会改变正文内容，只在当前 python-docx 不支持原生批注 API 时使用。
    正常环境下优先使用 Word 原生批注，做到正文一个字不改。
    """
    run = paragraph.add_run(f" {text}")
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def _split_run_at(run, offset):
    """把 run 在 offset 处分成左右两个 run，并保留原 run 的格式。

    只改变 Word XML 的 run 边界，不改变可见文字。返回 (left_run, right_run)，
    offset 在边界时其中一侧为 None。
    """
    text = run.text
    if offset <= 0:
        return None, run
    if offset >= len(text):
        return run, None

    right_r = deepcopy(run._r)
    run._r.addnext(right_r)
    right_run = Run(right_r, run._parent)

    run.text = text[:offset]
    right_run.text = text[offset:]
    return run, right_run


def _runs_between(paragraph, start_run, end_run):
    """返回 paragraph 中从 start_run 到 end_run 的连续 runs。"""
    selected = []
    collecting = False
    for run in paragraph.runs:
        if run._r is start_run._r:
            collecting = True
        if collecting and run.text:
            selected.append(run)
        if run._r is end_run._r:
            break
    return selected


def _find_quote_anchor_runs(paragraph, quote):
    """优先把批注锚定到 quote 对应的短句 runs。

    quote 必须是段落原文中的连续片段。必要时会拆分 run 边界，但不改变可见文字；
    定位不到时返回空列表，由调用方回退到整段批注。
    """
    quote = str(quote or "").strip()
    if not quote:
        return []

    runs = [run for run in paragraph.runs if run.text]
    if not runs:
        return []

    full_text = "".join(run.text for run in runs)
    start = full_text.find(quote)
    if start < 0:
        return []
    end = start + len(quote)

    spans = []
    cursor = 0
    for run in runs:
        run_start = cursor
        run_end = cursor + len(run.text)
        spans.append((run, run_start, run_end))
        cursor = run_end

    start_info = None
    end_info = None
    for run, run_start, run_end in spans:
        if start_info is None and run_start <= start < run_end:
            start_info = (run, start - run_start)
        if run_start < end <= run_end:
            end_info = (run, end - run_start)
            break

    if not start_info or not end_info:
        return []

    start_run, start_offset = start_info
    end_run, end_offset = end_info

    if start_run._r is end_run._r:
        left, _after = _split_run_at(start_run, end_offset)
        if left is None:
            return []
        _before, target = _split_run_at(left, start_offset)
        return [target] if target and target.text else []

    end_left, _after = _split_run_at(end_run, end_offset)
    _before, start_right = _split_run_at(start_run, start_offset)
    if start_right is None or end_left is None:
        return []
    return _runs_between(paragraph, start_right, end_left)


def _get_comment_anchor_runs(paragraph):
    """获取 Word 批注锚点 runs。

    正常情况下用段落中已有的有文本 run 作为锚点，批注会标在正文对应段落上。
    少数 docx 的 paragraph.text 有内容但 paragraph.runs 为空或没有文本 run 时不在这里
    重建正文，避免破坏原有格式；调用方会对这类段落走内联提示降级。
    只有段落本身没有任何文字时才返回空列表，由调用方跳过。
    """
    runs = [run for run in paragraph.runs if run.text]
    if runs:
        return runs
    return []


def annotate_document(doc, issues):
    """把 AI 问题写回传入的 Document 对象。

    Args:
        doc: python-docx 的 Document 对象。函数只修改这个内存对象，不保存、不覆盖原文件。
        issues: check_document 返回的问题列表，每条包含 para/category/quote/message/suggestion。

    Returns:
        统计 dict：
        - annotated: 实际插入条数
        - skipped: 因段落越界、内容缺失、批注锚点缺失等原因跳过的条数
        - mode: 'comment' 表示全部使用 Word 原生批注；'inline' 表示全部或部分使用内联红字
        - comment: 使用 Word 原生批注的条数
        - inline: 使用内联红字提示的条数
    """
    use_comment_api = _comment_api_available(doc)
    annotated = 0
    skipped = 0
    comment_count = 0
    inline_count = 0

    for issue in issues:
        para = issue.get("para")
        if isinstance(para, str) and para.isdigit():
            para = int(para)
        if not isinstance(para, int) or para < 1 or para > len(doc.paragraphs):
            skipped += 1
            continue

        comment_text = _format_issue_comment(issue)
        if not comment_text.strip() or "】" == comment_text[-1:]:
            skipped += 1
            continue

        paragraph = doc.paragraphs[para - 1]

        if use_comment_api:
            # 原生方案：python-docx 1.2.0 的 Document.add_comment。
            # 优先锚定到 quote 对应的短句 runs，定位不到再锚定到该段落的全部 runs；
            # 正文文本不增加、不删除、不替换。
            runs = _find_quote_anchor_runs(paragraph, issue.get("quote", ""))
            if not runs:
                runs = _get_comment_anchor_runs(paragraph)
            if not runs:
                if paragraph.text.strip():
                    # 安全兜底：段落有文字但没有可锚定 run 时，不清空、不重建正文，
                    # 只在段末追加红色提示，避免破坏原文格式。
                    _add_inline_hint(paragraph, comment_text)
                    annotated += 1
                    inline_count += 1
                else:
                    skipped += 1
                continue
            try:
                doc.add_comment(runs, text=comment_text, author=COMMENT_AUTHOR, initials="AI")
                annotated += 1
                comment_count += 1
            except Exception as exc:
                logger.warning("第 %s 段插入 Word 批注失败: %s", para, exc)
                skipped += 1
        else:
            # 降级方案：旧版 python-docx 没有批注 API 时，不手写底层 XML，
            # 而是在段末追加醒目的红色提示，保证生成的 docx 能正常打开。
            _add_inline_hint(paragraph, comment_text)
            annotated += 1
            inline_count += 1

    mode = "comment" if inline_count == 0 else "inline"
    return {
        "annotated": annotated,
        "skipped": skipped,
        "mode": mode,
        "comment": comment_count,
        "inline": inline_count,
    }


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("用法: python scripts/ai_checker.py <待检查.docx>")
        sys.exit(1)

    docx_path = sys.argv[1]
    document = Document(docx_path)
    mode = "假实现" if ai_config.USE_FAKE or not ai_config.API_KEY else "真实API"
    print(f"本次走的是{mode}")
    print(f"术语表字符数：{len(ai_config.load_glossary())}")

    def print_progress(percent, message):
        print(f"[{percent:3d}%] {message}")

    result = check_document(document, progress_fn=print_progress)
    issues = result["issues"]
    pprint(issues, width=120)
    print(f"运行信息：total_chunks={result['total_chunks']}, failed_chunks={result['failed_chunks']}, used_fake={result['used_fake']}")

    stats = annotate_document(document, issues)
    input_path = Path(docx_path)
    output_path = input_path.with_name(f"{input_path.stem}_AI批注{input_path.suffix}")
    document.save(output_path)
    print(f"批注写入统计：{stats}")
    print(f"已另存为：{output_path}")
